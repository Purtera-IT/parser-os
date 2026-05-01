from __future__ import annotations

import csv
import json
import re
import zipfile
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

from app.core.ids import stable_id
from app.core.normalizers import (
    detect_speaker,
    normalize_text,
    normalize_transcript_text,
    split_transcript_segments,
)
from app.core.schemas import ArtifactType, SourceRef
from app.core.segments import ArtifactSegment, make_segment

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

_EMAIL_BOUNDARY_RE = re.compile(r"^(On .+ wrote:|-----Original Message-----)$", flags=re.IGNORECASE)
_QUOTE_HEADER_ALIASES = {
    "part_number": {"part", "part number", "sku", "item number"},
    "description": {"description", "item", "product", "device"},
    "quantity": {"qty", "quantity", "count"},
    "unit_price": {"unit price", "price", "cost"},
    "lead_time": {"lead time", "eta"},
}
_XLSX_HEADER_ALIASES = {
    "site": {"site", "location", "facility", "store", "building"},
    "device": {"device", "asset", "equipment", "camera", "ap", "reader", "item"},
    "quantity": {"qty", "qty.", "quantity", "count", "#"},
    "floor": {"floor", "level"},
    "room": {"room", "area", "zone"},
    "scope": {"scope", "included", "work type"},
    "access": {"access", "access window", "hours", "site access"},
}


def _source_ref(
    *,
    artifact_id: str,
    artifact_type: ArtifactType,
    filename: str,
    locator: dict[str, Any],
    method: str,
    parser_version: str,
) -> SourceRef:
    return SourceRef(
        id=stable_id("src", artifact_id, artifact_type.value, locator),
        artifact_id=artifact_id,
        artifact_type=artifact_type,
        filename=filename,
        locator=locator,
        extraction_method=method,
        parser_version=parser_version,
    )


def _detect_header(rows: list[list[Any]], aliases: dict[str, set[str]]) -> tuple[int | None, dict[str, int]]:
    scan_limit = min(25, len(rows))
    best_idx: int | None = None
    best_map: dict[str, int] = {}
    best_score = -1
    for idx in range(scan_limit):
        row = rows[idx]
        current_map: dict[str, int] = {}
        for col_idx, cell in enumerate(row):
            cell_text = normalize_text(str(cell or "")).strip(".:")
            if not cell_text:
                continue
            for canonical, alias_values in aliases.items():
                if cell_text in alias_values and canonical not in current_map:
                    current_map[canonical] = col_idx
        if len(current_map) > best_score:
            best_score = len(current_map)
            best_idx = idx
            best_map = current_map
    if best_score <= 0:
        return None, {}
    return best_idx, best_map


def segment_xlsx(
    *,
    project_id: str,
    artifact_id: str,
    path: Path,
    parser_version: str = "segmenter_xlsx_v1",
) -> list[ArtifactSegment]:
    suffix = path.suffix.lower()
    rows_by_sheet: list[tuple[str, list[list[Any]], ArtifactType]] = []
    if suffix == ".xlsx":
        workbook = load_workbook(path, read_only=True, data_only=True)
        for sheet in workbook.worksheets:
            rows_by_sheet.append((sheet.title, [list(row) for row in sheet.iter_rows(values_only=True)], ArtifactType.xlsx))
    elif suffix == ".csv":
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            rows_by_sheet.append(("csv", [list(row) for row in csv.reader(handle)], ArtifactType.csv))
    else:
        return []

    segments: list[ArtifactSegment] = []
    for sheet_name, rows, artifact_type in rows_by_sheet:
        header_idx, header_map = _detect_header(rows, _XLSX_HEADER_ALIASES)
        if header_idx is None:
            continue
        columns = {key: get_column_letter(index + 1) for key, index in header_map.items()}
        for row_idx in range(header_idx + 1, len(rows)):
            row = rows[row_idx]
            if all(str(cell or "").strip() == "" for cell in row):
                continue
            locator = {"sheet": sheet_name, "row": row_idx + 1, "columns": columns}
            source_ref = _source_ref(
                artifact_id=artifact_id,
                artifact_type=artifact_type,
                filename=path.name,
                locator=locator,
                method="segment_spreadsheet_rows",
                parser_version=parser_version,
            )
            row_values = {key: str(row[col]).strip() if col < len(row) and row[col] is not None else "" for key, col in header_map.items()}
            text = " | ".join(f"{key}={value}" for key, value in row_values.items() if value)
            if not text:
                continue
            segments.append(
                make_segment(
                    project_id=project_id,
                    artifact_id=artifact_id,
                    artifact_type=artifact_type,
                    segment_type="spreadsheet_row",
                    text=text,
                    locator=locator,
                    source_ref=source_ref,
                    metadata={"sheet_name": sheet_name, "header_map": columns, "row_values": row_values},
                )
            )
    return segments


def _extract_email_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".eml":
        raw = path.read_bytes()
        msg = BytesParser(policy=policy.default).parsebytes(raw)
        body = msg.get_body(preferencelist=("plain", "html"))
        content = body.get_content() if body is not None else raw.decode("utf-8", errors="ignore")
    else:
        content = path.read_text(encoding="utf-8", errors="ignore")
    if "<html" in content.lower():
        soup = BeautifulSoup(content, "html.parser")
        return soup.get_text(separator="\n", strip=True)
    return content


def segment_email(
    *,
    project_id: str,
    artifact_id: str,
    path: Path,
    parser_version: str = "segmenter_email_v1",
) -> list[ArtifactSegment]:
    text = _extract_email_text(path)
    lines = text.splitlines()
    if not lines:
        return []
    blocks: list[list[tuple[int, str]]] = []
    current: list[tuple[int, str]] = []
    for idx, line in enumerate(lines, start=1):
        stripped = line.strip()
        is_boundary = bool(_EMAIL_BOUNDARY_RE.match(stripped))
        is_from_after_body = (
            stripped.lower().startswith("from:")
            and current
            and any(not entry.strip().lower().startswith(("from:", "sent:", "date:", "subject:")) for _, entry in current)
        )
        if current and (is_boundary or is_from_after_body):
            blocks.append(current)
            current = []
        current.append((idx, line))
    if current:
        blocks.append(current)

    segments: list[ArtifactSegment] = []
    artifact_type = ArtifactType.email if path.suffix.lower() == ".eml" else ArtifactType.txt
    for message_index, block in enumerate(blocks):
        stripped_lines = [value.strip() for _, value in block]
        quoted = any(row.startswith(">") for row in stripped_lines) or message_index > 0
        locator = {"message_index": message_index, "line_start": block[0][0], "line_end": block[-1][0], "quoted": quoted}
        source_ref = _source_ref(
            artifact_id=artifact_id,
            artifact_type=artifact_type,
            filename=path.name,
            locator=locator,
            method="segment_email_messages",
            parser_version=parser_version,
        )
        message_text = "\n".join(stripped_lines).strip()
        if message_text:
            segments.append(
                make_segment(
                    project_id=project_id,
                    artifact_id=artifact_id,
                    artifact_type=artifact_type,
                    segment_type="email_message",
                    text=message_text,
                    locator=locator,
                    source_ref=source_ref,
                    metadata={"quoted": quoted},
                )
            )
        for line_index, line in block:
            cleaned = line.strip()
            if not cleaned:
                continue
            line_locator = {"message_index": message_index, "line_start": line_index, "line_end": line_index, "quoted": cleaned.startswith(">")}
            line_ref = _source_ref(
                artifact_id=artifact_id,
                artifact_type=artifact_type,
                filename=path.name,
                locator=line_locator,
                method="segment_email_lines",
                parser_version=parser_version,
            )
            segments.append(
                make_segment(
                    project_id=project_id,
                    artifact_id=artifact_id,
                    artifact_type=artifact_type,
                    segment_type="email_line",
                    text=cleaned,
                    locator=line_locator,
                    source_ref=line_ref,
                    metadata={"message_index": message_index},
                )
            )
    return segments


def segment_docx(
    *,
    project_id: str,
    artifact_id: str,
    path: Path,
    parser_version: str = "segmenter_docx_v1",
) -> list[ArtifactSegment]:
    document = Document(path)
    segments: list[ArtifactSegment] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        locator = {"paragraph_index": paragraph_index, "table_index": None, "row": None, "cell": None, "tracked_change": None}
        source_ref = _source_ref(
            artifact_id=artifact_id,
            artifact_type=ArtifactType.docx,
            filename=path.name,
            locator=locator,
            method="segment_docx_paragraph",
            parser_version=parser_version,
        )
        segments.append(
            make_segment(
                project_id=project_id,
                artifact_id=artifact_id,
                artifact_type=ArtifactType.docx,
                segment_type="docx_paragraph",
                text=text,
                locator=locator,
                source_ref=source_ref,
                metadata={"heading": bool(paragraph.style and paragraph.style.name.lower().startswith("heading"))},
            )
        )

    for table_index, table in enumerate(document.tables):
        for row_index, row_cells in enumerate(table.rows):
            cell_values = [cell.text.strip() for cell in row_cells.cells]
            row_text = " | ".join(value for value in cell_values if value)
            if not row_text:
                continue
            locator = {"paragraph_index": None, "table_index": table_index, "row": row_index, "cell": None, "tracked_change": None}
            source_ref = _source_ref(
                artifact_id=artifact_id,
                artifact_type=ArtifactType.docx,
                filename=path.name,
                locator=locator,
                method="segment_docx_table_row",
                parser_version=parser_version,
            )
            segments.append(
                make_segment(
                    project_id=project_id,
                    artifact_id=artifact_id,
                    artifact_type=ArtifactType.docx,
                    segment_type="docx_table_row",
                    text=row_text,
                    locator=locator,
                    source_ref=source_ref,
                    metadata={"cell_count": len(cell_values)},
                )
            )

    with zipfile.ZipFile(path) as archive:
        xml_raw = archive.read("word/document.xml")
    root = ET.fromstring(xml_raw)
    for tracked_index, node in enumerate(root.findall(".//w:del", WORD_NS)):
        text_parts = [part.text for part in node.findall(".//w:delText", WORD_NS) if part.text]
        if not text_parts:
            text_parts = [part.text for part in node.findall(".//w:t", WORD_NS) if part.text]
        text = " ".join(part.strip() for part in text_parts if part and part.strip()).strip()
        if not text:
            continue
        locator = {
            "paragraph_index": None,
            "table_index": None,
            "row": None,
            "cell": None,
            "tracked_change": "deleted",
            "tracked_index": tracked_index,
        }
        source_ref = _source_ref(
            artifact_id=artifact_id,
            artifact_type=ArtifactType.docx,
            filename=path.name,
            locator=locator,
            method="segment_docx_tracked_change",
            parser_version=parser_version,
        )
        segments.append(
            make_segment(
                project_id=project_id,
                artifact_id=artifact_id,
                artifact_type=ArtifactType.docx,
                segment_type="docx_tracked_deletion",
                text=text,
                locator=locator,
                source_ref=source_ref,
                metadata={"tracked_change": "deleted"},
            )
        )

    return segments


def _clean_vtt(raw_text: str) -> str:
    rows = []
    for line in raw_text.splitlines():
        stripped = line.strip()
        if stripped.upper() == "WEBVTT" or "-->" in stripped or not stripped:
            continue
        rows.append(line)
    return "\n".join(rows)


def _clean_srt(raw_text: str) -> str:
    rows = []
    for line in raw_text.splitlines():
        stripped = line.strip()
        if stripped.isdigit() or "-->" in stripped or not stripped:
            continue
        rows.append(line)
    return "\n".join(rows)


def _speaker_role(speaker: str | None, text: str) -> str | None:
    source = normalize_text(f"{speaker or ''} {text}")
    if any(token in source for token in ("customer", "client")):
        return "customer"
    if any(token in source for token in ("purtera", "pm", "project manager", "coordinator")):
        return "internal"
    return "unknown" if speaker else None


def segment_transcript(
    *,
    project_id: str,
    artifact_id: str,
    path: Path,
    parser_version: str = "segmenter_transcript_v1",
) -> list[ArtifactSegment]:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    suffix = path.suffix.lower()
    artifact_type = ArtifactType.transcript if suffix in {".vtt", ".srt", ".json"} else ArtifactType.txt
    if suffix == ".json":
        payload = json.loads(raw)
        source_items = payload.get("utterances") if isinstance(payload, dict) else payload
        source_items = source_items if isinstance(source_items, list) else []
        utterances: list[dict[str, Any]] = []
        for idx, item in enumerate(source_items):
            if not isinstance(item, dict):
                continue
            text = str(item.get("text", "")).strip()
            if not text:
                continue
            utterances.append(
                {
                    "utterance_index": idx,
                    "line_start": idx + 1,
                    "line_end": idx + 1,
                    "speaker": item.get("speaker"),
                    "timestamp_start": item.get("start") or item.get("timestamp"),
                    "timestamp_end": item.get("end"),
                    "section": item.get("section"),
                    "text": text,
                }
            )
    else:
        if suffix == ".vtt":
            raw = _clean_vtt(raw)
        elif suffix == ".srt":
            raw = _clean_srt(raw)
        utterances = split_transcript_segments(normalize_transcript_text(raw))

    segments: list[ArtifactSegment] = []
    for utterance in utterances:
        text = str(utterance.get("text", "")).strip()
        if not text:
            continue
        speaker = utterance.get("speaker")
        role = _speaker_role(speaker, text)
        locator = {
            "line_start": utterance.get("line_start"),
            "line_end": utterance.get("line_end"),
            "speaker": speaker,
            "speaker_role": role,
            "timestamp_start": utterance.get("timestamp_start"),
            "timestamp_end": utterance.get("timestamp_end"),
            "section": utterance.get("section"),
            "utterance_index": utterance.get("utterance_index"),
        }
        source_ref = _source_ref(
            artifact_id=artifact_id,
            artifact_type=artifact_type,
            filename=path.name,
            locator=locator,
            method="segment_transcript_utterance",
            parser_version=parser_version,
        )
        segments.append(
            make_segment(
                project_id=project_id,
                artifact_id=artifact_id,
                artifact_type=artifact_type,
                segment_type="transcript_utterance",
                text=text,
                locator=locator,
                source_ref=source_ref,
                speaker=speaker,
                speaker_role=role,
                section=utterance.get("section"),
                metadata={"has_speaker_marker": bool(detect_speaker(text))},
            )
        )

    sections: dict[str, list[ArtifactSegment]] = {}
    for segment in segments:
        if not segment.section:
            continue
        sections.setdefault(segment.section, []).append(segment)
    for section_name, section_segments in sections.items():
        line_start = min((item.locator.get("line_start") or 0) for item in section_segments)
        line_end = max((item.locator.get("line_end") or line_start) for item in section_segments)
        locator = {"line_start": line_start, "line_end": line_end, "section": section_name}
        source_ref = _source_ref(
            artifact_id=artifact_id,
            artifact_type=artifact_type,
            filename=path.name,
            locator=locator,
            method="segment_transcript_section",
            parser_version=parser_version,
        )
        section_text = "\n".join(item.text for item in section_segments)
        segments.append(
            make_segment(
                project_id=project_id,
                artifact_id=artifact_id,
                artifact_type=artifact_type,
                segment_type="transcript_section",
                text=section_text,
                locator=locator,
                source_ref=source_ref,
                section=section_name,
                metadata={"utterance_count": len(section_segments)},
            )
        )

    return segments


def segment_quote(
    *,
    project_id: str,
    artifact_id: str,
    path: Path,
    parser_version: str = "segmenter_quote_v1",
) -> list[ArtifactSegment]:
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        workbook = load_workbook(path, read_only=True, data_only=True)
        source_rows = [
            (sheet.title, [list(row) for row in sheet.iter_rows(values_only=True)], ArtifactType.xlsx)
            for sheet in workbook.worksheets
        ]
    elif suffix == ".csv":
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            source_rows = [("csv", [list(row) for row in csv.reader(handle)], ArtifactType.csv)]
    elif suffix == ".txt":
        content = path.read_text(encoding="utf-8", errors="ignore")
        source_rows = [("txt", [re.split(r"[,\t|]", line) for line in content.splitlines() if line.strip()], ArtifactType.txt)]
    else:
        return []

    segments: list[ArtifactSegment] = []
    for sheet_name, rows, artifact_type in source_rows:
        header_idx, header_map = _detect_header(rows, _QUOTE_HEADER_ALIASES)
        if header_idx is None:
            continue
        columns = {key: get_column_letter(index + 1) for key, index in header_map.items()}
        for row_idx in range(header_idx + 1, len(rows)):
            row = rows[row_idx]
            values = {key: str(row[col]).strip() if col < len(row) and row[col] is not None else "" for key, col in header_map.items()}
            if all(not value for value in values.values()):
                continue
            locator = {"sheet": sheet_name, "row": row_idx + 1, "columns": columns}
            source_ref = _source_ref(
                artifact_id=artifact_id,
                artifact_type=artifact_type,
                filename=path.name,
                locator=locator,
                method="segment_quote_line_item",
                parser_version=parser_version,
            )
            text = " | ".join(f"{key}={value}" for key, value in values.items() if value)
            segments.append(
                make_segment(
                    project_id=project_id,
                    artifact_id=artifact_id,
                    artifact_type=ArtifactType.vendor_quote,
                    segment_type="quote_line_item",
                    text=text,
                    locator=locator,
                    source_ref=source_ref,
                    metadata={"values": values},
                )
            )
    return segments


def segment_text(
    *,
    project_id: str,
    artifact_id: str,
    path: Path,
    parser_version: str = "segmenter_text_v1",
) -> list[ArtifactSegment]:
    artifact_type = ArtifactType.txt
    text = path.read_text(encoding="utf-8", errors="ignore")
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    segments: list[ArtifactSegment] = []
    line_cursor = 1
    for index, block in enumerate(blocks):
        block_lines = block.splitlines()
        locator = {"block_index": index, "line_start": line_cursor, "line_end": line_cursor + len(block_lines) - 1}
        line_cursor += len(block_lines) + 1
        source_ref = _source_ref(
            artifact_id=artifact_id,
            artifact_type=artifact_type,
            filename=path.name,
            locator=locator,
            method="segment_text_block",
            parser_version=parser_version,
        )
        segments.append(
            make_segment(
                project_id=project_id,
                artifact_id=artifact_id,
                artifact_type=artifact_type,
                segment_type="text_block",
                text=block,
                locator=locator,
                source_ref=source_ref,
                metadata={"line_count": len(block_lines)},
            )
        )
    return segments
