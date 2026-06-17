from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as _DocxParagraph
from docx.table import Table as _DocxTable

from app.core.ids import stable_id
from app.core.normalizers import normalize_entity_key, normalize_text
from app.core.segments import ArtifactSegment
from app.core.schemas import (
    ArtifactType,
    AtomType,
    AuthorityClass,
    EvidenceAtom,
    ParserOutput,
    ReviewStatus,
    SourceRef,
    ParserCapability,
    ParserMatch,
)
from app.parsers.base import BaseParser
from app.parsers.segmenters import segment_docx
from app.parsers.structured_projection import (
    derived_files_for,
    make_bullet_list,
    make_page,
    make_paragraph,
    make_section,
    make_structured_document,
    make_table,
    stamp_section_and_block_ids,
)
from app.domain.schemas import DomainPack

STRUCTURED_SCHEMA_DOCX = "orbitbrief.docx.structured.v1"

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
WORD_TBL_TAG = f"{{{WORD_NS['w']}}}tbl"


def _iter_block_items(parent):
    """Yield ``("p", element)`` / ``("tbl", element)`` for the block-level
    paragraphs and tables under ``parent``, IN READING ORDER, **descending into
    ``w:sdt`` content controls**.

    Root-cause fix: python-docx's ``document.paragraphs`` / ``document.tables``
    only iterate direct ``w:body`` children and silently skip everything wrapped
    in a ``<w:sdt>`` structured-document-tag (Word content control / template
    cover-page block). Templated SOWs routinely wrap their entire first page
    (intro, exec summary, contact tables) in one ``w:sdt`` — so the standard
    iterators never see it and the content is lost (or recovered header-less by a
    lossy fallback). Walking ``iterchildren`` and recursing through
    ``w:sdtContent`` restores that content in true document order. Universal:
    any docx, any layout, any nesting depth of content controls."""
    P, TBL, SDT, SDTC = qn("w:p"), qn("w:tbl"), qn("w:sdt"), qn("w:sdtContent")
    for child in parent.iterchildren():
        tag = child.tag
        if tag == P:
            yield ("p", child)
        elif tag == TBL:
            yield ("tbl", child)
        elif tag == SDT:
            content = child.find(SDTC)
            if content is not None:
                yield from _iter_block_items(content)


def _all_paragraphs(document):
    """All body paragraphs in reading order, including those inside content
    controls (drop-in replacement for ``document.paragraphs``)."""
    return [_DocxParagraph(el, document) for kind, el in _iter_block_items(document.element.body) if kind == "p"]


def _all_tables(document):
    """All body tables in reading order, including those inside content controls
    (drop-in replacement for ``document.tables``)."""
    return [_DocxTable(el, document) for kind, el in _iter_block_items(document.element.body) if kind == "tbl"]

SCOPE_PATTERNS = [r"\bscope includes\b", r"\binclude\b", r"\binstallation\b", r"\binstall\b"]
EXCLUSION_PATTERNS = [r"\bexclude\b", r"\bexcluded\b", r"\bout of scope\b", r"\bnot in scope\b"]
# Strong constraint cues: load-bearing phrases that are reliably a constraint.
STRONG_CONSTRAINT_PATTERNS = [
    r"\bcustomer is responsible\b",
    r"\bescort required\b",
    r"\bbadge required\b",
]
# Weak constraint cues: single common words that fire on any mention (e.g.
# "access windows" in a descriptive list). On their own these are low-trust —
# they are typed provisionally and routed to the label queue, not shipped
# confident. See _weak_lexical_types.
WEAK_CONSTRAINT_PATTERNS = [r"\baccess\b"]
CONSTRAINT_PATTERNS = STRONG_CONSTRAINT_PATTERNS + WEAK_CONSTRAINT_PATTERNS
ASSUMPTION_PATTERNS = [r"\bassum(?:e|ption|ing)\b"]


class DocxParser(BaseParser):
    parser_name = "docx"
    parser_version = "docx_parser_v1"
    capability = ParserCapability(
        parser_name=parser_name,
        parser_version=parser_version,
        supported_extensions=[".docx"],
        supported_artifact_types=[ArtifactType.docx],
        emitted_atom_types=[AtomType.scope_item, AtomType.exclusion, AtomType.constraint, AtomType.assumption, AtomType.open_question],
        supported_domain_packs=["*"],
        requires_binary=True,
        supports_source_replay=True,
    )

    def match(self, path: Path, sample_text: str | None, domain_pack: DomainPack | None) -> ParserMatch:
        del sample_text, domain_pack
        suffix = path.suffix.lower()
        confidence = 0.94 if suffix == ".docx" else 0.0
        reasons = ["docx_extension"] if suffix == ".docx" else []
        return ParserMatch(
            parser_name=self.parser_name,
            confidence=confidence,
            reasons=reasons,
            artifact_type=ArtifactType.docx,
        )

    def parse(self, artifact_path: Path) -> list[Any]:
        artifact_id = stable_id("art", str(artifact_path))
        return self.parse_artifact("unknown_project", artifact_id, artifact_path)

    def segment_artifact(self, project_id: str, artifact_id: str, path: Path) -> list[ArtifactSegment]:
        return segment_docx(project_id=project_id, artifact_id=artifact_id, path=path, parser_version=self.parser_version)

    def parse_artifact(
        self,
        project_id: str,
        artifact_id: str,
        path: Path,
        domain_pack: DomainPack | None = None,
    ) -> list[EvidenceAtom]:
        return self.parse_artifact_full(
            project_id=project_id,
            artifact_id=artifact_id,
            path=path,
            domain_pack=domain_pack,
        ).atoms

    def parse_artifact_full(
        self,
        project_id: str,
        artifact_id: str,
        path: Path,
        domain_pack: DomainPack | None = None,
    ) -> ParserOutput:
        del domain_pack
        document = Document(path)
        atoms: list[EvidenceAtom] = []
        # Universal reading-order section map: every paragraph/table learns the
        # heading chain it lives under, so site/section attribution has real
        # signal instead of empty section_path.
        para_section, table_section, _heading_paras, para_order, table_order = self._build_section_index(document)

        # v54 ROOT-CAUSE FIX: skip paragraphs that belong to a table cell so
        # their content is handled by the table loop below with proper row
        # structure and column context preserved (the original v48 intent).
        #
        # The previous implementation precomputed a set of id(paragraph) for
        # every table-cell paragraph and skipped main-loop paragraphs whose
        # id() was in that set. That is fundamentally broken: python-docx
        # builds throwaway Paragraph proxy objects on every access and they are
        # garbage-collected immediately, so the main loop's freshly-allocated
        # proxies reused the same memory addresses → id() collisions → real
        # body paragraphs (e.g. the SOW overview carrying "~110 TVs") were
        # silently skipped as if they were table cells. We now decide table
        # membership STRUCTURALLY by walking the lxml element's ancestors,
        # which is stable and correct across python-docx versions.
        for idx, paragraph in enumerate(_all_paragraphs(document)):
            if self._paragraph_in_table(paragraph):
                continue
            text = paragraph.text.strip()
            if not text:
                continue
            is_list_item = self._paragraph_is_list_item(paragraph)
            # _build_section_index is the single source of truth for what's
            # structure (style heading / bold sub-heading / short colon list-intro)
            # vs content, so the heading-drop decision can never diverge from the
            # section-path computation.
            is_heading = idx in getattr(self, "_structure_idxs", set())
            atoms.extend(
                self._emit_atoms_for_text(
                    project_id=project_id,
                    artifact_id=artifact_id,
                    filename=path.name,
                    text=text,
                    paragraph_index=idx,
                    table_index=None,
                    row=None,
                    cell=None,
                    tracked_change=None,
                    heading=is_heading,
                    is_list_item=is_list_item,
                    section_path=para_section.get(idx, []),
                    lead_in=getattr(self, "_para_lead_in", {}).get(idx, []),
                )
            )

        # Build all-document text once for ``kind=physical_site`` declarations.
        # Exclude table-cell paragraphs so the surrounding-text heuristic stays
        # accurate (otherwise table text bleeds in twice).
        document_text = " ".join(
            (p.text or "").strip()
            for p in _all_paragraphs(document)
            if not self._paragraph_in_table(p)
        )

        for table_idx, table in enumerate(_all_tables(document)):
            # Build a column/rows view of the table for the site_roster
            # extractor (and a per-row atom emitter below).
            table_rows: list[list[str]] = []
            for row_cells in table.rows:
                table_rows.append([c.text.strip() for c in row_cells.cells])

            # Site-roster fast path — when the first non-empty row
            # looks like roster headers + a roster-specific column,
            # emit one physical_site atom per data row.
            roster_atoms = self._maybe_emit_docx_site_roster_atoms(
                project_id=project_id,
                artifact_id=artifact_id,
                filename=path.name,
                table_index=table_idx,
                rows=table_rows,
                surrounding_text=document_text,
                section_path=table_section.get(table_idx, []),
            )
            if roster_atoms:
                atoms.extend(roster_atoms)
                # v48 FIX 3: Only skip per-row fallback when the roster
                # extraction is HIGH-confidence (≥2 atoms, all confidence
                # ≥0.75). Low-confidence rosters fall through so the
                # per-row emitter ALSO processes the table — duplicate
                # coverage beats silent data loss. entity_resolution
                # collapses the dupes later.
                _high_conf_roster = (
                    len(roster_atoms) >= 2
                    and all(getattr(a, "confidence", 0.0) >= 0.75 for a in roster_atoms)
                )
                if _high_conf_roster:
                    continue
                # Low confidence: fall through to per-row emission below.

            # Per-row atom — concatenate all cells so the row's
            # context (Site + Part + Qty) survives as one atom for
            # entity extraction. Skip the all-cell header row when
            # row 0 looks like field labels.
            header_cells = table_rows[0] if table_rows else []
            for row_idx, row_cells in enumerate(table.rows):
                cell_texts = [c.text.strip() for c in row_cells.cells if c.text.strip()]
                if not cell_texts:
                    continue
                # Skip the header row (first row) ONLY when its cells are PURE
                # COLUMN LABELS — short AND value-free (no digit / @ / $). A
                # header-less table whose first row is real data ("Dan Pratt |
                # dan@x.com", "V1 | 5/27/26") must NOT be skipped, or that row is
                # silently lost. Real headers ("FULL NAME | JOB TITLE | EMAIL",
                # "SOW VERSION | QUOTED BY | DATE") carry no values, so they skip.
                if (
                    row_idx == 0
                    and len(cell_texts) >= 2
                    and all(len(c) <= 30 for c in cell_texts)
                    and not any(re.search(r"[\d@$]", c) for c in cell_texts)
                ):
                    continue
                row_text = " | ".join(cell_texts)
                # v49.2: emit a raw_table_row atom alongside the legacy
                # row blob. The centralized _enrich_table_atoms() in
                # entity_extraction will classify all raw_table_row
                # atoms in one pass using the column schema registry.
                if header_cells and row_idx > 0:
                    _row_cells_full = [c.text.strip() for c in row_cells.cells]
                    _rtr_id = stable_id("atm", artifact_id, "raw_table_row", table_idx, row_idx)
                    _rtr_src = SourceRef(
                        id=stable_id("src", _rtr_id),
                        artifact_id=artifact_id,
                        artifact_type=ArtifactType.docx,
                        filename=path.name,
                        locator={"table_index": table_idx, "row": row_idx, "extraction": "raw_table_row_v49_2", "section_path": table_section.get(table_idx, [])},
                        extraction_method="raw_table_row_v49_2",
                        parser_version=self.parser_version,
                    )
                    atoms.append(EvidenceAtom(
                        id=_rtr_id,
                        project_id=project_id,
                        artifact_id=artifact_id,
                        atom_type=AtomType.raw_table_row,
                        raw_text=row_text[:4000],
                        normalized_text=row_text.lower()[:4000],
                        value={
                            "_columns": list(header_cells),
                            "_row": _row_cells_full,
                            "_table_idx": table_idx,
                            "_row_idx": row_idx,
                            "_filename": path.name,
                            "_artifact_type": "docx",
                        },
                        entity_keys=[],
                        source_refs=[_rtr_src],
                        receipts=[],
                        authority_class=AuthorityClass.contractual_scope,
                        confidence=0.80,
                        confidence_raw=0.80,
                        calibrated_confidence=0.80,
                        review_status=ReviewStatus.auto_accepted,
                        review_flags=[],
                        parser_version=self.parser_version,
                    ))
                # Table rows carry structured data even without
                # scope/constraint verbs — emit unconditionally as
                # a scope_item (the classifier path is for prose).
                row_atom_id = stable_id(
                    "atm", artifact_id, "docx_row",
                    table_idx, row_idx, row_text
                )
                row_src = SourceRef(
                    id=stable_id("src", row_atom_id),
                    artifact_id=artifact_id,
                    artifact_type=ArtifactType.docx,
                    filename=path.name,
                    locator={
                        "table_index": table_idx,
                        "row": row_idx,
                        "extraction": "docx_table_row_v1",
                        "section_path": table_section.get(table_idx, []),
                    },
                    extraction_method="docx_table_row_v1",
                    parser_version=self.parser_version,
                )
                atoms.append(
                    EvidenceAtom(
                        id=row_atom_id,
                        project_id=project_id,
                        artifact_id=artifact_id,
                        atom_type=AtomType.scope_item,
                        raw_text=row_text,
                        normalized_text=row_text.lower(),
                        value={
                            "kind": "table_row",
                            "columns": header_cells,
                            "cells": dict(zip(header_cells, cell_texts)) if header_cells else {f"col_{i}": v for i, v in enumerate(cell_texts)},
                        },
                        entity_keys=[],
                        source_refs=[row_src],
                        receipts=[],
                        authority_class=AuthorityClass.contractual_scope,
                        confidence=0.85,
                        confidence_raw=0.85,
                        calibrated_confidence=0.85,
                        review_status=ReviewStatus.auto_accepted,
                        review_flags=[],
                        parser_version=self.parser_version,
                    )
                )

        atoms.extend(
            self._extract_tracked_change_atoms(
                project_id=project_id,
                artifact_id=artifact_id,
                filename=path.name,
                path=path,
            )
        )

        # Comments — sidebar Word comments often carry scope-relevant
        # notes ("we're cutting this from scope", "verify w/ vendor",
        # "approved by Jane"). Extract them so they don't get
        # silently dropped.
        atoms.extend(
            self._extract_comment_atoms(
                project_id=project_id,
                artifact_id=artifact_id,
                filename=path.name,
                path=path,
            )
        )

        # Never-detected recovery: pull in content that lives inside content
        # controls / textboxes (invisible to Document.paragraphs/.tables) and
        # mark every embedded binary region. Dedup against what we already
        # emitted so AlternateContent fallbacks don't double-count.
        already_emitted = {
            a.normalized_text for a in atoms if getattr(a, "normalized_text", "")
        }
        atoms.extend(
            self._recover_nested_region_atoms(
                project_id=project_id,
                artifact_id=artifact_id,
                filename=path.name,
                document=document,
                already_emitted=already_emitted,
            )
        )
        atoms.extend(
            self._emit_embedded_media_markers(
                project_id=project_id,
                artifact_id=artifact_id,
                filename=path.name,
                path=path,
            )
        )

        # Reading-order sort: python-docx yields all paragraphs then all tables,
        # so a table that sits right under a heading in the document otherwise
        # surfaces far down the atom list. Re-order atoms to true document order
        # using the body sequence computed in _build_section_index. Stable sort
        # preserves the emission sub-order within one paragraph/table (e.g. the
        # raw_table_row alongside its row blob) and parks order-less atoms
        # (comments, tracked changes, recovered/embedded regions) at the end.
        def _body_key(a: Any) -> tuple[int, int]:
            try:
                refs = getattr(a, "source_refs", None) or []
                loc = (getattr(refs[0], "locator", None) or {}) if refs else {}
            except Exception:
                loc = {}
            pi = loc.get("paragraph_index") if isinstance(loc, dict) else None
            ti = loc.get("table_index") if isinstance(loc, dict) else None
            if pi is not None and pi in para_order:
                return (para_order[pi], 0)
            if ti is not None and ti in table_order:
                return (table_order[ti], 1)
            return (10**9, 2)
        atoms.sort(key=_body_key)

        structured_doc = self._build_structured_doc(filename=path.name, document=document)
        stamp_section_and_block_ids(structured_doc, artifact_seed=artifact_id)
        return ParserOutput(
            atoms=atoms,
            derived_files=derived_files_for(artifact_path=path, structured_doc=structured_doc),
        )

    def _maybe_emit_docx_site_roster_atoms(
        self,
        *,
        project_id: str,
        artifact_id: str,
        filename: str,
        table_index: int,
        rows: list[list[str]],
        surrounding_text: str,
        section_path: list[str] | None = None,
    ) -> list[EvidenceAtom]:
        """Emit ``physical_site`` entity atoms when a DOCX table looks
        like a site roster (header row with Site ID + Facility +
        Address / MDF / Access / Escort columns). ``section_path`` (the heading
        chain the roster table lives under) is stamped on each atom so roster
        sites carry their section like every other table row.
        """
        try:
            from app.parsers.site_roster_extractor import (
                extract_site_roster,
                looks_like_site_roster,
                map_columns_to_fields,
            )
        except Exception:  # pragma: no cover
            return []
        if not rows or len(rows) < 2:
            return []
        header = [(c or "").strip() for c in rows[0]]
        data_rows: list[dict[str, str]] = []
        for r in rows[1:]:
            cells = {
                header[i] if i < len(header) and header[i] else f"col_{i}":
                (r[i] if i < len(r) else "") or ""
                for i in range(max(len(header), len(r)))
            }
            if any(v.strip() for v in cells.values()):
                data_rows.append(cells)
        if not data_rows:
            return []
        try:
            if not looks_like_site_roster(
                columns=header, rows=data_rows, surrounding_text=surrounding_text
            ):
                return []
            field_map = map_columns_to_fields(header)
            roster_specific = {
                "facility_name", "street_address", "mdf_idf",
                "access_window", "escort_owner", "city_state",
            }
            if not (set(field_map.values()) & roster_specific):
                return []
            roster_rows = extract_site_roster(
                columns=header, rows=data_rows, surrounding_text=surrounding_text
            )
        except Exception:  # pragma: no cover
            return []
        if not roster_rows:
            return []

        out: list[EvidenceAtom] = []
        for site_row in roster_rows:
            sid = (site_row.site_id or "").strip()
            canon_id = sid or site_row.facility_name or ""
            if not canon_id:
                continue
            row_index = site_row.row_index + 1  # +1 for header
            text_parts: list[str] = []
            for label, val in [
                ("site_id", sid or site_row.site_id),
                ("facility", site_row.facility_name),
                ("address", site_row.street_address),
                ("mdf_idf", site_row.mdf_idf),
                ("access", site_row.access_window),
                ("escort", site_row.escort_owner),
                ("contact", site_row.contact),
                ("phone", site_row.phone),
                ("email", site_row.email),
                ("city_state", site_row.city_state),
                ("zip", site_row.zip),
                ("sqft", site_row.sqft),
                ("users", site_row.occupancy),
                ("notes", site_row.notes),
            ]:
                if val:
                    text_parts.append(f"{label}: {val}")
            # Surface every remaining column the extractor could not map to a
            # canonical field (e.g. "Rooms") so NO column is invisible to the head.
            for _k, _v in (site_row.extra_fields or ()):
                if _v and str(_v).strip():
                    text_parts.append(f"{_k}: {_v}")
            row_text = " | ".join(text_parts) or canon_id
            entity_keys: list[str] = []
            if sid:
                slug = re.sub(r"[^a-z0-9]+", "_", sid.lower()).strip("_")
                if slug:
                    entity_keys.append(f"site:{slug}")
            atom_id = stable_id(
                "atm", artifact_id, "docx_site_roster",
                table_index, row_index, canon_id
            )
            src = SourceRef(
                id=stable_id("src", atom_id),
                artifact_id=artifact_id,
                artifact_type=ArtifactType.docx,
                filename=filename,
                locator={
                    "table_index": table_index,
                    "row": row_index,
                    "extraction": "docx_site_roster_v1",
                    "section_path": list(section_path) if section_path else [],
                },
                extraction_method="docx_site_roster_v1",
                parser_version=self.parser_version,
            )
            out.append(
                EvidenceAtom(
                    id=atom_id,
                    project_id=project_id,
                    artifact_id=artifact_id,
                    # v53.2 ROOT-CAUSE FIX: physical_site (not entity).
                    atom_type=AtomType.physical_site,
                    raw_text=row_text,
                    normalized_text=row_text.lower(),
                    value={
                        "kind": "physical_site",
                        "id": sid or site_row.site_id,
                        "site_id": sid or site_row.site_id,
                        "name": site_row.facility_name,
                        "facility_name": site_row.facility_name,
                        "address": site_row.street_address,
                        "street_address": site_row.street_address,
                        "mdf_idf": site_row.mdf_idf,
                        "access_window": site_row.access_window,
                        "escort_owner": site_row.escort_owner,
                        "sqft": site_row.sqft,
                        "occupancy": site_row.occupancy,
                        "contact": site_row.contact,
                        "phone": site_row.phone,
                        "email": site_row.email,
                        "city_state": site_row.city_state,
                        "zip": site_row.zip,
                        "notes": site_row.notes,
                        "extras": dict(site_row.extra_fields),
                    },
                    entity_keys=sorted(set(entity_keys)),
                    source_refs=[src],
                    receipts=[],
                    authority_class=AuthorityClass.contractual_scope,
                    confidence=site_row.confidence,
                    confidence_raw=site_row.confidence,
                    calibrated_confidence=site_row.confidence,
                    review_status=ReviewStatus.auto_accepted,
                    review_flags=[],
                    parser_version=self.parser_version,
                )
            )
        return out

    def _build_structured_doc(
        self,
        *,
        filename: str,
        document: Any,
    ) -> dict[str, Any]:
        """Build a structured doc from the DOCX, walking the body element
        sequentially so paragraphs, bullets, and tables stay in source
        order.  Headings open new sections; consecutive bullets fuse
        into a bullet_list block.
        """
        sections: list[dict[str, Any]] = []
        current_blocks: list[dict[str, Any]] = []
        current_heading = filename
        current_level = 1
        pending_bullets: list[dict[str, Any]] = []

        def flush_bullets() -> None:
            nonlocal pending_bullets
            if pending_bullets:
                current_blocks.append(make_bullet_list(items=pending_bullets))
                pending_bullets = []

        def flush_section() -> None:
            nonlocal current_blocks, current_heading, current_level
            flush_bullets()
            if current_blocks or current_heading:
                sections.append(
                    make_section(
                        heading=current_heading,
                        level=current_level,
                        blocks=current_blocks,
                    )
                )
            current_blocks = []

        for paragraph in _all_paragraphs(document):
            text = (paragraph.text or "").strip()
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            is_heading = style_name.startswith("heading")
            is_list = "list" in style_name or "bullet" in style_name
            if is_heading and text:
                flush_section()
                current_heading = text
                # Heading 1 -> level 2, Heading 2 -> level 3, etc.
                m = re.search(r"\d+", style_name)
                level = (int(m.group()) + 1) if m else 2
                current_level = max(2, min(level, 6))
                continue
            if not text:
                continue
            if is_list:
                pending_bullets.append({"text": text, "children": []})
            else:
                flush_bullets()
                current_blocks.append(make_paragraph(text))

        for table in _all_tables(document):
            try:
                cells = [[(c.text or "").strip() for c in row.cells] for row in table.rows]
            except Exception:
                continue
            if not cells:
                continue
            columns = cells[0] or [f"col_{i + 1}" for i in range(len(cells[0]))]
            columns = [c or f"col_{i + 1}" for i, c in enumerate(columns)]
            rows: list[dict[str, Any]] = []
            for raw in cells[1:]:
                if all(not v for v in raw):
                    continue
                rows.append({columns[i]: raw[i] if i < len(raw) else "" for i in range(len(columns))})
            flush_bullets()
            current_blocks.append(make_table(columns=columns, rows=rows))

        flush_section()

        if not sections:
            sections.append(
                make_section(
                    heading=filename,
                    level=2,
                    blocks=[make_paragraph("(empty document)")],
                )
            )
        page = make_page(page=0, title=filename, sections=sections)
        return make_structured_document(
            schema_version=STRUCTURED_SCHEMA_DOCX,
            filename=filename,
            artifact_type=ArtifactType.docx.value,
            title=filename,
            metadata=[],
            pages=[page],
        )

    def _extract_comment_atoms(
        self,
        project_id: str,
        artifact_id: str,
        filename: str,
        path: Path,
    ) -> list[EvidenceAtom]:
        """Pull Word comments out of word/comments.xml. Each comment
        becomes a low-confidence scope_item atom flagged so the
        reviewer can decide whether the side-note is in-scope."""
        atoms: list[EvidenceAtom] = []
        try:
            with zipfile.ZipFile(path) as zf:
                if "word/comments.xml" not in zf.namelist():
                    return []
                xml_raw = zf.read("word/comments.xml")
        except Exception:
            return []
        try:
            root = ET.fromstring(xml_raw)
        except Exception:
            return []
        for idx, node in enumerate(root.findall(".//w:comment", WORD_NS)):
            text_parts = [t.text for t in node.findall(".//w:t", WORD_NS) if t.text]
            text = " ".join(part.strip() for part in text_parts if part.strip()).strip()
            if not text:
                continue
            author = node.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", "")
            initials = node.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}initials", "")
            comment_id = node.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", str(idx))
            atom_text = text if not author else f"[Comment by {author}] {text}"
            atom_id = stable_id("atm", project_id, artifact_id, "docx_comment", comment_id, atom_text)
            src = SourceRef(
                id=stable_id("src", atom_id),
                artifact_id=artifact_id,
                artifact_type=ArtifactType.docx,
                filename=filename,
                locator={
                    "comment_id": comment_id,
                    "author": author,
                    "initials": initials,
                    "extraction": "docx_comment_v1",
                },
                extraction_method="docx_comment_v1",
                parser_version=self.parser_version,
            )
            atoms.append(
                EvidenceAtom(
                    id=atom_id,
                    project_id=project_id,
                    artifact_id=artifact_id,
                    atom_type=AtomType.open_question,
                    raw_text=atom_text,
                    normalized_text=atom_text.lower(),
                    value={
                        "kind": "docx_comment",
                        "author": author,
                        "initials": initials,
                        "comment_id": comment_id,
                    },
                    entity_keys=[],
                    source_refs=[src],
                    receipts=[],
                    authority_class=AuthorityClass.meeting_note,
                    confidence=0.55,
                    confidence_raw=0.55,
                    calibrated_confidence=0.55,
                    review_status=ReviewStatus.needs_review,
                    review_flags=["docx_sidebar_comment"],
                    parser_version=self.parser_version,
                )
            )
        return atoms

    def _extract_tracked_change_atoms(
        self,
        project_id: str,
        artifact_id: str,
        filename: str,
        path: Path,
    ) -> list[EvidenceAtom]:
        atoms: list[EvidenceAtom] = []
        with zipfile.ZipFile(path) as zf:
            xml_raw = zf.read("word/document.xml")
        root = ET.fromstring(xml_raw)

        for idx, node in enumerate(root.findall(".//w:del", WORD_NS)):
            text_parts = [t.text for t in node.findall(".//w:delText", WORD_NS) if t.text]
            if not text_parts:
                text_parts = [t.text for t in node.findall(".//w:t", WORD_NS) if t.text]
            text = " ".join(part.strip() for part in text_parts if part.strip()).strip()
            if not text:
                continue
            atoms.extend(
                self._emit_atoms_for_text(
                    project_id=project_id,
                    artifact_id=artifact_id,
                    filename=filename,
                    text=text,
                    paragraph_index=None,
                    table_index=None,
                    row=None,
                    cell=None,
                    tracked_change="deleted",
                    heading=False,
                    tracked_index=idx,
                )
            )

        for idx, node in enumerate(root.findall(".//w:ins", WORD_NS)):
            text_parts = [t.text for t in node.findall(".//w:t", WORD_NS) if t.text]
            text = " ".join(part.strip() for part in text_parts if part.strip()).strip()
            if not text:
                continue
            atoms.extend(
                self._emit_atoms_for_text(
                    project_id=project_id,
                    artifact_id=artifact_id,
                    filename=filename,
                    text=text,
                    paragraph_index=None,
                    table_index=None,
                    row=None,
                    cell=None,
                    tracked_change="inserted",
                    heading=False,
                    tracked_index=idx,
                )
            )
        return atoms

    def _extract_entity_keys(self, text: str) -> list[str]:
        lowered = normalize_text(text)
        keys: list[str] = []
        if "main campus" in lowered:
            keys.append(normalize_entity_key("site", "Main Campus"))
        if "west wing" in lowered:
            keys.append(normalize_entity_key("site", "West Wing"))
        if "camera" in lowered:
            keys.append(normalize_entity_key("device", "IP Cameras"))
        return keys

    def _classify_text(self, text: str) -> list[AtomType]:
        lowered = normalize_text(text)
        atom_types: list[AtomType] = []
        if any(re.search(pattern, lowered) for pattern in SCOPE_PATTERNS):
            atom_types.append(AtomType.scope_item)
        if any(re.search(pattern, lowered) for pattern in EXCLUSION_PATTERNS):
            atom_types.append(AtomType.exclusion)
        if any(re.search(pattern, lowered) for pattern in CONSTRAINT_PATTERNS):
            atom_types.append(AtomType.constraint)
        if any(re.search(pattern, lowered) for pattern in ASSUMPTION_PATTERNS):
            atom_types.append(AtomType.assumption)
        if "?" in text:
            atom_types.append(AtomType.open_question)
        # Stakeholder / signature-block pattern: a role keyword
        # adjacent to a proper-noun name. Emit a scope_item so the
        # downstream stakeholder extractor sees the name in atom
        # entity_keys. Without this, signature paragraphs like
        # "OPTBOT - Director of Workplace Technology: Jane Roe"
        # produce zero atoms despite carrying a load-bearing
        # stakeholder.
        if not atom_types and re.search(
            r"\b(?:director|vp|cio|cto|cfo|coo|ceo|president|"
            r"manager|architect|engineer|owner|sponsor|"
            r"program\s+manager|project\s+manager|principal|"
            r"foreman|superintendent|approver|signatory|"
            r"officer|administrator|consultant)\b",
            lowered,
        ):
            # Must also contain at least one likely name token
            # (Capitalized word, 2+ chars).
            if re.search(r"\b[A-Z][a-z]{1,}\b", text):
                atom_types.append(AtomType.scope_item)
        return atom_types

    def _weak_lexical_types(self, text: str) -> set[AtomType]:
        """Lexical types that rest ONLY on a brittle single-word cue.

        ``constraint`` and ``assumption`` are the two flip-floppy lexical types:
        a bare "access" (in e.g. "…access windows…") fires the constraint rule,
        and a lone "assume" fires the assumption rule, even in descriptive prose
        that is neither. These guesses are low-trust — the parser ships them
        PROVISIONAL (weak_label -> needs_review, low confidence) so they (a)
        stop polluting the brief as confident facts and (b) become rows in the
        PM labelling queue that trains the eventual supervised type head. Strong
        constraint phrases ("escort required", "badge required") are unaffected.
        """
        lowered = normalize_text(text)
        weak: set[AtomType] = set()
        if any(re.search(p, lowered) for p in WEAK_CONSTRAINT_PATTERNS) and not any(
            re.search(p, lowered) for p in STRONG_CONSTRAINT_PATTERNS
        ):
            weak.add(AtomType.constraint)
        if any(re.search(p, lowered) for p in ASSUMPTION_PATTERNS):
            weak.add(AtomType.assumption)
        return weak

    @staticmethod
    def _paragraph_is_list_item(paragraph: Any) -> bool:
        """Whether a paragraph is a bullet/numbered list item.

        Decided STRUCTURALLY via the ``w:numPr`` numbering reference, not just
        the style name: Word frequently leaves list paragraphs on the ``Normal``
        style while carrying real numbering (``<w:pPr><w:numPr>``). The old
        style-name-only check missed those, so the prose gate then dropped short
        bullets like "Connecting power where available" (4 words, no digit).
        List items are deliberate scope content and must fail OPEN regardless of
        length — this restores that for every numbering-driven bullet."""
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        if "list" in style or "bullet" in style:
            return True
        el = getattr(paragraph, "_p", None)
        if el is None:
            return False
        pPr = el.find(qn("w:pPr"))
        return pPr is not None and pPr.find(qn("w:numPr")) is not None

    @staticmethod
    def _list_level(paragraph: Any) -> int:
        """Nesting depth of a list item (0 = top bullet, 1 = sub-bullet, ...), read
        STRUCTURALLY from ``w:numPr/w:ilvl``. This is the parent->child relation
        between a bullet and its sub-bullets — it lives in the indent level, not the
        words, so it's structural (no embedding)."""
        el = getattr(paragraph, "_p", None)
        if el is None:
            return 0
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            return 0
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return 0
        ilvl = numPr.find(qn("w:ilvl"))
        if ilvl is not None:
            try:
                return max(0, int(ilvl.get(qn("w:val"))))
            except Exception:
                return 0
        return 0

    # A sentence whose grammatical job is to ANNOUNCE a following list / section
    # ("PurTera will provide field technicians to perform the following services.",
    # "The scope is as follows."). It carries no standalone fact — it frames its
    # children — so it's lifted onto them as lead_in context rather than emitted.
    _FRAMING_LEAD_IN_RE = re.compile(r"\b(the following|as follows)\b", re.I)
    _LEAD_IN_RULE = None  # lazily-built SemanticRule (see _lead_in_rule)

    @classmethod
    def _lead_in_lexical(cls, t: str) -> bool:
        """Keyword fallback for the lead-in judgment, used ONLY when the embedder
        is offline. Presence of the 'the following'/'as follows' cue — the
        structural gate (a bullet directly follows, no intervening heading) is
        what actually constrains it, so the offline net just needs the cue."""
        return bool(cls._FRAMING_LEAD_IN_RE.search(t))

    @classmethod
    def _lead_in_rule(cls):
        """SEMANTIC lead-in detector: fires on MEANING, not keywords, so a lead-in
        worded "the vendor's responsibilities encompass:" fires the same as "...the
        following services." Falls back to the keyword rule when embeddings are
        down, so a parse never breaks. Examples are the only knob — a correction
        becomes a new example, not a new regex."""
        if cls._LEAD_IN_RULE is None:
            from app.core.semantic_rules import SemanticRule

            cls._LEAD_IN_RULE = SemanticRule(
                name="docx_list_lead_in",
                # POLARITY-AGNOSTIC: "does this announce a following list" — true for
                # service, exclusion, customer, deliverable intros alike. (Polarity
                # is handled separately by the contradiction gate, not here.)
                positives=[
                    "PurTera will provide field technicians to perform the following services.",
                    "Subject to the other provisions of this SOW, Provider will perform the following services.",
                    "The vendor shall complete the following tasks:",
                    "Services include:",
                    "Scope of work consists of the following activities:",
                    "The contractor will perform the work as follows:",
                    "PurTera will provide the following deliverables:",
                    "The vendor's responsibilities encompass the items below:",
                    "The following items are excluded from this SOW unless separately quoted:",
                    "The following are out of scope:",
                    "Customer responsibilities include the following:",
                    "The customer is responsible for the following:",
                    # short section labels (a header over a following list)
                    "Deliverables:", "Assumptions:", "Requirements:",
                    "Notes:", "Exclusions:", "Scope of work:",
                ],
                # NOT list-intros (any polarity): standalone facts, AND key->value
                # lines that also end in ':' but carry a value, not a list header.
                negatives=[
                    "This SOW does not include predictive wireless design or spectrum analysis.",
                    "The school currently receives 5 Gbps of internet bandwidth.",
                    "Access point placement validation is limited to confirming locations align with floor plans.",
                    "All work will be performed during normal business hours.",
                    "The vendor agrees to hold the client harmless from any liability.",
                    "Payment is due within thirty days of invoice receipt.",
                    "The total contract value is fixed at the agreed amount.",
                    "Address: 123 Main Street, Macon GA",
                    "Phone: 555-0100", "Total: $5,000", "Date: January 1, 2026",
                ],
                threshold=0.62,
                lexical_fallback=cls._lead_in_lexical,
            )
        return cls._LEAD_IN_RULE

    @classmethod
    def _is_framing_lead_in(cls, text: str) -> bool:
        t = (text or "").strip()
        # Cheap structural prefilter: bounds what we embed, and is the SHAPE any
        # list lead-in has regardless of wording. The lead-in judgment itself is
        # semantic (embedding) with a keyword fallback when the embedder is down.
        if not t or len(t) > 200 or not t.endswith((".", ":")):
            return False
        words = re.findall(r"[A-Za-z][A-Za-z'\-]*", t)
        # floor of 1 so short section labels ("Deliverables:", "Assumptions:")
        # also get the semantic judgment instead of a word-count shortcut.
        if not (1 <= len(words) <= 25):
            return False
        return cls._lead_in_rule().fires(t)

    _SUBSECTION_BLOCK_RE = re.compile(
        r"out\s*of\s*scope|exclusion|excluded|not\s+included|"
        r"customer\s+(?:responsib|oblig)|client\s+(?:responsib|oblig)|by\s+others",
        re.I,
    )
    _SUBSECTION_BLOCK_RULE = None

    @classmethod
    def _subsection_blocks_lift(cls, heading: str) -> bool:
        """Does this sub-heading CONTRADICT a 'vendor will provide' preamble — i.e.
        is it an exclusion ('Out of Scope') or other-party ('Customer
        Responsibilities') section? If so, the section preamble must NOT lift onto
        its bullets. Semantic (fires=True means BLOCK): positives are contradiction
        headings, negatives are vendor/service headings; the nearest wins. Falls
        back to a keyword check offline."""
        h = (heading or "").strip()
        if not h or len(h) > 120:
            return False
        if cls._SUBSECTION_BLOCK_RULE is None:
            from app.core.semantic_rules import SemanticRule

            cls._SUBSECTION_BLOCK_RULE = SemanticRule(
                name="docx_subsection_blocks_lift",
                positives=[  # contradiction sections -> BLOCK the vendor preamble
                    "Out of Scope", "Exclusions", "Not Included",
                    "Items provided by others", "Customer Responsibilities",
                    "Client Obligations", "Customer will provide",
                ],
                negatives=[  # vendor / service sections -> allow the lift
                    "Vendor scope of work", "Services the provider performs",
                    "Installation and configuration tasks",
                    "Deliverables the vendor provides", "Provider responsibilities",
                ],
                threshold=0.50,
                lexical_fallback=lambda t: bool(cls._SUBSECTION_BLOCK_RE.search(t)),
            )
        return cls._SUBSECTION_BLOCK_RULE.fires(h)

    @staticmethod
    def _is_bold_subheading(paragraph: Any) -> bool:
        """A short, fully-bold, non-list line that Word left on the ``Normal``
        style is a VISUAL sub-heading (e.g. "Configuration Support") the author
        used to open a sub-section. The style-name heading check misses it, so it
        was dropped by the prose gate AND failed to open a section for the bullets
        beneath it. Detect it structurally (whole line bold, few words, not a
        sentence/label) so it becomes section structure like any other heading.
        Conservative on purpose: bold emphasis inside a real sentence (which ends
        with terminal punctuation, or is long) is not promoted."""
        text = (paragraph.text or "").strip()
        if not text or text.endswith((".", ":", ";", "?", "!")):
            return False
        words = re.findall(r"[A-Za-z][A-Za-z'\-]*", text)
        if not (1 <= len(words) <= 8):
            return False
        runs = [r for r in paragraph.runs if (r.text or "").strip()]
        if not runs:
            return False

        def _bold(r: Any) -> bool:
            if r.bold is not None:
                return bool(r.bold)
            return bool(r.font is not None and r.font.bold)

        return all(_bold(r) for r in runs)

    @staticmethod
    def _paragraph_in_table(paragraph: Any) -> bool:
        """Whether a python-docx paragraph lives inside a table cell.

        Decided STRUCTURALLY by walking the underlying lxml element's
        ancestors for a ``<w:tbl>`` tag. This replaces an ``id(paragraph)``
        membership test that was unsound: python-docx creates throwaway
        Paragraph proxy objects on every access and they are GC'd at once, so
        ``id()`` values get reused across the table scan and the main loop,
        causing false-positive matches that silently dropped body paragraphs.
        Ancestry never collides and is stable across python-docx versions.
        """
        el = getattr(paragraph, "_p", None)
        if el is None:
            return False
        parent = el.getparent()
        while parent is not None:
            if parent.tag == WORD_TBL_TAG:
                return True
            parent = parent.getparent()
        return False

    @staticmethod
    def _span_id(artifact_id: str, paragraph_index, table_index, row, cell, tracked_change, tracked_index) -> str:
        """Stable id for a raw source unit, used by the span-provenance ledger."""
        if tracked_change is not None:
            return f"{artifact_id}:{tracked_change}{tracked_index}"
        if table_index is not None:
            return f"{artifact_id}:t{table_index}.r{row}.c{cell}"
        return f"{artifact_id}:p{paragraph_index}"

    @staticmethod
    def _prose_drop_reason(text: str, *, heading: bool) -> str:
        """One-line diagnosis for why the prose GATE dropped a paragraph.

        Doubles as the fix pointer in the lost-content report: it names the
        exact condition in ``_is_substantive_prose`` that rejected the line.
        """
        if heading:
            return "heading/label fragment, no scope/exclusion verb"
        words = re.findall(r"[A-Za-z][A-Za-z'\-]*", text)
        if len(words) < 5:
            return f"short fragment ({len(words)} words), no numeric fact"
        # 5+ word lines now fail open (kept as prose_fallback), so a GATE drop
        # of a multi-word line should not occur; if it ever does, surface it.
        return "multi-word line unexpectedly dropped at prose gate"

    @staticmethod
    def _is_substantive_prose(text: str) -> bool:
        """Whether a paragraph is load-bearing narrative prose worth keeping
        even when it matches none of the scope/exclusion/constraint patterns.

        The lexical classifier is a *type hint*, not a keep/drop gate: an SOW
        overview sentence ("The customer requires onsite field services to
        replace approximately 110 existing TVs across 23 dwellings...") carries
        the deal's headline facts yet uses no scope verb. Dropping it silently
        loses data. This is a universal, content-derived signal — a real
        multi-word sentence — not a keyword whitelist. Downstream semantic
        classification and dedup refine the type and prune true boilerplate."""
        words = re.findall(r"[A-Za-z][A-Za-z'\-]*", text)
        has_number = bool(re.search(r"\d", text))
        if len(words) < 5:
            # Short fragment. Keep it ONLY when it states a concrete numeric
            # fact with a little surrounding context, e.g. a "label: value"
            # line like "Estimated quantity: 110 units" or "Project duration:
            # 2 weeks". These carry load-bearing deal facts yet are too short
            # to be a sentence. Bare headings/labels ("PROJECT OVERVIEW",
            # "Name:") have no digit and are dropped; a lone number with no
            # words ("110") lacks context and is dropped. This is a
            # content-derived signal, not a keyword whitelist.
            return has_number and len(words) >= 2
        # 5+ real words -> KEEP. A multi-word line is load-bearing prose: SOW
        # narrative, an exclusion bullet ("Procurement or supply of TVs,
        # mounts, brackets..."), a tooling list item ("Battery powered drill
        # or impact driver"). The deterministic GATE must NOT silently kill
        # these just because they lack terminal punctuation or a digit —
        # exclusion lists ("what's NOT in scope") and equipment lists are
        # exactly the load-bearing facts a PM needs. We fail OPEN here and let
        # the learnable decide() SEAM (semantic dedup + boilerplate drop) prune
        # true boilerplate, where a human correction can teach it. That turns a
        # silent recall loss into a visible, learnable decision. Universal,
        # content-derived (a multi-word line), not a keyword whitelist.
        return True

    @staticmethod
    def _heading_level(style_name: str | None) -> int | None:
        """Heading depth from a paragraph style name, layout-agnostic.
        'Heading 1'->1, 'Heading 2'->2, 'Title'->0; non-headings -> None."""
        s = (style_name or "").strip().lower()
        if s in ("title", "subtitle"):
            return 0
        if s.startswith("heading"):
            m = re.search(r"(\d+)", s)
            return int(m.group(1)) if m else 1
        return None

    def _build_section_index(
        self, document: Any
    ) -> tuple[dict[int, list[str]], dict[int, list[str]], dict[int, tuple[int, list[str]]]]:
        """ONE reading-order pass over the document body, maintaining a heading
        stack, so EVERY paragraph and table knows the section (heading chain) it
        lives under — universal across any layout (heading→table, heading→prose→
        table, nested sub-headings, etc.). python-docx exposes paragraphs and
        tables as separate sequences, which is exactly why section context was
        being lost; walking ``body.iterchildren()`` restores true document order.

        Returns ``(para_section, table_section, heading_paras)``:
          - para_section[i]   = section_path for the i-th body paragraph
          - table_section[i]  = section_path for the i-th table
          - heading_paras[i]  = (level, ancestor_section_path) for heading paras
        Indices align with ``document.paragraphs`` / ``document.tables`` order.
        """
        para_section: dict[int, list[str]] = {}
        table_section: dict[int, list[str]] = {}
        heading_paras: dict[int, tuple[int, list[str]]] = {}
        # Connective tissue: the governing lead-in sentence(s) above an atom
        # ("PurTera will provide field technicians to perform the following
        # services.") that frame what its children ARE. Attached to every child
        # so the heads see the framing without it being a wasted standalone atom.
        para_lead_in: dict[int, list[str]] = {}
        # global reading-order sequence per body paragraph / table, so atoms can
        # be emitted/sorted into true document order (python-docx otherwise yields
        # all paragraphs, then all tables — losing the interleave).
        para_order: dict[int, int] = {}
        table_order: dict[int, int] = {}
        # paragraph indices that are STRUCTURE (a heading / section-only label),
        # not content — the single source of truth the main loop uses to decide
        # which paragraphs are dropped as atoms (so it never diverges from here).
        structure_idxs: set[int] = set()
        # (level, breadcrumb_label, is_list_intro, lead_in_text). breadcrumb_label
        # is "" for a pure framing lead-in (it must NOT pollute the section path);
        # lead_in_text is None for a normal heading/short label.
        stack: list[tuple[int, str, bool, str | None]] = []
        # Bullet hierarchy: last bullet text seen at each list level, so a sub-bullet
        # ("After Hours: 50% increase") carries its PARENT bullet ("All Services will
        # be performed during normal Business Hours...") as context instead of
        # reading in isolation. Cleared when the list ends (heading / non-list line).
        bullet_by_level: dict[int, str] = {}
        pidx = -1
        tidx = -1
        seq = 0
        try:
            # descend into w:sdt content controls so indices align with
            # _all_paragraphs / _all_tables (the main extraction loops).
            children = list(_iter_block_items(document.element.body))
        except Exception:
            self._structure_idxs = structure_idxs
            self._para_lead_in = para_lead_in
            return para_section, table_section, heading_paras, para_order, table_order

        def _next_is_bullet(k: int) -> bool:
            j = k + 1
            if j < len(children) and children[j][0] == "p":
                try:
                    return self._paragraph_is_list_item(_DocxParagraph(children[j][1], document))
                except Exception:
                    return False
            return False

        def _list_follows_within(k: int, n: int) -> bool:
            """A bullet appears within the next ``n`` paragraphs — headings ALLOWED
            in between, because a section PREAMBLE ('PurTera will provide field
            technicians to perform the following services.') sits a sub-heading
            above its list. What stops the preamble bleeding into a contradictory
            subsection ('Out of Scope', 'Customer Responsibilities') is the SEMANTIC
            contradiction gate applied at lift time, not this structural scan."""
            j, seen = k + 1, 0
            while j < len(children) and seen < n:
                if children[j][0] == "p":
                    try:
                        if self._paragraph_is_list_item(_DocxParagraph(children[j][1], document)):
                            return True
                    except Exception:
                        pass
                    seen += 1
                j += 1
            return False

        for k, (kind, child) in enumerate(children):
            if kind == "p":
                pidx += 1
                para_order[pidx] = seq
                seq += 1
                try:
                    para = _DocxParagraph(child, document)
                    text = (para.text or "").strip()
                    style = (para.style.name or "") if para.style is not None else ""
                    is_list = self._paragraph_is_list_item(para)
                except Exception:
                    para_section[pidx] = [t for _, t, _, _, _ in stack if t]
                    para_lead_in[pidx] = []
                    continue
                lvl = self._heading_level(style)
                if lvl is None and text and not is_list and self._is_bold_subheading(para):
                    # bold sub-heading Word left on Normal style — nest it below
                    # style headings so its following bullets inherit the section.
                    lvl = 3
                # LIST-INTRO: a colon-ending label/bullet immediately followed by
                # sub-bullets ("PMO Responsibilities:", "Services include:") is a
                # sub-section over those bullets — promote it so the bullets carry
                # the "...> PMO Responsibilities" path the heads need. A SHORT label
                # is structure only (no atom); a longer intro sentence stays an atom
                # AND opens the section (keeps its clause, still organizes children).
                is_intro = lvl is None and bool(text) and text.endswith(":") and _next_is_bullet(k)
                intro_section_only = False
                if is_intro:
                    # Structure-only (no standalone atom) when it's a list HEADER by
                    # MEANING (embedding), short or long, any polarity — so neither a
                    # bare "Deliverables:" nor a long "The following items are
                    # excluded ...:" becomes a duplicate atom alongside its breadcrumb.
                    # No word-count heuristic: the semantic rule is the judge.
                    intro_section_only = is_list or self._is_framing_lead_in(text)
                    lvl = (stack[-1][0] if stack else 0) + 1
                # FRAMING LEAD-IN: a non-heading sentence that announces a following
                # list / sub-section ("...will perform the following services.").
                # It's connective tissue, not a standalone fact — so it becomes
                # structure (no wasted atom) but is LIFTED onto every child it
                # governs as lead_in context. Scoped like a heading (persists across
                # an intervening sub-heading until the section closes), not like a
                # tight list-intro (which pops on the first non-bullet).
                # NB ordering: the cheap STRUCTURAL gates (not a heading/list, a
                # bullet directly follows) run first so the EXPENSIVE semantic
                # lead-in test (it embeds the text) only runs for the handful of
                # paragraphs that actually sit just above a list — never every
                # prose paragraph.
                is_framing = (
                    lvl is None
                    and bool(text)
                    and not is_list
                    and _list_follows_within(k, 8)
                    and self._is_framing_lead_in(text)
                )
                if is_framing:
                    lvl = (stack[-1][0] if stack else 0) + 1
                if lvl is not None and text:
                    while stack and stack[-1][0] >= lvl:
                        stack.pop()
                    ancestors = [t for _, t, _, _, _ in stack if t]
                    para_section[pidx] = ancestors
                    para_lead_in[pidx] = []
                    # a heading / list-intro starts a fresh bullet context — a
                    # sub-bullet's parent must come from the SAME list, not a prior one.
                    bullet_by_level.clear()
                    if is_framing:
                        # section preamble / list lead-in: structure (no atom),
                        # lifted onto descendant list items as lead_in. Empty
                        # breadcrumb label so the sentence never enters the path.
                        heading_paras[pidx] = (lvl, ancestors)
                        structure_idxs.add(pidx)
                        stack.append((lvl, "", False, text, False))
                    else:
                        if not is_intro or intro_section_only:
                            # real heading or short label -> structure (no atom)
                            heading_paras[pidx] = (lvl, ancestors)
                            structure_idxs.add(pidx)
                        # else: long intro sentence stays an atom (not structure)
                        label = text.rstrip(":").strip() if is_intro else text
                        # CONTRADICTION GATE: a real sub-heading meaning the OPPOSITE
                        # of a "vendor will provide" preamble ("Out of Scope",
                        # "Customer Responsibilities") blocks that preamble from being
                        # lifted onto its bullets. Semantic + cached; colon list-intros
                        # ("Services include:") never block.
                        blocks = (not is_intro) and self._subsection_blocks_lift(label)
                        stack.append((lvl, label, is_intro, None, blocks))
                else:
                    # plain content: if we've left the bullet list, close any open
                    # tight list-intro sub-section(s) so a following paragraph doesn't
                    # wrongly inherit "...> PMO Responsibilities". Framing lead-ins
                    # (is_list_intro=False) are NOT popped here — they're section
                    # scoped and close only when their heading level closes.
                    if not is_list:
                        while stack and stack[-1][2]:
                            stack.pop()
                    para_section[pidx] = [t for _, t, _, _, _ in stack if t]
                    # Lift the governing preamble onto LIST ITEMS — UNLESS a
                    # contradiction subsection is active (an exclusion / other-party
                    # section), in which case a vendor "will provide" preamble must
                    # not apply to these bullets.
                    leads = (
                        [li for _, _, _, li, _ in stack if li]
                        if (is_list and not any(b for *_, b in stack))
                        else []
                    )
                    if is_list:
                        # PARENT BULLET (structural list hierarchy): a sub-bullet
                        # carries its nearest shallower bullet so it doesn't read in
                        # isolation. Always valid (local hierarchy) — independent of
                        # the contradiction gate.
                        lev = self._list_level(para)
                        if lev > 0:
                            parent = next(
                                (bullet_by_level[l] for l in range(lev - 1, -1, -1)
                                 if l in bullet_by_level),
                                None,
                            )
                            if parent:
                                leads = leads + [parent]
                        bullet_by_level[lev] = text
                        for _l in [x for x in bullet_by_level if x > lev]:
                            del bullet_by_level[_l]
                    else:
                        bullet_by_level.clear()
                    para_lead_in[pidx] = leads
            elif kind == "tbl":
                tidx += 1
                table_order[tidx] = seq
                seq += 1
                table_section[tidx] = [t for _, t, _, _, _ in stack if t]
        self._structure_idxs = structure_idxs
        self._para_lead_in = para_lead_in
        return para_section, table_section, heading_paras, para_order, table_order

    # Section-heading -> atom type. The document's OWN heading is the authority:
    # an unmatched bullet under a "Deliverables" heading IS a deliverable, under
    # "Acceptance Criteria" IS an acceptance criterion, etc. Applied only to
    # fail-open prose (the unsure bucket); the head can still override. Universal
    # (keys on the structural heading the author wrote, not on content tokens).
    _SECTION_TYPE_HINTS: tuple[tuple[str, "AtomType"], ...] = (
        ("deliverable", AtomType.deliverable),
        ("acceptance criteria", AtomType.acceptance_criterion),
        ("assumptions and dependencies", AtomType.assumption),
    )

    def _section_type_hint(self, section_path: list[str] | None) -> "AtomType | None":
        joined = " > ".join(str(s) for s in (section_path or [])).lower()
        for kw, atom_type in self._SECTION_TYPE_HINTS:
            if kw in joined:
                return atom_type
        return None

    def _emit_atoms_for_text(
        self,
        project_id: str,
        artifact_id: str,
        filename: str,
        text: str,
        paragraph_index: int | None,
        table_index: int | None,
        row: int | None,
        cell: int | None,
        tracked_change: str | None,
        heading: bool,
        tracked_index: int | None = None,
        is_list_item: bool = False,
        section_path: list[str] | None = None,
        lead_in: list[str] | None = None,
    ) -> list[EvidenceAtom]:
        # Span-provenance ledger (passive side-channel; only active when a
        # ledger is attached). Register this raw unit so the lost-content
        # report can attribute any drop to this exact GATE.
        ledger = getattr(self, "_ledger", None)
        span_id = self._span_id(
            artifact_id, paragraph_index, table_index, row, cell, tracked_change, tracked_index
        )
        if ledger is not None:
            ledger.register_span(span_id, text)

        # A heading is STRUCTURE, never a content atom — even when its title text
        # happens to match a lexical pattern (e.g. "7. Out of Scope" matches the
        # exclusion regex, "2. Sites and Scope" the scope regex). Its text is
        # already preserved as section_path on every child beneath it, so it never
        # needs to be its own atom. Drop it here (recorded) BEFORE lexical typing,
        # so a heading can't leak in as content just because of a title keyword.
        if heading and text:
            if ledger is not None:
                from app.core.span_ledger import StageKind

                ledger.record_drop(
                    span_id=span_id,
                    stage="docx_parse.heading_is_structure",
                    kind=StageKind.GATE,
                    rule="heading_not_content",
                    reason="section heading preserved as section_path on children, not a content atom",
                    raw_text=text,
                    artifact=artifact_id,
                )
            return []

        atom_types = self._classify_text(text)
        # Weak-label flag: when the lexical regex assigns a type that CONTRADICTS
        # the governing section heading (e.g. "...provides site access..." -> the
        # \baccess\b pattern says constraint, but the heading is "Assumptions and
        # Dependencies"), that word-level guess is low-trust. Don't ship it as a
        # confident fact — flag it weak_label / needs_review so it (a) stops
        # polluting the output and (b) becomes a candidate for PM hand-labelling,
        # i.e. the training queue for the eventual supervised type head.
        weak_label = False
        # Per-type weak lexical cues (bare "access" -> constraint, lone "assume"
        # -> assumption). Computed from the lexical match; only applied below to
        # types that actually came from the lexical classifier (not section_typed
        # / prose_fallback, which carry their own confidence + flags).
        weak_lexical = self._weak_lexical_types(text) if atom_types else set()
        if atom_types:
            _shint = self._section_type_hint(section_path)
            if _shint is not None and _shint not in atom_types:
                weak_label = True
        # Fail OPEN, not closed: a paragraph that matches no lexical pattern
        # but is substantive narrative prose is captured as a scope_item
        # (lower confidence, flagged) so no load-bearing fact is silently
        # dropped at parse time. Headings stay out — they are NOT facts, they
        # are structure: their text is preserved as section_path on every child
        # atom beneath them, so a heading never needs to be its own (content)
        # atom. (Matches the PDF parser, which treats headings as section
        # context only.)
        prose_fallback = False
        section_typed = False
        if not atom_types:
            # Bullet list items are deliberate, load-bearing content (deliverables,
            # assumptions, checklists) — fail OPEN regardless of length, even when
            # they are too short to pass the substantive-prose heuristic. Otherwise
            # a 3-word bullet like "Hardware order tracking" is silently dropped.
            if not heading and (self._is_substantive_prose(text) or is_list_item):
                hint = self._section_type_hint(section_path)
                if hint is not None:
                    # The governing heading names the type (Deliverables ->
                    # deliverable, etc.) — propose it instead of an unsure
                    # scope_item. Still flagged section_typed so the head may revise.
                    atom_types = [hint]
                    section_typed = True
                else:
                    atom_types = [AtomType.scope_item]
                    prose_fallback = True
            else:
                if ledger is not None:
                    from app.core.span_ledger import StageKind

                    ledger.record_drop(
                        span_id=span_id,
                        stage="docx_parse.prose_gate",
                        kind=StageKind.GATE,
                        rule="_is_substantive_prose",
                        reason=self._prose_drop_reason(text, heading=heading),
                        raw_text=text,
                        artifact=artifact_id,
                    )
                return []
        locator = {
            "paragraph_index": paragraph_index,
            "table_index": table_index,
            "row": row,
            "cell": cell,
            "tracked_change": tracked_change,
            "section_path": list(section_path) if section_path else [],
        }
        if lead_in:
            locator["lead_in"] = list(lead_in)
        source_ref = SourceRef(
            id=stable_id("src", artifact_id, paragraph_index, table_index, row, cell, tracked_change, tracked_index),
            artifact_id=artifact_id,
            artifact_type=ArtifactType.docx,
            filename=filename,
            locator=locator,
            extraction_method="docx_text_and_ooxml",
            parser_version=self.parser_version,
        )

        atoms: list[EvidenceAtom] = []
        for atom_type in atom_types:
            authority_class = AuthorityClass.contractual_scope if heading else AuthorityClass.meeting_note
            review_status = ReviewStatus.auto_accepted
            review_flags: list[str] = []
            confidence = 0.84
            if tracked_change == "deleted":
                authority_class = AuthorityClass.deleted_text
                review_status = ReviewStatus.rejected
                review_flags = ["tracked_change_deleted_text"]
                confidence = 0.2
            elif tracked_change == "inserted":
                confidence = 0.72
            elif section_typed:
                # Typed from the governing section heading (e.g. a bullet under
                # "Deliverables" -> deliverable). Moderate confidence, flagged so
                # the head can revise; authority is the written doc structure.
                authority_class = AuthorityClass.contractual_scope
                confidence = 0.7
                review_flags = ["section_typed"]
            elif prose_fallback:
                # Captured by fail-open prose rule, not a lexical match:
                # lower confidence and flag so downstream semantic stages
                # reclassify/prune. Provenance preserved, data not lost.
                confidence = 0.5
                review_flags = ["prose_fallback_capture"]
            # A type is weak when it contradicts the section heading OR rests on
            # a brittle single-word lexical cue. The latter only counts for types
            # that came straight from the lexical classifier (section_typed /
            # prose_fallback have their own provenance + confidence).
            is_weak = weak_label or (
                not section_typed and not prose_fallback and atom_type in weak_lexical
            )
            if is_weak:
                # Low-trust guess — provisional: route to review + the PM
                # labelling queue rather than ship as a confident fact.
                review_flags = review_flags + ["weak_label"]
                review_status = ReviewStatus.needs_review
                confidence = min(confidence, 0.45)
            atoms.append(
                EvidenceAtom(
                    id=stable_id(
                        "atm",
                        project_id,
                        artifact_id,
                        atom_type.value,
                        text,
                        paragraph_index,
                        table_index,
                        row,
                        cell,
                        tracked_change,
                        tracked_index,
                    ),
                    project_id=project_id,
                    artifact_id=artifact_id,
                    atom_type=atom_type,
                    raw_text=text,
                    normalized_text=normalize_text(text),
                    value={"text": text, "tracked_change": tracked_change, "prose_fallback": prose_fallback},
                    entity_keys=self._extract_entity_keys(text),
                    source_refs=[source_ref],
                    authority_class=authority_class,
                    confidence=confidence,
                    review_status=review_status,
                    review_flags=review_flags,
                    parser_version=self.parser_version,
                )
            )
        if ledger is not None and atoms:
            ledger.mark_represented(span_id)
        return atoms

    # -- never-detected recovery: content python-docx's body view can't see --
    def _recover_nested_region_atoms(
        self,
        *,
        project_id: str,
        artifact_id: str,
        filename: str,
        document: Any,
        already_emitted: set[str],
    ) -> list[EvidenceAtom]:
        """Recover paragraphs/tables that live inside content controls
        (``w:sdt``) or textboxes (``w:txbxContent``).

        ``Document.paragraphs`` / ``Document.tables`` only enumerate the body's
        *direct* ``w:p`` / ``w:tbl`` children, so anything nested under an sdt
        or a textbox (executive summaries, revision histories, sales/customer
        contact blocks) is silently invisible to the main parse loop — the
        *never-detected* loss class. We walk the body lxml tree, find those
        orphaned regions, and emit atoms for them:

        * orphan **tables** -> one ``scope_item`` per row, emitted
          UNCONDITIONALLY (mirroring the main table-row path) so contact rows
          with no scope verb / digit survive the prose gate that would
          otherwise drop them;
        * orphan **paragraphs** -> routed through ``_emit_atoms_for_text``.

        Dedup is by normalized text against ``already_emitted`` so VML/DrawingML
        ``mc:AlternateContent`` fallbacks (the same textbox stored twice) don't
        double-count. This is a structural, modality-agnostic walk — no
        keyword lists, no per-deal tuning.
        """
        body = getattr(getattr(document, "element", None), "body", None)
        if body is None:
            return []

        def _localname(tag: Any) -> str:
            t = str(tag)
            return t.rsplit("}", 1)[-1] if "}" in t else t

        def _nested_ancestors(el: Any) -> set[str]:
            out: set[str] = set()
            try:
                ancestors = el.iterancestors()
            except Exception:  # pragma: no cover - non-lxml element
                return out
            for a in ancestors:
                out.add(_localname(a.tag))
            return out

        def _el_text(el: Any) -> str:
            parts = [t.text for t in el.iter(f"{{{WORD_NS['w']}}}t") if t.text]
            return " ".join(parts).strip()

        ledger = getattr(self, "_ledger", None)
        atoms: list[EvidenceAtom] = []
        tbl_tag = f"{{{WORD_NS['w']}}}tbl"
        p_tag = f"{{{WORD_NS['w']}}}p"
        tr_tag = f"{{{WORD_NS['w']}}}tr"
        tc_tag = f"{{{WORD_NS['w']}}}tc"

        base_tbl = len(_all_tables(document))
        base_p = len(_all_paragraphs(document))

        # --- orphan tables (e.g. the SALES / CUSTOMER CONTACTS blocks) ------
        orphan_tbl_idx = 0
        for tbl in body.iter(tbl_tag):
            anc = _nested_ancestors(tbl)
            # sdt content controls are now walked by the MAIN path
            # (_iter_block_items descends into them) — with header-row skipping,
            # cross-type dedup and reading order. This lossy fallback must only
            # cover TEXT BOXES (txbxContent), which the main path can't reach;
            # re-scraping sdt here re-emitted bare header rows and duplicated the
            # executive summary (the exact issues Benjamin reported).
            if "txbxContent" not in anc:
                continue
            if "tbl" in anc:  # nested table; outer one already handled it
                continue
            for row_idx, tr in enumerate(tbl.iter(tr_tag)):
                cells = [_el_text(tc) for tc in tr.iter(tc_tag)]
                cells = [c for c in cells if c]
                if not cells:
                    continue
                row_text = " | ".join(cells)
                norm = normalize_text(row_text)
                if norm in already_emitted:
                    continue
                already_emitted.add(norm)
                t_index = base_tbl + orphan_tbl_idx
                span_id = self._span_id(artifact_id, None, t_index, row_idx, None, None, None)
                if ledger is not None:
                    ledger.register_span(span_id, row_text)
                atom_id = stable_id("atm", artifact_id, "docx_nested_row", t_index, row_idx, row_text)
                src = SourceRef(
                    id=stable_id("src", atom_id),
                    artifact_id=artifact_id,
                    artifact_type=ArtifactType.docx,
                    filename=filename,
                    locator={
                        "table_index": t_index,
                        "row": row_idx,
                        "extraction": "docx_nested_region_v1",
                        "region": "sdt" if "sdtContent" in anc else "textbox",
                    },
                    extraction_method="docx_nested_region_v1",
                    parser_version=self.parser_version,
                )
                atoms.append(EvidenceAtom(
                    id=atom_id,
                    project_id=project_id,
                    artifact_id=artifact_id,
                    atom_type=AtomType.scope_item,
                    raw_text=row_text,
                    normalized_text=norm,
                    value={
                        "kind": "table_row",
                        "region": "sdt" if "sdtContent" in anc else "textbox",
                        "cells": {f"col_{i}": v for i, v in enumerate(cells)},
                    },
                    entity_keys=self._extract_entity_keys(row_text),
                    source_refs=[src],
                    receipts=[],
                    authority_class=AuthorityClass.contractual_scope,
                    confidence=0.7,
                    confidence_raw=0.7,
                    calibrated_confidence=0.7,
                    review_status=ReviewStatus.auto_accepted,
                    review_flags=["recovered_nested_region"],
                    parser_version=self.parser_version,
                ))
                if ledger is not None:
                    ledger.mark_represented(span_id)
            orphan_tbl_idx += 1

        # --- orphan paragraphs (executive summary, revision notes, etc.) ----
        orphan_p_idx = 0
        for p in body.iter(p_tag):
            anc = _nested_ancestors(p)
            # sdt content controls are now walked by the MAIN path
            # (_iter_block_items descends into them) — with header-row skipping,
            # cross-type dedup and reading order. This lossy fallback must only
            # cover TEXT BOXES (txbxContent), which the main path can't reach;
            # re-scraping sdt here re-emitted bare header rows and duplicated the
            # executive summary (the exact issues Benjamin reported).
            if "txbxContent" not in anc:
                continue
            if "tbl" in anc:  # already covered by its table region above
                continue
            txt = _el_text(p)
            if not txt:
                continue
            norm = normalize_text(txt)
            if norm in already_emitted:
                continue
            already_emitted.add(norm)
            recovered = self._emit_atoms_for_text(
                project_id=project_id,
                artifact_id=artifact_id,
                filename=filename,
                text=txt,
                paragraph_index=base_p + orphan_p_idx,
                table_index=None,
                row=None,
                cell=None,
                tracked_change=None,
                heading=False,
            )
            for a in recovered:
                a.review_flags = list(a.review_flags) + ["recovered_nested_region"]
            atoms.extend(recovered)
            orphan_p_idx += 1

        return atoms

    def _emit_embedded_media_markers(
        self,
        *,
        project_id: str,
        artifact_id: str,
        filename: str,
        path: Path,
    ) -> list[EvidenceAtom]:
        """Emit a located ``image_marker`` / ``embedded_object_marker`` atom for
        every binary part in the .docx zip (``word/media/*``,
        ``word/embeddings/*``).

        Detection of a binary region is *total and guaranteed* even when we
        cannot yet read it: the region never silently vanishes — it becomes a
        located marker the PM sees ("image/object here, vision/OLE pass
        required"). Extraction quality is a separate, improving frontier. The
        ``region_ref`` in each marker's value lets the content census reconcile
        the region as MARKED rather than UNCOVERED.
        """
        atoms: list[EvidenceAtom] = []
        try:
            zf = zipfile.ZipFile(path)
        except Exception:  # pragma: no cover - not a zip / unreadable
            return []
        with zf:
            for name in sorted(zf.namelist()):
                if name.startswith("word/media/"):
                    kind, atype = "image_marker", "image"
                elif name.startswith("word/embeddings/"):
                    kind, atype = "embedded_object_marker", "embedded object"
                else:
                    continue
                rel = name[len("word/"):]
                try:
                    size = zf.getinfo(name).file_size
                except KeyError:  # pragma: no cover
                    size = 0
                marker_text = (
                    f"[{atype.capitalize()} awaiting OCR / vision / OLE extraction] "
                    f"{rel} in {filename} — {size:,} bytes. A vision or embedded-"
                    f"object pass is required to recover its content."
                )
                atom_id = stable_id("atm", artifact_id, "docx_media_marker", rel)
                src = SourceRef(
                    id=stable_id("src", atom_id),
                    artifact_id=artifact_id,
                    artifact_type=ArtifactType.docx,
                    filename=filename,
                    locator={"region_ref": rel, "extraction": "docx_media_marker_v1"},
                    extraction_method="docx_media_marker_v1",
                    parser_version=self.parser_version,
                )
                atoms.append(EvidenceAtom(
                    id=atom_id,
                    project_id=project_id,
                    artifact_id=artifact_id,
                    atom_type=AtomType.open_question,
                    raw_text=marker_text,
                    normalized_text=normalize_text(marker_text),
                    value={"kind": kind, "region_ref": rel, "size_bytes": size},
                    entity_keys=[],
                    source_refs=[src],
                    receipts=[],
                    authority_class=AuthorityClass.meeting_note,
                    confidence=0.5,
                    confidence_raw=0.5,
                    calibrated_confidence=0.5,
                    review_status=ReviewStatus.needs_review,
                    review_flags=["binary_region_marker"],
                    parser_version=self.parser_version,
                ))
        return atoms
