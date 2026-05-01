from __future__ import annotations

import tempfile
import zipfile
from pathlib import Path
from random import Random
from xml.etree import ElementTree as ET

from docx import Document
from openpyxl import Workbook

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": WORD_NS}


def _inject_deleted_tracked_change(docx_path: Path, deleted_text: str) -> None:
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        with zipfile.ZipFile(docx_path, "r") as source:
            source.extractall(temp_root)
        doc_xml = temp_root / "word" / "document.xml"
        root = ET.fromstring(doc_xml.read_bytes())
        body = root.find(".//w:body", NSMAP)
        if body is None:
            raise ValueError("Invalid DOCX structure: missing word body")
        p = ET.Element(f"{{{WORD_NS}}}p")
        r = ET.SubElement(p, f"{{{WORD_NS}}}r")
        del_node = ET.SubElement(r, f"{{{WORD_NS}}}del")
        del_run = ET.SubElement(del_node, f"{{{WORD_NS}}}r")
        del_text_node = ET.SubElement(del_run, f"{{{WORD_NS}}}delText")
        del_text_node.text = deleted_text
        body.append(p)
        doc_xml.write_bytes(ET.tostring(root, encoding="utf-8", xml_declaration=True))
        with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as target:
            for file in temp_root.rglob("*"):
                if file.is_file():
                    target.write(file, file.relative_to(temp_root).as_posix())


def write_spreadsheet_fixture(
    path: Path,
    *,
    included_site: str,
    excluded_site: str,
    scoped_device: str,
    main_qty: str,
    west_qty: str,
    access_constraint: str,
    mutation: str,
    rng: Random,
) -> None:
    quantity_header = {
        "header_qty_synonym": "Qty",
        "header_count_synonym": "Count",
        "header_hash_synonym": "#",
        "header_quantity_synonym": "Quantity",
    }.get(mutation, "Quantity")
    west_site = {
        "site_alias_dash": "West-Wing",
        "site_alias_bldg": "Bldg A West",
    }.get(mutation, excluded_site)
    if mutation == "mixed_case":
        headers = ["sItE", "fLoOr", "dEvIcE", quantity_header, "AcCeSs WiNdOw", "sCoPe"]
    else:
        headers = ["Site", "Floor", "Device", quantity_header, "Access Window", "Scope"]
    if mutation == "extra_irrelevant_columns":
        headers = headers + ["Owner", "Ticket", "Comment"]

    rows = [
        [included_site, "1", scoped_device, main_qty, "Weekdays 8am-5pm", "Install"],
        [west_site, "2", scoped_device, west_qty, access_constraint, "Install"],
    ]
    if mutation == "row_order_shuffle":
        rng.shuffle(rows)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "site_roster"
    if mutation == "shifted_header_row":
        sheet.append(["export generated", "", "", "", "", ""])
    if mutation == "blank_rows":
        sheet.append(["", "", "", "", "", ""])
    sheet.append(headers)
    for row in rows:
        if mutation == "extra_irrelevant_columns":
            row = row + ["ops", "TKT-1001", "n/a"]
        sheet.append(row)
        if mutation == "blank_rows":
            sheet.append(["", "", "", "", "", "", "", "", ""])
    if mutation == "subtotal_rows":
        sheet.append(["Subtotal", "", "", "91", "", ""])
    sheet.append(["TOTAL", "", "", "91", "", ""])
    if mutation == "hidden_columns":
        sheet.column_dimensions["G"].hidden = True
    workbook.save(path)


def write_email_fixture(
    path: Path,
    *,
    excluded_site: str,
    included_site: str,
    access_constraint: str,
    mutation: str,
) -> None:
    current_line = f"Please remove {excluded_site} from scope and hold off until confirmed."
    quoted_line = f"> Please include {excluded_site} in the camera rollout."
    if mutation == "internal_only_note":
        sender = "From: pm.internal@purtera.com"
    else:
        sender = "From: jane.customer@example.com"
    prefix = "FW: " if mutation == "forwarding_prefix" else ""
    body = [
        sender,
        "Sent: 2026-02-01 09:15",
        f"Subject: {prefix}Scope direction",
        "",
        current_line,
        f"{included_site} requires {access_constraint}.",
        "",
    ]
    if mutation in {"quoted_on_date", "multiple_quoted_levels"}:
        body.extend(
            [
                "On 2026-01-28, Jane Customer wrote:",
                quoted_line,
            ]
        )
        if mutation == "multiple_quoted_levels":
            body.extend(
                [
                    "On 2026-01-20, PM wrote:",
                    ">> Do not proceed at West Wing until approval.",
                ]
            )
    else:
        body.extend(
            [
                "-----Original Message-----",
                "From: jane.customer@example.com",
                "Sent: 2026-01-20 11:00",
                "Subject: Earlier scope",
                quoted_line,
            ]
        )
    path.write_text("\n".join(body) + "\n", encoding="utf-8")


def write_docx_fixture(
    path: Path,
    *,
    included_site: str,
    excluded_site: str,
    scoped_device: str,
    mutation: str,
) -> None:
    doc = Document()
    heading = "SCOPE EXCLUSIONS" if mutation == "all_caps_exclusion" else "Statement of Work Draft"
    doc.add_heading(heading, level=1)
    if mutation == "scope_in_table":
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Scope"
        table.rows[0].cells[1].text = f"Install {scoped_device} at {included_site} and {excluded_site}."
        table.rows[1].cells[0].text = "Notes"
        table.rows[1].cells[1].text = "Customer is responsible for lift access."
    else:
        doc.add_paragraph(f"Scope includes installation of {scoped_device} at {included_site} and {excluded_site}.")
    exclusion_line = f"{excluded_site} is excluded from scope pending customer confirmation."
    if mutation == "all_caps_exclusion":
        exclusion_line = exclusion_line.upper()
    doc.add_paragraph(exclusion_line)
    doc.add_heading("Assumptions", level=2)
    doc.add_paragraph("Assumption: Existing network closets remain available.")
    doc.add_heading("Customer Responsibility", level=2)
    doc.add_paragraph("Customer is responsible for lift access and escort coordination.")
    doc.save(path)
    _inject_deleted_tracked_change(path, f"Install AV displays at {excluded_site}.")


def write_transcript_fixture(
    path: Path,
    *,
    included_site: str,
    excluded_site: str,
    scoped_device: str,
    access_constraint: str,
    open_question: str,
    mutation: str,
) -> None:
    if mutation == "meeting_notes_sections":
        content = "\n".join(
            [
                "Decisions:",
                f"- Customer: remove {excluded_site} from scope for now.",
                "Action Items:",
                f"- Customer to provide access escort plan for {included_site}.",
                "Open Questions:",
                f"- {open_question}?",
                f"- Maybe add 5 {scoped_device}s after pilot.",
            ]
        )
    elif mutation == "unknown_speaker":
        content = "\n".join(
            [
                "[00:00:01] Unknown: We may add 5 cameras at main campus.",
                "[00:00:50] Jane Customer: Do not proceed at West Wing.",
                "[00:01:20] Purtera PM: Open question: MDF badge access?",
            ]
        )
    else:
        content = "\n".join(
            [
                "[00:00:01] Purtera PM: Starting kickoff for camera rollout.",
                f"[00:01:10] Jane Customer: {included_site} requires {access_constraint}.",
                f"[00:02:00] Jane Customer: Please remove {excluded_site} from scope for now.",
                f"[00:03:25] Purtera PM: Open question: {open_question}?",
                f"[00:04:00] Jane Customer: maybe add 5 {scoped_device}s after pilot.",
            ]
        )
    path.write_text(content + "\n", encoding="utf-8")


def write_quote_fixture(
    path: Path,
    *,
    scoped_device: str,
    vendor_total: str,
    mutation: str,
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "quote"
    part_number = "cam-ip-001" if mutation == "part_number_variation" else "CAM-IP-001"
    description = "IP Cam" if mutation == "device_alias_ip_cam" else scoped_device
    quantity = vendor_total if mutation != "quantity_as_string" else str(vendor_total)
    unit_price = "$300.00" if mutation == "unit_price_dollar" else "300.00"
    lead_time = "" if mutation == "lead_time_missing" else "2 weeks"
    sheet.append(["Part Number", "Description", "Quantity", "Unit Price", "Lead Time"])
    sheet.append([part_number, description, quantity, unit_price, lead_time])
    if mutation == "quote_total_row":
        sheet.append(["TOTAL", "", vendor_total, "", ""])
    workbook.save(path)
