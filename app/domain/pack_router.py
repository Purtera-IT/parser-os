"""Project-level domain-pack router.

Picks the best ``DomainPack`` for a project directory.  Order of
precedence:

1. **Explicit override** — ``--domain-pack <pack_id>`` from the CLI
   (handled at the caller level by passing ``domain_pack`` to
   ``load_domain_pack``); this module only fires when the caller
   passes ``None``.
2. **Project config** — ``<project>/project.yaml`` with
   ``service_line:`` or ``domain_pack:`` keys.
3. **SOURCE_NOTES.md** — a free-text declaration like ``Service line:
   security_camera`` (case-insensitive, multiple synonyms accepted).
4. **Filename heuristics** — keywords in the artifact filenames that
   strongly imply a service line (e.g. ``cabling`` → ``copper_cabling``,
   ``camera`` → ``security_camera``, ``wireless`` → ``wireless``).
5. **Content scoring** — read the first ~5 KB of each parseable file
   (PDF, XLSX, txt) and tally hits against each pack's
   ``device_aliases`` and ``entity_types[].aliases``; pick the
   highest-scoring pack with a margin over default.
6. **Fallback** — ``default_pack``.

The router is conservative: when scores tie or are too low to be
confident, it returns the default pack so downstream stages stay
predictable.  All decisions are reported via ``RoutingDecision`` so
operators can see *why* a pack was selected.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml

from app.domain.loader import (
    DEFAULT_PACK_ID,
    _candidate_pack_path,
    load_domain_pack,
)
from app.domain.schemas import DomainPack


# Service-line synonyms commonly seen in RFP / SOURCE_NOTES text.
# Each canonical pack id maps to the surface forms a human might write.
_SERVICE_LINE_SYNONYMS: dict[str, list[str]] = {
    "security_camera": [
        "security_camera",
        "security camera",
        "video surveillance",
        "vms",
        "cctv",
        "video_surveillance",
        "surveillance",
    ],
    "access_control": [
        "access_control",
        "access control",
        "eacs",
        "card access",
        "door access",
        "physical access control",
    ],
    "wireless": [
        "wireless",
        "wifi",
        "wi-fi",
        "wlan",
        "e-rate wireless",
    ],
    "networking": [
        "networking",
        "managed vpn",
        "vpn",
        "lan",
        "wan",
        "ethernet switching",
        "structured network",
    ],
    "copper_cabling": [
        "copper_cabling",
        "copper cabling",
        "cabling",
        "structured cabling",
        "cat6",
        "cat 6",
        "cat-6",
        "ip phone cabling",
    ],
    "av": [
        "av",
        "audio visual",
        "audio-visual",
        "audiovisual",
        "audio_visual",
        "av integration",
        "video conferencing",
    ],
    "bms": [
        "bms",
        "building management",
        "building automation",
        "bas",
        "ddc",
        "facility controls",
        "integrated automation",
    ],
    "paging": [
        "paging",
        "mass notification",
        "emergency notification",
        "mass communication",
        "emns",
        "wide-area public address",
    ],
    "fire_safety": [
        "fire_safety",
        "fire safety",
        "fire alarm",
        "fire detection",
        "life safety",
        "nfpa 72",
    ],
    "das": [
        "das",
        "distributed antenna system",
        "in-building das",
        "public safety das",
    ],
    "electrical": [
        "electrical",
        "electrical power",
        "panel schedule",
        "power distribution",
    ],
    "itad": [
        "itad",
        "it asset disposition",
        "asset disposition",
        "secure data destruction",
        "e-waste recycling",
    ],
}

# Filename keywords that bias pack selection.  Used as a tiebreaker
# when content scoring is ambiguous.
_FILENAME_KEYWORDS: dict[str, list[str]] = {
    "security_camera": ["camera", "surveillance", "cctv", "vms", "video_analytics"],
    "access_control": ["access_control", "card_reader", "door_access", "eacs"],
    "wireless": ["wireless", "wifi", "wlan", "e-rate", "erate", "ap_install"],
    "networking": ["network", "vpn", "lan", "wan", "switch", "managed_vpn"],
    "copper_cabling": ["cabling", "cat6", "cat-6", "structured_cabling", "ip_phone"],
    "av": ["audio_visual", "audiovisual", "av_", "boardroom", "conference_room"],
    "bms": ["bas", "bms", "building_automation", "hvac_controls", "integrated_automation"],
    "paging": ["mass_notification", "emergency_notification", "paging", "mass_comm"],
    "fire_safety": ["fire_alarm", "fire_safety", "nfpa72", "smoke_detector"],
    "das": ["das", "distributed_antenna", "ibw", "in_building_wireless"],
    "electrical": ["panel_schedule", "electrical_distribution"],
    "itad": ["itad", "asset_disposition", "asset_inventory", "data_destruction"],
}


@dataclass
class RoutingDecision:
    """Why we picked a particular pack — for telemetry and review folders."""

    pack_id: str
    source: str  # 'explicit' | 'project_yaml' | 'source_notes' | 'filename' | 'content' | 'default'
    confidence: float
    rationale: str
    alternatives: list[tuple[str, float]] = field(default_factory=list)


def _read_project_yaml(project_dir: Path) -> dict[str, Any] | None:
    for name in ("project.yaml", "project.yml"):
        candidate = project_dir / name
        if candidate.is_file():
            try:
                payload = yaml.safe_load(candidate.read_text(encoding="utf-8"))
                if isinstance(payload, dict):
                    return payload
            except Exception:  # pragma: no cover — never fail compile on config read
                return None
    return None


def _read_source_notes(project_dir: Path) -> str:
    candidate = project_dir / "SOURCE_NOTES.md"
    if not candidate.is_file():
        return ""
    try:
        return candidate.read_text(encoding="utf-8", errors="ignore")
    except Exception:  # pragma: no cover
        return ""


def _service_line_to_pack_id(service_line: str) -> str | None:
    """Normalize free-text service line → canonical pack id."""
    if not service_line:
        return None
    needle = service_line.strip().lower().replace("-", "_").replace(" ", "_")
    # Direct match against canonical names
    for pack_id in _SERVICE_LINE_SYNONYMS:
        if needle == pack_id:
            return pack_id
    # Synonym match (using normalized synonyms)
    raw = service_line.strip().lower()
    for pack_id, surfaces in _SERVICE_LINE_SYNONYMS.items():
        for surface in surfaces:
            if raw == surface or surface in raw:
                return pack_id
    return None


def _route_from_project_yaml(project_dir: Path) -> RoutingDecision | None:
    payload = _read_project_yaml(project_dir)
    if not payload:
        return None
    pack_id_explicit = payload.get("domain_pack") or payload.get("pack_id")
    if isinstance(pack_id_explicit, str) and pack_id_explicit.strip():
        pack_id = pack_id_explicit.strip()
        return RoutingDecision(
            pack_id=pack_id,
            source="project_yaml",
            confidence=1.0,
            rationale=f"project.yaml declared domain_pack: {pack_id}",
        )
    service_line = payload.get("service_line")
    if isinstance(service_line, str) and service_line.strip():
        pack_id = _service_line_to_pack_id(service_line) or DEFAULT_PACK_ID
        return RoutingDecision(
            pack_id=pack_id,
            source="project_yaml",
            confidence=0.95,
            rationale=f"project.yaml declared service_line: {service_line!r} → {pack_id}",
        )
    return None


_SOURCE_NOTES_PATTERNS = [
    # **Service line**: `copper_cabling`
    re.compile(r"\*\*service\s*line\*\*\s*:\s*`?([a-z0-9_\- ]{2,40})`?", re.IGNORECASE),
    # Service line: copper_cabling
    re.compile(r"^[\s\-\*]*service\s*line\s*[:=]\s*([a-z0-9_\- ]{2,40})", re.IGNORECASE | re.MULTILINE),
    # service_line: copper_cabling
    re.compile(r"\bservice_line\s*[:=]\s*([a-z0-9_\- ]{2,40})", re.IGNORECASE),
]


def _route_from_source_notes(project_dir: Path) -> RoutingDecision | None:
    notes = _read_source_notes(project_dir)
    if not notes:
        return None
    for pattern in _SOURCE_NOTES_PATTERNS:
        match = pattern.search(notes)
        if not match:
            continue
        raw_value = match.group(1).strip().strip("`'\"")
        pack_id = _service_line_to_pack_id(raw_value)
        if pack_id:
            return RoutingDecision(
                pack_id=pack_id,
                source="source_notes",
                confidence=0.90,
                rationale=f"SOURCE_NOTES.md declared service line {raw_value!r} → {pack_id}",
            )
    return None


def _filename_hint_score(filenames: list[str]) -> dict[str, float]:
    scores: dict[str, float] = {}
    haystack = " ".join(name.lower().replace(".", "_").replace("-", "_") for name in filenames)
    for pack_id, keywords in _FILENAME_KEYWORDS.items():
        hits = sum(1 for kw in keywords if kw in haystack)
        if hits:
            scores[pack_id] = float(hits)
    return scores


def _docx_text_preview(path: Path, max_chars: int) -> str:
    """First paragraphs + table cells of a .docx, capped at ``max_chars``.

    Office-doc deals (the common case for real RFP/SOW packages) carry
    their authoritative scope in .docx/.xlsx, not in side-car .txt. If we
    can't read them, content scoring sees nothing and routing falls back
    to incidental filename/customer-name tokens — exactly how a
    residential TV-install deal got routed as a datacenter/cabling job.
    """
    try:
        from docx import Document  # python-docx
    except Exception:
        return ""
    try:
        doc = Document(str(path))
    except Exception:
        return ""
    parts: list[str] = []
    total = 0
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue
        parts.append(t)
        total += len(t)
        if total >= max_chars:
            return "\n".join(parts)[:max_chars]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if not t:
                    continue
                parts.append(t)
                total += len(t)
                if total >= max_chars:
                    return "\n".join(parts)[:max_chars]
    return "\n".join(parts)[:max_chars]


def _xlsx_text_preview(path: Path, max_chars: int) -> str:
    """Text of the first worksheet's cells, capped at ``max_chars``."""
    try:
        from openpyxl import load_workbook
    except Exception:
        return ""
    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
    except Exception:
        return ""
    parts: list[str] = []
    total = 0
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is None:
                        continue
                    t = str(cell).strip()
                    if not t:
                        continue
                    parts.append(t)
                    total += len(t)
                    if total >= max_chars:
                        return " ".join(parts)[:max_chars]
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return " ".join(parts)[:max_chars]


def _read_text_preview(path: Path, max_bytes: int = 6 * 1024) -> str:
    """Best-effort text preview for filename + lightweight content scoring.

    Reads plain text directly; extracts a cheap preview from .docx/.xlsx so
    office-doc deals route on their actual scope content. Never raises —
    routing must not fail on a single unreadable artifact.
    """
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")[:max_bytes]
        except Exception:
            return ""
    if suffix == ".docx":
        return _docx_text_preview(path, max_bytes)
    if suffix == ".xlsx":
        return _xlsx_text_preview(path, max_bytes)
    return ""


_ALIAS_RE_CACHE: dict[str, re.Pattern[str]] = {}

# Single-word aliases that are common English / cross-trade vocabulary.
# They appear in nearly any deal document regardless of trade, so they
# must never be a routing signal — "area"/"battery" made a TV-install
# SOW score for fire_safety, "tape" for itad.  Multi-word aliases that
# merely contain these ("fire panel", "server rack") are unaffected;
# the stoplist only suppresses the bare single-token form.
_GENERIC_ALIAS_STOPWORDS: frozenset[str] = frozenset({
    "area", "level", "site", "zone", "room", "floor", "unit", "units",
    "building", "battery", "tape", "channel", "board", "point", "line",
    "system", "device", "section", "phase", "field", "panel", "outlet",
})


def _alias_in_text(alias: str, haystack: str) -> bool:
    """Word-boundary alias match.

    Plain substring matching produced false positives that flipped
    routing — "ap" matched *ap*proximately, "tr" matched cen*tr*e,
    "site"/"level" matched freely — so a TV-install deal scored higher
    for ``wireless`` than ``av``. Anchoring on word boundaries kills
    those incidental collisions (and lets distinctive 2-char aliases
    like "tv" count safely). Single-character aliases stay too
    ambiguous to trust and are skipped.
    """
    if len(alias) < 2:
        return False
    pat = _ALIAS_RE_CACHE.get(alias)
    if pat is None:
        pat = re.compile(r"(?<!\w)" + re.escape(alias) + r"(?!\w)")
        _ALIAS_RE_CACHE[alias] = pat
    return pat.search(haystack) is not None


# Markers that a blob is authoritative scope prose (a SOW / scope of
# work / project overview) vs. a commercial deal-kit / pricing summary.
# Domain routing must follow the scope document, not the money sheet —
# a deal-kit xlsx mentions generic infra vocabulary ("cat6", "switch")
# regardless of the actual trade.
_SCOPE_AUTHORITY_MARKERS = (
    "statement of work",
    "scope of work",
    "project overview",
    "scope:",
    "work will be performed",
    "deliverables",
    "responsibilities",
)
_COMMERCIAL_MARKERS = (
    "deal kit",
    "net margin",
    "total deal revenue",
    "total deal cost",
    "total deal margin",
    "margin %",
    "gross margin",
    "rate card",
    "bill of materials",
)
_SCOPE_AUTHORITY_WEIGHT = 3.0
_COMMERCIAL_AUTHORITY_WEIGHT = 0.4
_DEFAULT_AUTHORITY_WEIGHT = 1.0


def _blob_authority(filename: str, text: str) -> float:
    """Weight a content blob by how authoritative it is for *domain*.

    A statement-of-work / scope document is the real signal; a
    financial deal-kit summary's incidental infra words are noise.
    """
    lo = text.lower()
    fn = filename.lower()
    if any(m in lo for m in _SCOPE_AUTHORITY_MARKERS):
        return _SCOPE_AUTHORITY_WEIGHT
    if any(m in lo for m in _COMMERCIAL_MARKERS) or "deal kit" in fn:
        return _COMMERCIAL_AUTHORITY_WEIGHT
    return _DEFAULT_AUTHORITY_WEIGHT


def _content_score(project_dir: Path) -> dict[str, float]:
    """Score each pack against the artifact filenames + textual content.

    For PDFs we'd ideally parse the first page text, but that's
    expensive at routing time.  Instead we score:
    - filenames (cheap, almost always informative)
    - any plain-text / markdown / csv files we can read directly

    Each pack's ``device_aliases`` keys + each ``entity_types[].aliases``
    entry contribute to the haystack-needle scoring.

    We require an ``artifacts/`` subdirectory to exist; if it doesn't,
    we return empty scores so callers fall through to default_pack
    instead of accidentally scoring against project metadata files.
    """
    scores: dict[str, float] = {}
    pack_ids = list(_SERVICE_LINE_SYNONYMS.keys())
    # Cheap pass — filenames.  Restricted to ``artifacts/`` so we
    # never accidentally score against ``labels/`` gold standards or
    # documentation in the project root.
    artifact_dir = project_dir / "artifacts"
    if not artifact_dir.is_dir():
        return scores
    filenames: list[str] = []
    text_blobs: list[tuple[str, float]] = []  # (lowercased text, authority weight)
    skip_names = {
        "source_notes.md",
        "readme.md",
        "license",
        "license.md",
        "license.txt",
        "structured.json",  # parser-emitted derived artifact
        "structured.md",
        "orbitbrief.input.json",
        "orbitbrief.input.md",
    }
    skip_substrings = ("gold_standard", ".gold.", "_gold.", "_review.", ".review.")
    skip_dir_suffixes = (".derived",)
    skip_dir_names = {".orbitbrief", ".cache", ".git", ".github", "labels", "node_modules", "__pycache__"}
    for path in artifact_dir.rglob("*"):
        if not path.is_file():
            continue
        # Skip files inside parser-managed derived / output dirs
        try:
            rel_parts = path.relative_to(artifact_dir).parts
        except ValueError:
            rel_parts = (path.name,)
        if any(part.endswith(skip_dir_suffixes) for part in rel_parts[:-1]):
            continue
        if any(part.lower() in skip_dir_names for part in rel_parts[:-1]):
            continue
        name_lower = path.name.lower()
        if name_lower in skip_names:
            continue
        if any(token in name_lower for token in skip_substrings):
            continue
        filenames.append(path.name)
        text_preview = _read_text_preview(path)
        if text_preview:
            text_blobs.append(
                (text_preview.lower(), _blob_authority(path.name, text_preview))
            )
    if not filenames:
        return scores

    # Filename keyword hits get weighted higher than content hits
    # (filenames are deliberate signals, content can have incidental
    # word collisions).  We also skip wide reference packs in the
    # filename pass — same rationale as the content pass below.
    filename_scores = _filename_hint_score(filenames)
    for pack_id, sc in filename_scores.items():
        try:
            pack = load_domain_pack(pack_id)
            service_lines = [sl.lower() for sl in (pack.service_lines or [])]
            if pack.reference_ontology_path:
                continue
            if not service_lines or service_lines == ["default"]:
                continue
        except Exception:  # pragma: no cover
            pass
        scores[pack_id] = scores.get(pack_id, 0.0) + sc * 2.0

    if text_blobs:
        # Build {pack_id: set(aliases)} for every narrow pack, plus a
        # frequency map of how many packs share each alias.  Aliases
        # shared across many packs ("switch", "site", "rack", "panel")
        # are generic low-voltage vocabulary and carry almost no
        # routing signal — a deal-kit financial summary mentions them
        # regardless of trade.  We down-weight them by inverse pack
        # frequency so distinctive aliases ("tv", "display mount")
        # dominate.  This is what keeps a TV-install deal from routing
        # to ``wireless`` just because its money sheet says "switch".
        pack_aliases: dict[str, set[str]] = {}
        alias_freq: dict[str, int] = {}
        for pack_id in pack_ids:
            try:
                pack = load_domain_pack(pack_id)
            except Exception:  # pragma: no cover
                continue
            service_lines = [sl.lower() for sl in (pack.service_lines or [])]
            if pack.reference_ontology_path:
                continue
            if not service_lines or service_lines == ["default"]:
                continue
            aliases: set[str] = set()
            for surfaces in (pack.device_aliases or {}).values():
                for surface in surfaces:
                    if surface:
                        aliases.add(surface.lower())
            for entity in pack.entity_types or []:
                for alias in entity.aliases or []:
                    if alias:
                        aliases.add(alias.lower())
            aliases = {
                a for a in aliases if a not in _GENERIC_ALIAS_STOPWORDS
            }
            pack_aliases[pack_id] = aliases
            for alias in aliases:
                alias_freq[alias] = alias_freq.get(alias, 0) + 1

        n_packs = len(pack_aliases) or 1
        for pack_id, aliases in pack_aliases.items():
            hits = 0.0
            for alias in aliases:
                freq = alias_freq.get(alias, 1)
                # Drop aliases that appear in more than half the packs —
                # they're shared vocabulary, not a discriminating signal.
                if freq > n_packs / 2:
                    continue
                # Strongest evidence comes from the most authoritative
                # blob that mentions this alias: a SOW outweighs a deal
                # kit.  Take the max authority among matching blobs.
                best_authority = 0.0
                for blob_text, authority in text_blobs:
                    if authority > best_authority and _alias_in_text(alias, blob_text):
                        best_authority = authority
                if best_authority <= 0.0:
                    continue
                # Multi-word phrases are far stronger evidence than a
                # bare token; inverse-frequency rewards distinctiveness.
                weight = 2.0 if " " in alias else 1.0
                hits += (weight / freq) * best_authority
            if hits:
                scores[pack_id] = scores.get(pack_id, 0.0) + hits
    return scores


def _route_from_content(project_dir: Path) -> RoutingDecision | None:
    scores = _content_score(project_dir)
    if not scores:
        return None
    ranked = sorted(scores.items(), key=lambda kv: (-kv[1], kv[0]))
    best_id, best_score = ranked[0]
    runner_score = ranked[1][1] if len(ranked) > 1 else 0.0
    margin = best_score - runner_score
    # Require both a positive score and a measurable margin to commit
    # to a non-default pack.  Otherwise fall through to default_pack.
    if best_score < 2.0 or margin < 1.0:
        return None
    confidence = min(0.85, 0.55 + 0.05 * best_score)
    return RoutingDecision(
        pack_id=best_id,
        source="content",
        confidence=confidence,
        rationale=(
            f"content scoring picked {best_id} (score {best_score:.1f}, "
            f"margin {margin:.1f} over runner-up)"
        ),
        alternatives=[(pid, score) for pid, score in ranked[1:6]],
    )


def auto_route_pack(
    project_dir: Path,
    *,
    explicit: str | Path | None = None,
) -> tuple[DomainPack, RoutingDecision]:
    """Pick the best DomainPack for ``project_dir``.

    ``explicit`` (if given) overrides all auto-detection — it's whatever
    the user passed via ``--domain-pack``.  Returns the loaded pack
    plus the decision rationale for telemetry.
    """
    if explicit is not None:
        # User-supplied pack id wins.  Validate by loading; on failure
        # fall through to default_pack.
        try:
            pack = load_domain_pack(explicit)
            pack_path = _candidate_pack_path(explicit)
            return pack, RoutingDecision(
                pack_id=pack.pack_id,
                source="explicit",
                confidence=1.0,
                rationale=f"user override --domain-pack={explicit} (resolved to {pack_path.name})",
            )
        except Exception as exc:  # pragma: no cover
            return load_domain_pack(None), RoutingDecision(
                pack_id=DEFAULT_PACK_ID,
                source="default",
                confidence=0.30,
                rationale=f"failed to load explicit pack {explicit!r}: {exc}; using default_pack",
            )

    for resolver in (_route_from_project_yaml, _route_from_source_notes, _route_from_content):
        decision = resolver(project_dir)
        if decision is None:
            continue
        try:
            pack = load_domain_pack(decision.pack_id)
        except Exception:  # pragma: no cover — bad pack id, fall through
            continue
        return pack, decision

    return load_domain_pack(None), RoutingDecision(
        pack_id=DEFAULT_PACK_ID,
        source="default",
        confidence=0.50,
        rationale="no project.yaml / SOURCE_NOTES service line / content scoring signal; using default_pack",
    )


__all__ = ["auto_route_pack", "RoutingDecision"]
