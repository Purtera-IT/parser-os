from __future__ import annotations

import csv
import json
import re
import unicodedata
from pathlib import Path
from typing import Any, Literal

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import column_index_from_string

from app.core.normalizers import normalize_text
from app.core.segments import ArtifactSegment
from app.core.schemas import EvidenceAtom, EvidenceReceipt, SourceRef


# ────────────── workbook cache (insanity-perf) ──────────────
# openpyxl ``load_workbook`` is the dominant cost in source_replay
# when a project has many spreadsheet atoms. Each atom triggered an
# independent load of the same .xlsx file — on STRESS_MULTI_CAM,
# 1349 spreadsheet atoms × ~100 ms per load = ~135 s of source_replay
# wall time. Cache by (path, mtime, size) so each unique workbook is
# loaded at most once per compile, while invalidating if the file
# changes on disk between calls.
_WORKBOOK_CACHE: dict[tuple[str, float, int], Any] = {}


def _load_workbook_cached(path: Path):
    """Return a memoized workbook for ``path``. Cached by (path, mtime,
    size) so editing the file between compiles invalidates the entry."""
    try:
        st = path.stat()
        key = (str(path), st.st_mtime, st.st_size)
    except FileNotFoundError:
        return None
    cached = _WORKBOOK_CACHE.get(key)
    if cached is not None:
        return cached
    try:
        wb = load_workbook(path, data_only=True)
    except Exception:
        return None
    _WORKBOOK_CACHE[key] = wb
    # Soft cap: keep the cache to ~16 distinct workbooks. Realistic
    # projects compile <16 unique .xlsx files.
    if len(_WORKBOOK_CACHE) > 16:
        # Drop the oldest entry. dict preserves insertion order so the
        # first key is the oldest.
        oldest = next(iter(_WORKBOOK_CACHE))
        if oldest != key:
            _WORKBOOK_CACHE.pop(oldest, None)
            # An evicted workbook's worksheets may be GC'd and their id() reused;
            # drop the memoized max_rows so a future worksheet can't read a stale
            # value under a recycled id.
            _MAXROW_CACHE.clear()
    return wb


# Memoized ``worksheet.max_row``. openpyxl recomputes ``max_row`` by scanning
# the worksheet's *entire* cell set on every single access. source_replay
# verifies one row per atom, calling ``max_row`` once per atom purely for an
# out-of-range check — so a big sheet with many atoms made this O(atoms × cells)
# (observed: a 39k-atom deal grinding for >8 min in ``max_row`` alone). The value
# can't change during a read-only verification pass, so cache it per worksheet
# identity. The workbook (and thus its worksheets) is held alive in
# ``_WORKBOOK_CACHE`` while in use, so ``id()`` is stable for the cache lifetime;
# we clear this alongside any workbook-cache eviction to avoid id reuse.
_MAXROW_CACHE: dict[int, int] = {}


def _cached_max_row(worksheet: Any) -> int:
    """``worksheet.max_row`` computed at most once per worksheet (see note)."""
    wid = id(worksheet)
    mr = _MAXROW_CACHE.get(wid)
    if mr is None:
        mr = worksheet.max_row or 0
        _MAXROW_CACHE[wid] = mr
    return mr


def clear_workbook_cache() -> None:  # pragma: no cover — used by tests
    _WORKBOOK_CACHE.clear()
    _MAXROW_CACHE.clear()


# ────────────── docx whole-document text cache ──────────────
# A large share of DOCX atoms carry a locator type the precise
# paragraph/table verifier can't resolve (tracked-change anchors,
# section/run offsets, list-item paths the parser emits but the
# verifier doesn't model). Those used to return "unsupported" — a
# receipt that proves nothing. Reading the whole document body once
# (cached) lets us at least confirm the atom's text genuinely appears
# in the file, upgrading "unsupported" to a (lower-precision but real)
# "verified" provenance receipt. Cache by (path, mtime, size).
_DOCX_TEXT_CACHE: dict[tuple[str, float, int], str] = {}


def _docx_full_text(path: Path) -> str:
    """Return the concatenated text of every paragraph and table cell
    in a .docx, memoized by (path, mtime, size)."""
    try:
        st = path.stat()
        key = (str(path), st.st_mtime, st.st_size)
    except FileNotFoundError:
        return ""
    cached = _DOCX_TEXT_CACHE.get(key)
    if cached is not None:
        return cached
    try:
        document = Document(path)
    except Exception:
        return ""
    parts: list[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for trow in table.rows:
            for tcell in trow.cells:
                if tcell.text:
                    parts.append(tcell.text)
    # python-docx's .paragraphs/.tables views omit text inside content
    # controls (w:sdt) and other body structures, so an atom whose source
    # lives in an SDT (e.g. a SOW intro clause) is invisible to the views
    # above and its receipt can't be verified. Harvest every raw text run
    # (w:t) from the body XML as the ground-truth "does this text exist in
    # the doc" corpus. Duplication is harmless for a contains-search.
    try:
        from docx.oxml.ns import qn

        for _t in document.element.body.iter(qn("w:t")):
            if _t.text:
                parts.append(_t.text)
    except Exception:
        pass
    body = "\n".join(parts)
    _DOCX_TEXT_CACHE[key] = body
    if len(_DOCX_TEXT_CACHE) > 16:
        oldest = next(iter(_DOCX_TEXT_CACHE))
        if oldest != key:
            _DOCX_TEXT_CACHE.pop(oldest, None)
    return body


def clear_docx_text_cache() -> None:  # pragma: no cover — used by tests
    _DOCX_TEXT_CACHE.clear()


def _match_excerpt(atom: EvidenceAtom, body: str, *, window: int = 90) -> str | None:
    """Best-effort short excerpt of ``body`` around where the atom text
    appears, for the receipt's ``extracted_snippet``. Returns ``None``
    when no anchor token can be located (the caller stores nothing
    rather than the whole document)."""
    if not body:
        return None
    haystack = body.lower()
    anchors = sorted(
        (a for a in _atom_match_candidates(atom) if a),
        key=len,
        reverse=True,
    )
    for anchor in anchors:
        idx = haystack.find(anchor)
        if idx >= 0:
            start = max(0, idx - window)
            end = min(len(body), idx + len(anchor) + window)
            prefix = "…" if start > 0 else ""
            suffix = "…" if end < len(body) else ""
            return f"{prefix}{body[start:end].strip()}{suffix}"
    return None


def _whole_document_text_fallback(
    atom: EvidenceAtom,
    source_ref: SourceRef,
    body: str,
    *,
    label: str,
) -> EvidenceReceipt | None:
    """Upgrade an otherwise-unsupported receipt to ``verified`` when the
    atom's text genuinely appears anywhere in the document body. Returns
    ``None`` when there is no match, so the caller keeps its original
    (unsupported/failed) verdict. This never *downgrades* — it is only
    ever called on paths that would otherwise be unsupported."""
    if body and _snippet_matches_atom(atom, body):
        return _receipt(
            atom,
            source_ref,
            "verified",
            f"{label} (whole-document text fallback; locator imprecise)",
            _match_excerpt(atom, body),
        )
    return None


def _replay_norm(text: str) -> str:
    """Normalize text for replay matching.

    Strips Unicode combining marks (so "café" matches "cafe", and
    "M\xa0Smith" matches "M Smith") then defers to the canonical
    ``normalize_text``. The PR8 spec calls this out specifically: a
    lot of "failed" replay receipts in the corpus were just NFKD
    drift between parser-extracted text and the spreadsheet/PDF
    re-read.
    """
    text = unicodedata.normalize("NFKD", text or "")
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return normalize_text(text)

VERIFIER_VERSION = "source_replay_v2"
_STOP_WORDS = {
    "the",
    "and",
    "for",
    "with",
    "from",
    "that",
    "this",
    "will",
    "was",
    "are",
    "after",
    "before",
    "please",
}


def _receipt(
    atom: EvidenceAtom,
    source_ref: SourceRef,
    replay_status: Literal["verified", "failed", "unsupported"],
    reason: str,
    extracted_snippet: str | None = None,
) -> EvidenceReceipt:
    return EvidenceReceipt(
        atom_id=atom.id,
        artifact_id=source_ref.artifact_id,
        filename=source_ref.filename,
        source_ref_id=source_ref.id,
        replay_status=replay_status,
        extracted_snippet=extracted_snippet,
        locator=dict(source_ref.locator),
        reason=reason,
        verifier_version=VERIFIER_VERSION,
    )


def _atom_match_candidates(atom: EvidenceAtom) -> list[str]:
    """Phrase-level candidates the snippet must contain to count as a hit.

    A candidate is "evidence-grade" only if it is *substantive enough*
    that finding it inside the snippet is meaningful on its own. Single
    short tokens like ``customer``, ``vendor``, ``owner`` (which leak
    out of ``atom.value['owner']`` / role tags) cause false-positive
    line-range matches: any speaker line that happens to contain the
    word "Customer" would falsely verify an atom whose actual text was
    deleted from the file. We require either whitespace (i.e. a real
    phrase) or length >= 12 (i.e. a long-enough single token like a
    site/part identifier) to qualify.
    """
    candidates = [normalize_text(atom.raw_text), normalize_text(atom.normalized_text)]
    for key in atom.entity_keys:
        _, _, tail = key.partition(":")
        if tail:
            candidates.append(normalize_text(tail.replace("_", " ")))
    for value in atom.value.values():
        if isinstance(value, (str, int, float)):
            candidates.append(normalize_text(str(value)))
    return sorted({
        c for c in candidates
        if c and (
            " " in c                       # multi-word phrases are always evidence
            or any(ch.isdigit() for ch in c)  # part numbers, quantities, measurements
            or len(c) >= 12                # long unique tokens (UUIDs, hashes, slugs)
        )
    })


def _important_terms(atom: EvidenceAtom) -> list[str]:
    tokens = re.findall(r"[a-z0-9]+", normalize_text(atom.normalized_text))
    terms = [t for t in tokens if (t.isdigit() or len(t) >= 4) and t not in _STOP_WORDS]
    if not terms:
        terms = [t for t in tokens if t not in _STOP_WORDS]
    deduped: list[str] = []
    for term in terms:
        if term not in deduped:
            deduped.append(term)
    return deduped[:8]


def _snippet_matches_atom(atom: EvidenceAtom, snippet: str) -> bool:
    snippet_norm = _replay_norm(snippet)
    if not snippet_norm:
        return False
    if atom.normalized_text and _replay_norm(atom.normalized_text) in snippet_norm:
        return True
    candidates = [_replay_norm(c) for c in _atom_match_candidates(atom)]
    if any(c and c in snippet_norm for c in candidates):
        return True
    terms = [_replay_norm(t) for t in _important_terms(atom)]
    terms = [t for t in terms if t]
    if not terms:
        return False
    matched = sum(1 for term in terms if term in snippet_norm)
    required = 1 if len(terms) <= 2 else 2
    return matched >= required


def _spreadsheet_full_row_text(
    path: Path, sheet: str | None, row_number: int
) -> str:
    """Return ``"v1 | v2 | v3 ..."`` for every non-empty cell in
    ``row_number`` (1-based). Used as a fallback when the cited
    columns alone don't satisfy ``_snippet_matches_atom``: the row
    might have been authored with the relevant fact in a column the
    parser didn't cite (e.g. the parser cited Severity, but the
    important text is in the Mitigation column on the same row).
    """
    suffix = path.suffix.lower()
    parts: list[str] = []
    if suffix == ".xlsx":
        wb = _load_workbook_cached(path)
        if wb is None:
            return ""
        ws = (
            wb[sheet]
            if isinstance(sheet, str) and sheet in wb.sheetnames
            else wb.active
        )
        max_row = ws.max_row or 0
        if max_row and row_number > max_row:
            return ""
        for cell in ws[row_number]:
            if cell.value is None:
                continue
            text = str(cell.value).strip()
            if text:
                parts.append(text)
    else:
        delimiter = "," if suffix == ".csv" else None
        try:
            with path.open("r", encoding="utf-8", errors="ignore", newline="") as fh:
                if delimiter is None:
                    rows = [
                        re.split(r"[,\t|]", line.rstrip("\n"))
                        for line in fh.readlines()
                    ]
                else:
                    rows = list(csv.reader(fh))
        except Exception:
            return ""
        if row_number > len(rows):
            return ""
        for cell in rows[row_number - 1]:
            text = str(cell or "").strip()
            if text:
                parts.append(text)
    return " | ".join(parts)


def _verify_spreadsheet_row(atom: EvidenceAtom, source_ref: SourceRef, path: Path) -> EvidenceReceipt:
    locator = source_ref.locator
    sheet = locator.get("sheet")
    row_number = locator.get("row")
    columns = locator.get("columns") or {}
    if not isinstance(row_number, int) or row_number < 1:
        return _receipt(atom, source_ref, "failed", "Spreadsheet locator missing valid row number")

    row_values: dict[str, str] = {}
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        # Cached loader (insanity-perf). On STRESS_MULTI_CAM this
        # alone cuts source_replay from ~135 s to ~3 s.
        workbook = _load_workbook_cached(path)
        if workbook is None:
            return _receipt(atom, source_ref, "failed", f"Spreadsheet file unreadable: {path}")
        worksheet = workbook[sheet] if isinstance(sheet, str) and sheet in workbook.sheetnames else workbook.active
        max_row = _cached_max_row(worksheet)
        if max_row and row_number > max_row:
            return _receipt(atom, source_ref, "failed", f"Spreadsheet row {row_number} out of range")
        for key, col_letter in columns.items():
            try:
                col_index = column_index_from_string(str(col_letter))
                value = worksheet.cell(row=row_number, column=col_index).value
                row_values[str(key)] = "" if value is None else str(value)
            except Exception:
                continue
    else:
        delimiter = "," if suffix == ".csv" else None
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            if delimiter is None:
                rows = [re.split(r"[,\t|]", line.rstrip("\n")) for line in handle.readlines()]
            else:
                rows = list(csv.reader(handle))
        if row_number > len(rows):
            return _receipt(atom, source_ref, "failed", f"Spreadsheet row {row_number} out of range")
        row = rows[row_number - 1]
        for key, col_letter in columns.items():
            try:
                index = column_index_from_string(str(col_letter)) - 1
                row_values[str(key)] = row[index] if index < len(row) else ""
            except Exception:
                continue

    # Sort columns alphabetically so the snippet text is deterministic across
    # runs.  Parser-side dict insertion order has been observed to drift (e.g.
    # openpyxl backed by C accelerators), which silently broke output_signature
    # reproducibility — see manifest.compute_output_signature.
    snippet_parts = [
        f"{str(key).replace('_', ' ').title()}={str(value).strip()}"
        for key, value in sorted(row_values.items())
    ]
    extracted_snippet = " | ".join(part for part in snippet_parts if part and not part.endswith("="))
    if not extracted_snippet:
        return _receipt(atom, source_ref, "failed", "Spreadsheet row found but no cited cells were readable")
    if _snippet_matches_atom(atom, extracted_snippet):
        return _receipt(atom, source_ref, "verified", "Spreadsheet row and cited cells verified", extracted_snippet)

    # PR8 — full-row fallback. The row exists but the cited cells
    # alone don't carry the atom's text. This commonly happens when
    # the parser cited only one column (e.g. Severity) but the atom
    # text was authored from a different column on the same row
    # (e.g. Mitigation). Read the full row and try once more before
    # giving up — true mismatches still fail.
    full_row = _spreadsheet_full_row_text(path, sheet if isinstance(sheet, str) else None, row_number)
    if full_row and _snippet_matches_atom(atom, full_row):
        return _receipt(
            atom,
            source_ref,
            "verified",
            "Spreadsheet row verified via full-row fallback",
            full_row,
        )

    # Insanity-pass — value-dict match. The quote_parser synthesizes
    # atoms with text like ``"Line item TesiraFORTÉ X 400-SUP Biamp ..."``
    # whose "Line item" prefix doesn't appear in any cell, so the
    # cell-based matcher fails even though the row IS the source.
    # If the atom carries a value-dict with cell-derived fields
    # (part_number, description, manufacturer, etc.), check that at
    # least one of those values is present in the full row text.
    if full_row and isinstance(atom.value, dict):
        cell_like_keys = (
            "part_number", "description", "manufacturer", "vendor",
            "model", "asset_id", "serial", "section", "item",
            "support_note", "lead_time",
        )
        norm_row = _replay_norm(full_row)
        for key in cell_like_keys:
            v = atom.value.get(key)
            if not isinstance(v, str) or len(v.strip()) < 4:
                continue
            if _replay_norm(v) in norm_row:
                return _receipt(
                    atom,
                    source_ref,
                    "verified",
                    f"Spreadsheet row verified via value-dict field {key!r}",
                    full_row,
                )

    # Constraint-keyword fallback: xlsx parser emits synthetic atoms
    # like "After-hours access constraint" / "Site access constraint"
    # whose canonical phrase doesn't appear in any cell. Verify by
    # matching the atom's value.constraint_type against keyword
    # signals in the actual row text. This keeps receipts honest
    # (the row contains "after-hours" → after_hours constraint) while
    # not failing verification on parser-synthesized headlines.
    if full_row and isinstance(atom.value, dict):
        ctype = atom.value.get("constraint_type")
        if isinstance(ctype, str) and ctype:
            ctype_signals = {
                "after_hours": ("after-hours", "after hours", "nights only", "weekends only"),
                "access": ("badge", "escort", "ceiling access"),
                "lift": ("lift required", "elevator", "customer provides lift"),
                "certification": ("certification required", "certify", "test standard"),
            }
            keywords = ctype_signals.get(ctype, ())
            full_norm = _replay_norm(full_row)
            if any(_replay_norm(kw) in full_norm for kw in keywords):
                return _receipt(
                    atom,
                    source_ref,
                    "verified",
                    f"Spreadsheet row verified via constraint_type={ctype!r}",
                    full_row,
                )

    return _receipt(
        atom,
        source_ref,
        "failed",
        "Spreadsheet row exists but cited cells do not match atom content",
        extracted_snippet,
    )


def _verify_line_range(atom: EvidenceAtom, source_ref: SourceRef, path: Path) -> EvidenceReceipt:
    locator = source_ref.locator
    line_start = locator.get("line_start")
    line_end = locator.get("line_end")
    if not isinstance(line_start, int) or not isinstance(line_end, int) or line_start < 1 or line_end < line_start:
        # No usable line locator. Rather than fail outright, confirm the
        # atom text against the whole file — a real provenance receipt
        # even though the precise line offsets are absent.
        body = path.read_text(encoding="utf-8", errors="ignore")
        fallback = _whole_document_text_fallback(
            atom, source_ref, body, label="Text locator missing line range"
        )
        return fallback or _receipt(atom, source_ref, "unsupported", "Line-range locator missing or invalid")
    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    if line_end > len(lines):
        return _receipt(atom, source_ref, "failed", f"Line range {line_start}-{line_end} is out of bounds")
    snippet = "\n".join(lines[line_start - 1 : line_end])
    if _snippet_matches_atom(atom, snippet):
        return _receipt(atom, source_ref, "verified", "Line-range snippet verified", snippet)
    return _receipt(atom, source_ref, "failed", "Line-range snippet did not match atom content", snippet)


def _verify_docx_locator(atom: EvidenceAtom, source_ref: SourceRef, path: Path) -> EvidenceReceipt:
    locator = source_ref.locator
    tracked_change = locator.get("tracked_change")
    if tracked_change == "deleted":
        # A tracked deletion is a special provenance case: the text is
        # struck through, not asserted. python-docx still surfaces the
        # struck text in the body, so a whole-document fallback would
        # falsely "verify" it — deliberately keep tracked deletions
        # unsupported (OOXML tracked-change replay isn't modeled here).
        return _receipt(
            atom,
            source_ref,
            "unsupported",
            "Tracked deletion replay via OOXML is not enabled in this verifier",
        )

    document = Document(path)
    paragraph_index = locator.get("paragraph_index")
    table_index = locator.get("table_index")
    row = locator.get("row")
    cell = locator.get("cell")

    if isinstance(paragraph_index, int):
        if 0 <= paragraph_index < len(document.paragraphs):
            snippet = document.paragraphs[paragraph_index].text
            if _snippet_matches_atom(atom, snippet):
                return _receipt(atom, source_ref, "verified", "DOCX paragraph locator verified", snippet)
        # Index out of range or text mismatch. The parser numbers paragraphs
        # including table/SDT paragraphs that python-docx's body-only
        # .paragraphs view omits, so the index drifts (e.g. a SOW intro in a
        # content control indexed past len(paragraphs)). Fall back to
        # whole-document text: the receipt verifies iff the atom's text
        # exists ANYWHERE in the source — integrity preserved (truly-absent
        # text still fails), false "failed" from index drift eliminated.
        fb = _whole_document_text_fallback(
            atom, source_ref, _docx_full_text(path),
            label="DOCX paragraph index drift",
        )
        if fb is not None:
            return fb
        snippet = (
            document.paragraphs[paragraph_index].text
            if 0 <= paragraph_index < len(document.paragraphs)
            else ""
        )
        return _receipt(atom, source_ref, "failed", "DOCX paragraph did not match atom content", snippet)

    if isinstance(table_index, int) and isinstance(row, int) and isinstance(cell, int):
        if table_index < 0 or table_index >= len(document.tables):
            return _receipt(atom, source_ref, "failed", f"Table index {table_index} out of range")
        table = document.tables[table_index]
        if row < 0 or row >= len(table.rows):
            return _receipt(atom, source_ref, "failed", f"Table row {row} out of range")
        if cell < 0 or cell >= len(table.rows[row].cells):
            return _receipt(atom, source_ref, "failed", f"Table cell {cell} out of range")
        snippet = table.rows[row].cells[cell].text
        if _snippet_matches_atom(atom, snippet):
            return _receipt(atom, source_ref, "verified", "DOCX table locator verified", snippet)
        return _receipt(atom, source_ref, "failed", "DOCX table cell did not match atom content", snippet)

    fallback = _whole_document_text_fallback(
        atom, source_ref, _docx_full_text(path), label="DOCX locator unmodeled"
    )
    return fallback or _receipt(
        atom, source_ref, "unsupported", "DOCX locator type is not currently supported"
    )


def _verify_pdf_block(atom: EvidenceAtom, source_ref: SourceRef, path: Path) -> EvidenceReceipt:
    """Verify an OrbitBrief PDF atom against the persisted structured doc.

    The OrbitBrief PDF parser writes a deterministic
    ``<stem>.derived/structured.json`` next to every parsed PDF.  We
    look up the atom's ``block_id`` (or ``row_index`` for table cells,
    or ``bullet_path`` for bullets) inside that doc and check that the
    atom's text actually came from there.  This makes PDF receipts
    first-class — same status as XLSX rows and DOCX paragraphs.

    Schematic atoms (legend / detection / warning) instead point at a
    bbox region of the page.  The locator carries ``page``, ``bbox``,
    ``bbox_units="pdf_points"``, and ``crop_sha256``; we re-render the
    page at a fixed DPI, crop the bbox, and hash the pixels to confirm
    nothing has drifted.  This path is taken whenever the locator has
    no ``block_id`` but does have a bbox + crop hash.
    """
    locator = source_ref.locator
    block_id = locator.get("block_id")
    has_bbox = _locator_is_bbox_crop(locator)
    if not block_id and has_bbox:
        return _verify_pdf_bbox_crop(atom, source_ref, path)
    if not block_id and not has_bbox:
        return _receipt(atom, source_ref, "unsupported", "PDF locator missing block_id")
    # Both block_id and a bbox+crop_sha256 are present. Both must
    # verify; if either fails, the receipt fails so the locator's
    # over-broad provenance can't quietly half-verify.
    if block_id and has_bbox:
        bbox_receipt = _verify_pdf_bbox_crop(atom, source_ref, path)
        if bbox_receipt.replay_status != "verified":
            return bbox_receipt
        # fall through to block_id check below
    structured = _load_structured_doc(path)
    if structured is None:
        return _receipt(
            atom,
            source_ref,
            "unsupported",
            "Structured PDF doc not present (run orbitbrief_pdf parser first)",
        )
    block = _find_block(structured, block_id)
    if block is None:
        return _receipt(atom, source_ref, "failed", f"block_id {block_id} not found in structured doc")

    snippet = _block_to_snippet(block, locator)
    if not snippet:
        return _receipt(atom, source_ref, "failed", "Located PDF block contained no text")
    if _snippet_matches_atom(atom, snippet):
        return _receipt(atom, source_ref, "verified", "PDF block locator verified", snippet)
    return _receipt(atom, source_ref, "failed", "PDF block found but text did not match atom", snippet)


def _structured_doc_path_for(pdf_path: Path) -> Path:
    return pdf_path.with_name(f"{pdf_path.stem}.derived") / "structured.json"


def _load_structured_doc(pdf_path: Path) -> dict[str, Any] | None:
    out = _structured_doc_path_for(pdf_path)
    if not out.is_file():
        return None
    try:
        return json.loads(out.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None


def _find_block(structured: dict[str, Any], block_id: str) -> dict[str, Any] | None:
    for page in structured.get("pages", []) or []:
        match = _walk_sections_for_block(page.get("sections", []) or [], block_id)
        if match is not None:
            return match
    return None


def _walk_sections_for_block(sections: list[dict[str, Any]], block_id: str) -> dict[str, Any] | None:
    for section in sections:
        for block in section.get("blocks", []) or []:
            if block.get("id") == block_id:
                return block
        nested = _walk_sections_for_block(section.get("subsections", []) or [], block_id)
        if nested is not None:
            return nested
    return None


def _block_to_snippet(block: dict[str, Any], locator: dict[str, Any]) -> str:
    kind = block.get("kind")
    if kind == "paragraph":
        return str(block.get("text") or "")
    if kind == "note":
        return str(block.get("text") or "")
    if kind == "bullet_list":
        bullet_path = locator.get("bullet_path")
        if isinstance(bullet_path, list) and bullet_path:
            item = _bullet_at_path(block.get("items", []) or [], bullet_path)
            if item is not None:
                return str(item.get("text") or "")
        if locator.get("bullet_role") == "intro":
            return str(block.get("intro") or "")
        # Fallback: stitch the whole list together so the snippet still
        # contains the atom text somewhere.
        parts: list[str] = []
        intro = block.get("intro")
        if intro:
            parts.append(str(intro))
        parts.extend(_flatten_bullets(block.get("items", []) or []))
        return " | ".join(parts)
    if kind == "table":
        row_index = locator.get("row_index")
        rows = block.get("rows", []) or []
        if isinstance(row_index, int) and 0 <= row_index < len(rows):
            cells = rows[row_index]
            if isinstance(cells, dict):
                return " | ".join(f"{k}: {v}" for k, v in cells.items() if v is not None)
        return str(block.get("raw_text") or "")
    return ""


def _bullet_at_path(items: list[dict[str, Any]], path: list[int]) -> dict[str, Any] | None:
    if not path:
        return None
    current = items
    node: dict[str, Any] | None = None
    for index in path:
        if not isinstance(index, int) or index < 0 or index >= len(current):
            return None
        node = current[index]
        if not isinstance(node, dict):
            return None
        current = node.get("children", []) or []
    return node


def _locator_is_bbox_crop(locator: dict[str, Any]) -> bool:
    bbox = locator.get("bbox")
    if not (isinstance(bbox, (list, tuple)) and len(bbox) == 4):
        return False
    if locator.get("bbox_units") != "pdf_points":
        return False
    return bool(locator.get("crop_sha256"))


def _verify_pdf_bbox_crop(atom: EvidenceAtom, source_ref: SourceRef, path: Path) -> EvidenceReceipt:
    """Verify a schematic PDF atom by re-rendering its bbox and hashing pixels.

    Deterministic by construction: fixed DPI, fixed PyMuPDF render
    parameters, and the hash is namespaced with ``crop_sha256_of_pixels``
    so dimension differences cannot collide.  Returns ``verified`` only
    when the recomputed hash equals the stored ``crop_sha256``.
    """
    from app.parsers.schematic_models import SCHEMATIC_REPLAY_DPI, crop_sha256_of_pixels

    locator = source_ref.locator
    page_index = locator.get("page")
    if not isinstance(page_index, int) or page_index < 0:
        return _receipt(atom, source_ref, "failed", "Schematic locator missing valid page index")
    bbox = locator.get("bbox")
    if not (isinstance(bbox, (list, tuple)) and len(bbox) == 4):
        return _receipt(atom, source_ref, "failed", "Schematic locator missing valid bbox")
    try:
        x0, y0, x1, y1 = (float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3]))
    except (TypeError, ValueError):
        return _receipt(atom, source_ref, "failed", "Schematic locator bbox is not numeric")
    if not (x1 > x0 and y1 > y0):
        return _receipt(atom, source_ref, "failed", "Schematic locator bbox is not strictly positive")
    expected_hash = str(locator.get("crop_sha256") or "")
    if not expected_hash:
        return _receipt(atom, source_ref, "failed", "Schematic locator missing crop_sha256")

    try:
        import fitz  # type: ignore[import-not-found]
    except ImportError:  # pragma: no cover — pymupdf is in pyproject deps
        return _receipt(atom, source_ref, "unsupported", "PyMuPDF unavailable for schematic replay")

    try:
        doc = fitz.open(str(path))
    except Exception as exc:  # pragma: no cover
        return _receipt(atom, source_ref, "failed", f"Could not open PDF for replay: {exc}")
    try:
        if page_index >= doc.page_count:
            return _receipt(
                atom,
                source_ref,
                "failed",
                f"Page index {page_index} out of range (page_count={doc.page_count})",
            )
        page = doc.load_page(page_index)
        # Clamp bbox to page bounds so a malformed-but-positive locator
        # (e.g. coordinates that drift outside the page after rounding,
        # or a legend bbox that grew past the page edge) cannot throw
        # inside ``get_pixmap``. We only clamp when the entire bbox
        # intersects the page; an entirely off-page bbox stays an
        # explicit failure.
        page_rect = page.rect
        px0 = max(float(page_rect.x0), min(x0, float(page_rect.x1)))
        py0 = max(float(page_rect.y0), min(y0, float(page_rect.y1)))
        px1 = max(float(page_rect.x0), min(x1, float(page_rect.x1)))
        py1 = max(float(page_rect.y0), min(y1, float(page_rect.y1)))
        if not (px1 > px0 and py1 > py0):
            return _receipt(
                atom,
                source_ref,
                "failed",
                "Schematic bbox is entirely outside the page rectangle",
            )
        clip = fitz.Rect(px0, py0, px1, py1)
        zoom = SCHEMATIC_REPLAY_DPI / 72.0
        matrix = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=matrix, clip=clip, alpha=False, colorspace=fitz.csRGB)
        actual_hash = crop_sha256_of_pixels(pix.samples, pix.width, pix.height, pix.n)
    except Exception as exc:  # pragma: no cover
        return _receipt(atom, source_ref, "failed", f"Crop render failed: {exc}")
    finally:
        try:
            doc.close()
        except Exception:  # pragma: no cover
            pass

    if actual_hash == expected_hash:
        return _receipt(
            atom,
            source_ref,
            "verified",
            "Schematic bbox crop hash verified",
            extracted_snippet=f"crop_sha256={actual_hash}",
        )
    return _receipt(
        atom,
        source_ref,
        "failed",
        f"Schematic bbox crop hash mismatch (expected {expected_hash[:12]}…, got {actual_hash[:12]}…)",
        extracted_snippet=f"crop_sha256={actual_hash}",
    )


def _flatten_bullets(items: list[dict[str, Any]]) -> list[str]:
    out: list[str] = []
    for item in items:
        text = item.get("text")
        if text:
            out.append(str(text))
        children = item.get("children")
        if children:
            out.extend(_flatten_bullets(children))
    return out


def replay_source_ref(atom: EvidenceAtom, source_ref: SourceRef, artifact_paths: dict[str, Path]) -> EvidenceReceipt:
    path = artifact_paths.get(source_ref.artifact_id)
    if path is None or not path.exists():
        return _receipt(atom, source_ref, "unsupported", "Original artifact file is not available for replay")

    suffix = path.suffix.lower()
    try:
        if source_ref.locator.get("row") is not None and source_ref.locator.get("columns"):
            return _verify_spreadsheet_row(atom, source_ref, path)
        if source_ref.locator.get("line_start") is not None and source_ref.locator.get("line_end") is not None:
            return _verify_line_range(atom, source_ref, path)
        if suffix == ".docx":
            return _verify_docx_locator(atom, source_ref, path)
        if suffix == ".pdf":
            return _verify_pdf_block(atom, source_ref, path)
        if suffix in {".txt", ".md", ".eml", ".vtt", ".srt", ".json"}:
            return _verify_line_range(atom, source_ref, path)
    except Exception as exc:  # pragma: no cover
        return _receipt(atom, source_ref, "failed", f"Source replay failed with exception: {exc}")

    return _receipt(atom, source_ref, "unsupported", "No verifier available for this locator/artifact type")


def replay_atom_receipts(atom: EvidenceAtom, artifact_paths: dict[str, Path]) -> list[EvidenceReceipt]:
    receipts = [replay_source_ref(atom, source_ref, artifact_paths) for source_ref in atom.source_refs]
    return sorted(receipts, key=lambda receipt: receipt.source_ref_id)


def attach_receipts_to_atoms(atoms: list[EvidenceAtom], artifact_paths: dict[str, Path]) -> list[EvidenceAtom]:
    for atom in atoms:
        atom.receipts = replay_atom_receipts(atom, artifact_paths)
    return atoms


def summarize_receipts(atoms: list[EvidenceAtom]) -> dict[str, int]:
    summary = {"verified": 0, "unsupported": 0, "failed": 0}
    for atom in atoms:
        for receipt in atom.receipts:
            if receipt.replay_status in summary:
                summary[receipt.replay_status] += 1
    return summary


def verify_segment(segment: ArtifactSegment, artifact_path: Path) -> EvidenceReceipt:
    artifact_paths = {segment.artifact_id: artifact_path}
    synthetic_atom = EvidenceAtom.model_validate(
        {
            "id": f"segment_atom_{segment.id}",
            "project_id": segment.project_id,
            "artifact_id": segment.artifact_id,
            "atom_type": "entity",
            "raw_text": segment.text,
            "normalized_text": segment.normalized_text,
            "value": {"segment_type": segment.segment_type},
            "entity_keys": [],
            "source_refs": [segment.source_ref.model_dump(mode="json")],
            "authority_class": "machine_extractor",
            "confidence": 1.0,
            "review_status": "auto_accepted",
            "review_flags": [],
            "parser_version": segment.source_ref.parser_version,
        }
    )
    return replay_source_ref(synthetic_atom, segment.source_ref, artifact_paths)
