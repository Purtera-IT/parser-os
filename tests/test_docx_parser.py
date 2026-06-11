from __future__ import annotations

from pathlib import Path

from docx import Document

from app.core.schemas import AtomType, AuthorityClass, ReviewStatus
from app.parsers.docx_parser import DocxParser
from scripts.make_demo_fixtures import create_demo_project


def test_docx_parser_with_tracked_deletion(tmp_path: Path) -> None:
    base_root = tmp_path / "repo"
    base_root.mkdir(parents=True, exist_ok=True)
    project_dir = create_demo_project(base_root)
    file_path = project_dir / "sow_draft.docx"

    atoms = DocxParser().parse_artifact(
        project_id="proj_1",
        artifact_id="art_docx_1",
        path=file_path,
    )

    assert atoms
    assert all(atom.source_refs for atom in atoms)

    scope_atoms = [a for a in atoms if a.atom_type == AtomType.scope_item]
    assert any("scope includes installation of ip cameras" in a.normalized_text for a in scope_atoms)

    exclusion_atoms = [a for a in atoms if a.atom_type == AtomType.exclusion]
    assert any("av displays are excluded from scope" in a.normalized_text for a in exclusion_atoms)

    constraint_atoms = [a for a in atoms if a.atom_type == AtomType.constraint]
    assert any("customer is responsible for providing lift access" in a.normalized_text for a in constraint_atoms)

    deleted_atoms = [a for a in atoms if "install av displays in conference rooms" in a.normalized_text]
    assert deleted_atoms
    assert all(a.authority_class == AuthorityClass.deleted_text for a in deleted_atoms)
    assert all(a.review_status == ReviewStatus.rejected for a in deleted_atoms)
    assert all("tracked_change_deleted_text" in a.review_flags for a in deleted_atoms)
    assert all(a.source_refs[0].locator.get("tracked_change") == "deleted" for a in deleted_atoms)
    assert all(a.authority_class == AuthorityClass.deleted_text for a in deleted_atoms)


def _emit(text: str, *, heading: bool = False):
    return DocxParser()._emit_atoms_for_text(
        project_id="p",
        artifact_id="a",
        filename="sow.docx",
        text=text,
        paragraph_index=0,
        table_index=None,
        row=None,
        cell=None,
        tracked_change=None,
        heading=heading,
    )


def test_overview_prose_without_scope_verb_is_captured_not_dropped() -> None:
    # The SOW overview sentence carries the deal's headline quantity but uses
    # no scope/install/exclude verb — the lexical classifier matches nothing.
    # It must be captured (fail open), not silently dropped.
    text = (
        "The customer requires onsite field services support to replace "
        "approximately 110 existing TVs and mounts across 23 dwellings."
    )
    atoms = _emit(text)
    assert len(atoms) == 1
    a = atoms[0]
    assert a.atom_type == AtomType.scope_item
    assert "prose_fallback_capture" in a.review_flags
    assert a.value.get("prose_fallback") is True
    assert "110" in a.raw_text


def test_section_heading_captured_as_attribution_candidate() -> None:
    # Section headings are the PARENT label for everything beneath them, so they
    # are captured as atoms (flagged section_heading) to serve as attribution
    # candidates for site/section resolution — NOT dropped. Non-heading short
    # fragments still drop so the graph is not flooded.
    atoms = _emit("Project Overview", heading=True)
    assert len(atoms) == 1
    assert "section_heading" in atoms[0].review_flags
    assert atoms[0].atom_type == AtomType.scope_item
    # A non-heading bare label still fails the prose/list gate and is dropped.
    assert _emit("Scope") == []


def test_matched_scope_prose_keeps_full_confidence() -> None:
    # A paragraph that DOES match a lexical pattern is unaffected by the
    # fallback path (full confidence, no fallback flag).
    atoms = _emit("Installation of IP cameras at the main campus is in scope.")
    assert atoms
    assert all("prose_fallback_capture" not in a.review_flags for a in atoms)


def test_short_numeric_fact_line_is_kept(tmp_path: Path) -> None:
    # A short "label: value" line such as "Estimated quantity: 110 units" is
    # under the 5-word sentence threshold but states a concrete deal fact, so
    # the prose gate must keep it. A bare label with no digit stays dropped.
    assert _emit("Estimated quantity: 110 units") != []
    assert _emit("Project duration: 2 weeks") != []
    # No digit / no context -> dropped.
    assert _emit("Project Overview") == []
    assert _emit("110") == []


def test_full_parse_does_not_drop_body_paragraph_near_table(tmp_path: Path) -> None:
    # Regression for the id()-collision data-loss bug: python-docx creates
    # throwaway Paragraph proxies that are GC'd immediately, so an id()-based
    # "is this a table cell?" test produced false positives and silently
    # dropped real body paragraphs. Build a doc with MANY body paragraphs plus
    # a table (to churn proxy allocations) and assert a sentinel overview
    # paragraph survives the full parse, not just an isolated _emit.
    doc = Document()
    doc.add_heading("Statement of Work", level=1)
    sentinel = (
        "The customer requires onsite field services support to replace "
        "approximately 110 existing TVs and mounts across 23 dwellings."
    )
    doc.add_paragraph(sentinel)
    # Filler body paragraphs to force many proxy allocations / GC churn.
    for i in range(80):
        doc.add_paragraph(
            f"Filler narrative paragraph number {i} describing routine "
            "logistics and coordination between the parties."
        )
    # A real table whose cells must NOT bleed into the body loop.
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Site"
    table.rows[0].cells[1].text = "Quantity"
    table.rows[1].cells[0].text = "Main Campus"
    table.rows[1].cells[1].text = "50"
    path = tmp_path / "overview_with_table.docx"
    doc.save(path)

    atoms = DocxParser().parse_artifact(
        project_id="proj_x",
        artifact_id="art_x",
        path=path,
    )

    # The headline "110" overview must appear in some atom's raw text.
    assert any("110 existing" in (a.raw_text or "").lower() for a in atoms), (
        "overview body paragraph carrying the headline quantity was dropped"
    )
    # The table data must also be present (table loop still works).
    assert any("50" in (a.raw_text or "") for a in atoms)
