"""Tests for the span-keyed provenance ledger (lost-content attribution)."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from app.core.span_ledger import DropRecord, SpanLedger, StageKind
from app.parsers.docx_parser import DocxParser


def test_coverage_and_canary() -> None:
    led = SpanLedger(coverage_threshold=0.95)
    for i in range(10):
        led.register_span(f"s{i}", f"text {i}")
    for i in range(9):
        led.mark_represented(f"s{i}")
    rep, total, ratio = led.coverage()
    assert (rep, total) == (9, 10)
    assert abs(ratio - 0.9) < 1e-9
    assert led.canary_ok() is False  # 90% < 95% threshold


def test_gate_vs_seam_attribution() -> None:
    led = SpanLedger()
    led.record_drop(
        span_id="a:p1", stage="docx_parse.prose_gate", kind=StageKind.GATE,
        rule="_is_substantive_prose", reason="verb-less bullet", raw_text="Network design, firewall config",
    )
    led.record_drop(
        span_id="a:atm9", stage="semantic_dedup", kind=StageKind.SEAM,
        rule="decide", reason="duplicate", raw_text="some dup",
    )
    gate = led.gate_losses()
    seam = led.seam_losses()
    assert len(gate) == 1 and gate[0].kind is StageKind.GATE
    assert len(seam) == 1 and seam[0].kind is StageKind.SEAM
    # The kind drives the fix shape — that IS the parser-vs-decide answer.
    assert gate[0].fix_shape.startswith("code")
    assert "correction" in seam[0].fix_shape


def test_represented_span_is_not_a_loss() -> None:
    # A span dropped at one stage but represented elsewhere is not "lost".
    led = SpanLedger()
    led.register_span("a:p1", "kept text")
    led.mark_represented("a:p1")
    led.record_drop(
        span_id="a:p1", stage="x", kind=StageKind.GATE, rule="r", reason="why", raw_text="kept text",
    )
    assert led.lost_records() == []


def test_ingest_suppressed_atoms_as_seam() -> None:
    led = SpanLedger()
    n = led.ingest_suppressed_atoms([
        {
            "id": "atm1", "artifact_id": "art", "raw_text": "dup row",
            "decision_provenance": {"stage": "semantic_dedup", "reason": "duplicate"},
        },
        {"raw_text": "", "id": "atm2"},  # empty -> skipped
    ])
    assert n == 1
    seam = led.seam_losses()
    assert len(seam) == 1
    assert seam[0].kind is StageKind.SEAM
    assert seam[0].stage == "semantic_dedup"


def test_docx_parser_records_gate_drops(tmp_path: Path) -> None:
    # End-to-end: a short label fragment (no digit, < 5 words) is dropped by
    # the prose GATE and the ledger attributes it to the exact rule — proving
    # silent parser loss becomes a visible, classified signal. NOTE: multi-word
    # (5+) verb-less bullets now FAIL OPEN (kept as prose_fallback) and flow to
    # the learnable decide() SEAM, so they are no longer GATE losses; only true
    # short label fragments die at the GATE.
    doc = Document()
    doc.add_paragraph("Network Design")  # 2-word label, no digit -> GATE drop
    doc.add_paragraph(
        "Network design, firewall configuration, or switch configuration"
    )  # 5+ words -> now kept (fails open)
    doc.add_paragraph("Installation of IP cameras at the main campus is in scope.")  # kept
    path = tmp_path / "scope.docx"
    doc.save(path)

    led = SpanLedger()
    parser = DocxParser()
    parser._ledger = led
    parser.parse_artifact(project_id="p", artifact_id="art", path=path)

    losses = led.gate_losses()
    # The short label fragment is a GATE loss, attributed to the exact rule.
    assert any(d.raw_text.strip().lower() == "network design" for d in losses)
    assert all(d.rule == "_is_substantive_prose" for d in losses)
    # The 5+ word verb-less bullet now fails open — NOT a loss.
    assert not any("firewall configuration" in d.raw_text.lower() for d in led.lost_records())
    # The kept scope line must NOT appear as a loss.
    assert not any("ip cameras" in d.raw_text.lower() for d in led.lost_records())


def test_report_renders_both_buckets() -> None:
    led = SpanLedger()
    led.register_span("a:p1", "x")
    led.record_drop(span_id="a:p1", stage="prose_gate", kind=StageKind.GATE,
                    rule="_is_substantive_prose", reason="verb-less bullet", raw_text="x")
    led.ingest_suppressed_atoms([
        {"id": "atm1", "artifact_id": "a", "raw_text": "dup",
         "decision_provenance": {"stage": "semantic_dedup", "reason": "duplicate"}},
    ])
    text = led.report()
    assert "PARSER ISSUES" in text
    assert "DECIDE ISSUES" in text
    assert "coverage canary" in text
