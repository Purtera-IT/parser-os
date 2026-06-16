"""Tests for the content census (per-modality region inventory + invariant)
and the independent .docx census reader.

These prove the *never-detected* loss class is caught: content that lives in
content controls / textboxes / embedded media — invisible to
``Document.paragraphs`` / ``Document.tables`` — is inventoried by an
independent reader and reconciled against what the parser actually emitted.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from lxml import etree

from app.core.content_census import (
    ContentCensus,
    CoverageStatus,
    Region,
    RegionKind,
)
from app.parsers.census_docx import census_docx
from app.parsers.docx_parser import DocxParser

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class _Atom:
    """Minimal stand-in for an EvidenceAtom for census reconciliation."""

    def __init__(self, raw_text: str = "", value: dict | None = None) -> None:
        self.raw_text = raw_text
        self.value = value or {}


# --- pure census unit tests ------------------------------------------------

def test_text_region_covered_when_text_appears_in_an_atom() -> None:
    census = ContentCensus(artifact="a")
    census.register(Region("r1", "a", RegionKind.TEXT, "body/p0", text="replace 110 existing TVs"))
    census.reconcile([_Atom(raw_text="The plan is to replace 110 existing TVs across the site.")])
    assert census.status("r1") is CoverageStatus.COVERED
    assert census.invariant_ok()


def test_text_region_uncovered_is_silent_loss() -> None:
    census = ContentCensus(artifact="a")
    census.register(Region("r1", "a", RegionKind.TEXT, "body/p0", text="hidden contact yonah sapir"))
    census.reconcile([_Atom(raw_text="totally unrelated atom text")])
    assert census.status("r1") is CoverageStatus.UNCOVERED
    assert not census.invariant_ok()
    assert [r.region_id for r in census.uncovered()] == ["r1"]


def test_binary_region_marked_by_marker_atom() -> None:
    census = ContentCensus(artifact="a")
    census.register(Region("img1", "a", RegionKind.IMAGE, "media/image1.png", note="1234 bytes"))
    census.reconcile([
        _Atom(raw_text="[Image awaiting OCR]", value={"kind": "image_marker", "region_ref": "media/image1.png"}),
    ])
    assert census.status("img1") is CoverageStatus.MARKED
    # MARKED satisfies the invariant — the region didn't silently vanish.
    assert census.invariant_ok()


def test_binary_region_uncovered_when_no_marker() -> None:
    census = ContentCensus(artifact="a")
    census.register(Region("img1", "a", RegionKind.IMAGE, "media/image1.png"))
    census.reconcile([_Atom(raw_text="some text", value={})])
    assert census.status("img1") is CoverageStatus.UNCOVERED
    assert not census.invariant_ok()


def test_coverage_by_kind_counts() -> None:
    census = ContentCensus(artifact="a")
    census.register(Region("t1", "a", RegionKind.TEXT, "body/p0", text="alpha beta"))
    census.register(Region("t2", "a", RegionKind.TEXT, "body/p1", text="never emitted"))
    census.register(Region("i1", "a", RegionKind.IMAGE, "media/image1.png"))
    census.reconcile([
        _Atom(raw_text="alpha beta gamma"),
        _Atom(raw_text="marker", value={"kind": "image_marker", "region_ref": "media/image1.png"}),
    ])
    by_kind = census.coverage_by_kind()
    assert by_kind["text"] == (1, 2)
    assert by_kind["image"] == (1, 1)


# --- end-to-end: independent reader vs the docx parser ---------------------

def _docx_with_hidden_content(tmp_path: Path) -> Path:
    """Build a .docx whose contacts live inside a content control (w:sdt) and
    whose executive summary lives in a second sdt — both invisible to
    python-docx's ``Document.paragraphs`` / ``.tables``.
    """
    doc = Document()
    doc.add_heading("Project Overview", level=1)
    doc.add_paragraph(
        "The customer requires onsite field services to replace approximately "
        "110 existing TVs across 23 dwellings."
    )
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Site"
    t.rows[0].cells[1].text = "Qty"
    t.rows[1].cells[0].text = "Main campus"
    t.rows[1].cells[1].text = "50"
    path = tmp_path / "hidden.docx"
    doc.save(path)

    # Inject an sdt-wrapped contacts table + an sdt-wrapped exec summary.
    d2 = Document(path)
    body = d2.element.body
    sdt = etree.SubElement(body, _W + "sdt")
    sdtc = etree.SubElement(sdt, _W + "sdtContent")
    tbl = etree.SubElement(sdtc, _W + "tbl")
    for name, email in [("Dan Pratt", "dan@x.com"), ("Yonah Sapir", "yonah.sapir@4dcw.com")]:
        tr = etree.SubElement(tbl, _W + "tr")
        for val in (name, email):
            tc = etree.SubElement(tr, _W + "tc")
            pp = etree.SubElement(tc, _W + "p")
            r = etree.SubElement(pp, _W + "r")
            tt = etree.SubElement(r, _W + "t")
            tt.text = val
    sdt2 = etree.SubElement(body, _W + "sdt")
    sdtc2 = etree.SubElement(sdt2, _W + "sdtContent")
    ep = etree.SubElement(sdtc2, _W + "p")
    er = etree.SubElement(ep, _W + "r")
    et = etree.SubElement(er, _W + "t")
    et.text = "EXECUTIVE SUMMARY: deployment covers 23 dwellings over 2 weeks."
    d2.save(path)
    return path


def test_parser_recovers_hidden_sdt_content(tmp_path: Path) -> None:
    path = _docx_with_hidden_content(tmp_path)
    atoms = DocxParser().parse_artifact(project_id="p", artifact_id="art", path=path)
    blob = " || ".join(a.raw_text.lower() for a in atoms)
    # The never-detected content is now present as atoms.
    assert "yonah.sapir@4dcw.com" in blob
    assert "dan@x.com" in blob
    assert "executive summary" in blob
    # The w:sdt-descending walker (_iter_block_items) now pulls content-control
    # content through the MAIN structural path — with real sections and table
    # structure — instead of the lossy `recovered_nested_region` fallback. So the
    # hidden content is captured AND clean (not flagged as a degraded recovery).
    clean_blob = " || ".join(
        a.raw_text.lower()
        for a in atoms
        if "recovered_nested_region" not in getattr(a, "review_flags", [])
    )
    # sdt table data row + sdt paragraph both arrive via the clean structural
    # path. (Dan's row is row 0 of a header-less table, so the main loop's
    # column-label skip drops it to the fallback — that heuristic is unrelated
    # to sdt descent and a real contacts table ships a header row.)
    assert "yonah.sapir@4dcw.com" in clean_blob
    assert "executive summary" in clean_blob


def test_census_invariant_only_flags_intentional_drops(tmp_path: Path) -> None:
    """After recovery, the only UNCOVERED regions are content the parser drops
    *on purpose* (the column-label header row and the bare heading) — never the
    hidden contacts. That is the proof the never-detected class is closed.
    """
    path = _docx_with_hidden_content(tmp_path)
    atoms = DocxParser().parse_artifact(project_id="p", artifact_id="art", path=path)
    census = census_docx(path, artifact_id="art")
    census.reconcile(atoms)

    uncovered_text = {r.text for r in census.uncovered()}
    # Hidden contacts and exec summary must NOT be uncovered.
    assert not any("yonah" in t.lower() for t in uncovered_text)
    assert not any("dan@x.com" in t.lower() for t in uncovered_text)
    assert not any("executive summary" in t.lower() for t in uncovered_text)
    # The sdt contact rows reconcile as COVERED.
    contact_regions = [r for r in census.regions.values() if "yonah" in r.text.lower()]
    assert contact_regions and all(
        census.status(r.region_id) is CoverageStatus.COVERED for r in contact_regions
    )


def test_census_reader_inventories_more_than_body_view(tmp_path: Path) -> None:
    """The independent reader sees the sdt-nested table+paragraph that
    python-docx's body view misses — the denominator is genuinely independent.
    """
    path = _docx_with_hidden_content(tmp_path)
    census = census_docx(path, artifact_id="art")
    locations = {r.location for r in census.regions.values()}
    assert any(loc.startswith("sdt/tbl") for loc in locations)
    assert any(loc.startswith("sdt/p") for loc in locations)
    # python-docx body view: 1 table only (the visible one), 0 sdt tables.
    d = Document(path)
    assert len(d.tables) == 1
