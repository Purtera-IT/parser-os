"""Inside the parsers, the same rule as at the router: structure over strings.

The routing audit established the evidence order -- container, then content,
then extension, then filename as a prior that may raise a claim but never
create one. These are the places INSIDE the two big content parsers where the
same rule was being broken:

  * ``DocxParser._heading_level`` matched the English substring "heading", so
    section detection depended on the LOCALE of the Word install that wrote
    the file, and on whether anyone had used a corporate template.
  * ``detect_hybrid_summary_transcript`` accepted a filename as sufficient
    evidence that a PDF was a transcript, which changed the resulting atom
    TYPE.

Both are pinned here because both are invisible failures: nothing raises, the
document simply comes out flat or comes out mistyped.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.parsers.docx_parser import DocxParser

_SECTIONS = [
    ("Scope of Work", [
        "Contractor shall install forty IP cameras at the Atlanta facility.",
        "All drops terminate in IDF 3 and are certified to TIA-568.",
    ]),
    ("Site Access", [
        "Escort is required at all times inside the yard.",
        "Work occurs after hours on weekends only.",
    ]),
    ("Exclusions", [
        "Mid-turn jumpers are excluded from this bill of materials.",
        "Conduit above ten feet is by others.",
    ]),
]


def _build_docx(target: Path, heading_style_name: str | None) -> Path:
    """``None`` uses built-in 'Heading 1' -- the shape the whole corpus has."""
    doc = Document()
    if heading_style_name is None:
        use = "Heading 1"
    else:
        style = doc.styles.add_style(heading_style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 1"]
        # Word writes outlineLvl into the style itself; set it explicitly so
        # this is the real shape and not a python-docx default.
        ppr = style.element.get_or_add_pPr()
        lvl = OxmlElement("w:outlineLvl")
        lvl.set(qn("w:val"), "0")
        ppr.append(lvl)
        use = heading_style_name
    for title, paragraphs in _SECTIONS:
        doc.add_paragraph(title, style=use)
        for text in paragraphs:
            doc.add_paragraph(text)
    doc.save(target)
    return target


def _section_paths(path: Path) -> set[str]:
    output = DocxParser().parse_artifact(project_id="p", artifact_id="a", path=path)
    atoms = output.atoms if hasattr(output, "atoms") else output
    found: set[str] = set()
    for atom in atoms:
        locator = (atom.source_refs[0].locator if atom.source_refs else {}) or {}
        section = locator.get("section_path") or locator.get("section")
        if section:
            found.add("/".join(section) if isinstance(section, list) else str(section))
    return found


@pytest.mark.parametrize(
    "style_name, label",
    [
        (None, "built-in 'Heading 1'"),
        ("Uberschrift 1", "German Word"),
        ("Titre 1", "French Word"),
        ("Encabezado 1", "Spanish Word"),
        ("PurTera Section Head", "corporate template based on Heading 1"),
    ],
)
def test_headings_are_found_by_structure_not_by_an_english_string(
    tmp_path: Path, style_name: str | None, label: str
) -> None:
    """The same document, three sections deep, under five heading styles.

    ``_heading_level`` answered from ``style.name.startswith("heading")``.
    Every style above carries ``w:outlineLvl`` -- Word writes it into heading
    styles and a style based on Heading N inherits it -- but only the first
    contains the substring. Measured before the fix: 3 sections for the
    built-in style, ZERO for the other four, with every atom losing its
    section_path.

    This file had already learned the lesson for lists:
    ``_paragraph_is_list_item`` reads ``w:numPr`` off the XML because "Word
    frequently leaves list paragraphs on the Normal style while carrying real
    numbering". Headings never got the same treatment.
    """
    target = _build_docx(tmp_path / "sow.docx", style_name)
    paths = _section_paths(target)
    for expected, _body in _SECTIONS:
        assert any(expected in p for p in paths), (
            f"{label}: section {expected!r} lost; found {sorted(paths)}"
        )


def test_outline_fallback_never_overrules_the_style_name(tmp_path: Path) -> None:
    """The fallback is additive, and that is what makes it safe to ship.

    It runs only after the style name has declined to answer, so it can add a
    heading the name could not see and can never change one the name already
    found. Across all 13 real .docx available the two signals already agree
    exactly, and parser output is byte-identical before and after.
    """
    built_in = _section_paths(_build_docx(tmp_path / "a.docx", None))
    localised = _section_paths(_build_docx(tmp_path / "b.docx", "Uberschrift 1"))
    assert built_in == localised


# ── the PDF half: a filename must not decide the atom TYPE ───────────────


def _pdf(target: Path, lines: list[str]) -> Path:
    fitz = pytest.importorskip("fitz")
    doc = fitz.open()
    page = doc.new_page()
    y = 72
    for line in lines:
        page.insert_text((72, y), line, fontsize=11)
        y += 18
    doc.save(str(target))
    doc.close()
    return target


_SOW_LINES = [
    "Statement of Work",
    "",
    "Contractor shall install forty IP cameras at the Atlanta facility.",
    "All drops terminate in IDF 3 and are certified to TIA-568.",
    "Escort is required at all times inside the yard.",
    "Mid-turn jumpers are excluded from this bill of materials.",
    "Conduit above ten feet is by others.",
]

_TRANSCRIPT_LINES = [
    "Kickoff Call",
    "",
    "Cliff Creech [00:12] we need forty sites by Q3.",
    "Dana Whitfield [00:31] escort is required at the Atlanta dock.",
    "Cliff Creech [01:04] mid-turn jumpers are excluded.",
    "Dana Whitfield [01:22] confirm the badge list by Friday.",
]


def _plan_kind(path: Path) -> str | None:
    fitz = pytest.importorskip("fitz")
    from app.core.hybrid_summary_transcript import detect_hybrid_summary_transcript

    with fitz.open(str(path)) as doc:
        pages = [(page.get_text() or "") for page in doc]
    plan = detect_hybrid_summary_transcript(filename=path.name, title=None, page_texts=pages)
    return plan.kind if plan else None


def test_a_scope_pdf_named_transcript_is_not_re_atomised_as_conversation(
    tmp_path: Path,
) -> None:
    """``title_or_filename_transcript`` alone satisfied the gate.

    So a PDF whose NAME contained "transcript", with no speaker stamps and no
    Full Transcript marker, fell through to kind="transcript_only" and had its
    prose rebuilt as conversation turns. On one document rendered twice from
    identical text, the filename changed the atom TYPE:

        scope_of_work.pdf              -> no plan          exclusion
        kickoff_meeting_transcript.pdf -> transcript_only  action_item

    exclusion is scope-governing; action_item is a follow-up note.
    """
    assert _plan_kind(_pdf(tmp_path / "scope_of_work.pdf", _SOW_LINES)) is None
    assert _plan_kind(_pdf(tmp_path / "kickoff_meeting_transcript.pdf", _SOW_LINES)) is None


def test_a_real_transcript_pdf_is_still_detected_without_a_helpful_name(
    tmp_path: Path,
) -> None:
    """The other half: removing the shortcut must not cost the capability.

    Speaker-timestamp density is content and qualifies on its own, so a real
    transcript is found under a neutral name.
    """
    assert _plan_kind(_pdf(tmp_path / "attachment_c.pdf", _TRANSCRIPT_LINES)) == "transcript_only"
    assert _plan_kind(_pdf(tmp_path / "kickoff_transcript.pdf", _TRANSCRIPT_LINES)) == "transcript_only"


# ── xlsx: a tab NAME must not delete the sheet's contents ────────────────


_SCOPE_ROWS = [
    ["Site", "Building", "Device", "Qty", "Notes"],
    ["ATL-01", "Building C", "IP Camera", 24, "escort required"],
    ["ATL-02", "Building C", "IP Camera", 18, "after hours"],
    ["ATL-03", "Building D", "Access Point", 12, ""],
]
_LOOKUP_ROWS = [
    ["PS-L1-ENG-LABOR", "L1", "Hourly"],
    ["PS-L2-ENG-LABOR", "L2", "Hourly"],
    ["PS-L3-ENG-LABOR", "L3", "Hourly"],
    ["PS-L4-ENG-LABOR", "L4", "Fixed"],
    ["PS-L0-TECH-LABOR", "L0", "Hourly"],
]


@pytest.mark.parametrize(
    "tab", ["Lookup", "Helper", "DO NOT EDIT", "Dropdown", "Validation"]
)
def test_a_real_scope_table_survives_a_backing_data_tab_name(tab: str) -> None:
    """``suppress=True`` drops every atom on the sheet, and rule 1 decided it
    on the tab name -- the comment said "regardless of content".

    A tab name is as forgeable as a filename and gets renamed far more often:
    a template's "Lookup" tab that now holds the real BOM lost the whole
    table silently. The guard is the one this file already applies to the
    FINANCIAL name rule, whose comment says "the data-header guard already
    protects any genuine scope table that happens to live under such a name".
    """
    from app.parsers.sheet_classifier import classify_sheet

    assert not classify_sheet(tab, _SCOPE_ROWS).suppress, (
        f"a Site/Device/Qty table was dropped because the tab is called {tab!r}"
    )


@pytest.mark.parametrize(
    "tab", ["Lookup", "Helper", "DO NOT EDIT", "SELL RATES", "PriceList"]
)
def test_genuine_backing_data_is_still_suppressed(tab: str) -> None:
    """The other half: the rule exists for a reason and must keep working.

    ``_looks_like_data_header`` wants two or more tokens from a scope/BOM
    vocabulary (site, device, qty, part number, ...). A rate card's
    ``Code | Skill | Billing`` carries none of them, so rate cards and price
    books keep collapsing exactly as before -- which is the correct behaviour
    for boilerplate that appears on most deals.
    """
    from app.parsers.sheet_classifier import classify_sheet

    assert classify_sheet(tab, _LOOKUP_ROWS).suppress


# ── pptx: a slide title comes from a title placeholder, not a substring ──


def test_a_subtitle_is_not_a_title(tmp_path: Path) -> None:
    """``"title" in str(placeholder_format.type).lower()`` matches SUBTITLE.

    The scan broke on the first near-miss, so a subtitle sitting earlier in
    shape order became the slide title -- and a slide title becomes the
    section heading for every atom on that slide. Found on a real deck: slide
    1 titled itself "April 12, 2025" while its CENTER_TITLE, later in the
    shape tree, held the actual title.
    """
    pptx = pytest.importorskip("pptx")
    from app.parsers.pptx_parser import PptxParser

    presentation = pptx.Presentation()
    slide = presentation.slides.add_slide(presentation.slide_layouts[0])
    slide.shapes.title.text = "Structured Cabling Rollout"
    slide.placeholders[1].text = "April 12, 2025"

    # Reproduce the real deck's ordering: subtitle first in the shape tree.
    tree = slide.shapes._spTree
    element = slide.shapes.title._element
    tree.remove(element)
    tree.append(element)

    target = tmp_path / "deck.pptx"
    presentation.save(target)

    output = PptxParser().parse_artifact(project_id="p", artifact_id="a", path=target)
    atoms = output.atoms if hasattr(output, "atoms") else output
    titles = set()
    for atom in atoms:
        locator = (atom.source_refs[0].locator if atom.source_refs else {}) or {}
        found = locator.get("slide_title") or locator.get("section")
        if found:
            titles.add(found)
    assert "April 12, 2025" not in titles, "a subtitle was used as the slide title"
    assert "Structured Cabling Rollout" in titles


# ── html: a CSS class name must not silence the document ────────────────


_SOW_HTML = """<html><body>
<h1>Statement of Work</h1>
<h2>Scope</h2>
<p>Contractor shall install forty IP cameras at the Atlanta facility.</p>
<p>All drops terminate in IDF 3 and are certified to TIA-568.</p>
<h2>Exclusions</h2>
<ul><li>Mid-turn jumpers are excluded from this bill of materials.</li>
<li>Conduit above ten feet is by others.</li></ul>
<table><tr><th>Site</th><th>Qty</th></tr><tr><td>ATL-01</td><td>24</td></tr></table>
{extra}
</body></html>"""

# A stock Bulma notice. "message-body" is a standard component class in Bulma
# and several CMS themes -- nothing to do with chat.
_BULMA_NOTICE = (
    '<article class="message is-warning">'
    '<div class="message-body">Escort is required at all times.</div></article>'
)


def _html_atoms(target: Path, extra: str) -> int:
    from app.parsers.universal_parsers import HtmlParser

    target.write_text(_SOW_HTML.format(extra=extra), encoding="utf-8")
    output = HtmlParser().parse_artifact_full(
        project_id="p", artifact_id="a", path=target
    )
    return len(output.atoms)


def test_one_css_class_cannot_delete_the_whole_page(tmp_path: Path) -> None:
    """The chat detectors matched CSS class names and then RETURNED.

    So when a detector was wrong, the generic heading / paragraph / table walk
    never ran and the page produced nothing. Measured: an ordinary SOW page
    yields 11 atoms; adding a single ``<div class="message-body">`` -- one
    Bulma notice -- produced ZERO. Silent, total loss of the document.
    """
    plain = _html_atoms(tmp_path / "plain.html", "")
    with_notice = _html_atoms(tmp_path / "notice.html", _BULMA_NOTICE)
    assert plain == 11
    assert with_notice == plain, (
        f"a Bulma notice cost {plain - with_notice} atoms"
    )


def test_a_real_chat_export_still_takes_the_chat_path(tmp_path: Path) -> None:
    """The chat path has to keep winning when it actually has messages.

    It now has to PRODUCE something to win: the detector firing is no longer
    enough, because the detector reads class names and the extractor reads
    ``data-tid`` -- they could disagree, and when they did the document was
    lost rather than merely mis-scanned.
    """
    from app.parsers.universal_parsers import HtmlParser

    messages = "".join(
        '<div><div data-tid="messageHeader">Cliff Creech</div>'
        f'<div data-tid="message-body">we need forty sites by Q3, point {i}.</div></div>'
        for i in range(5)
    )
    target = tmp_path / "teams_export.html"
    target.write_text(f"<html><body><h1>Chat</h1>{messages}</body></html>", encoding="utf-8")
    output = HtmlParser().parse_artifact_full(project_id="p", artifact_id="a", path=target)
    kinds = {(atom.value or {}).get("kind") for atom in output.atoms}
    assert "teams_message" in kinds
    assert any("Cliff Creech" in atom.raw_text for atom in output.atoms)


# ── json: stepping aside is not the same as leaving nobody holding it ────


def test_a_json_deferral_does_not_drop_the_file(tmp_path: Path) -> None:
    """JsonParser returned 0.00 to hand the file to TranscriptParser.

    TranscriptParser's .json branch needs ``json.loads`` to SUCCEED, and the
    router hands it a 4000-character sample, so a large .json raises there and
    scores 0.00 too. Both parsers stood down and the artifact routed to NONE.

    Measured on the local corpus: 2062 of 2500 artifacts landed on NONE -- all
    valid .json that parse fine in full. They carry the key ``"segments"``,
    which in those files means document segments, and which is an ordinary
    business word in any case (network, customer, cable).
    """
    from app.parsers.registry import choose_parser

    payload = '{\n  "segments": [\n' + ",\n".join(
        f'    {{"id": {i}, "label": "cable segment {i}", "sites": ["ATL-{i:02d}"]}}'
        for i in range(400)
    ) + "\n  ]\n}\n"
    target = tmp_path / "artifact_cache.json"
    target.write_text(payload, encoding="utf-8")
    assert len(payload) > 4000, "the point of the case is that the sample truncates"

    parser, match, _all = choose_parser(target)
    assert parser is not None, "valid JSON routed to NONE"
    assert match.confidence >= 0.5


def test_a_real_transcript_json_still_reaches_the_transcript_parser(
    tmp_path: Path,
) -> None:
    """The deferral must still defer. 0.55 sits below TranscriptParser's 0.8."""
    from app.parsers.registry import choose_parser

    payload = (
        '{"utterances": ['
        + ",".join(
            f'{{"speaker": "Cliff", "text": "point {i}", "start": {i}}}'
            for i in range(5)
        )
        + "]}"
    )
    target = tmp_path / "meeting.json"
    target.write_text(payload, encoding="utf-8")
    parser, _match, _all = choose_parser(target)
    assert type(parser).__name__ == "TranscriptParser"


# ── email: the container must not change the evidence ───────────────────


_MAIL_HEADERS = [
    ("From", "jane.customer@acme.example"),
    ("To", "pm@purtera.example"),
    ("Subject", "Scope update"),
]
_MAIL_BODY = (
    "Please remove the West Wing from scope.\n"
    "Escort access is required at the Atlanta dock before 2pm.\n"
    "Mid-turn jumpers are excluded from this order.\n"
)


def _email_fingerprint(path: Path) -> tuple[int, tuple[tuple[str, str], ...]]:
    from app.parsers.email_parser import EmailParser

    output = EmailParser().parse_artifact_full(
        project_id="p", artifact_id="a", path=path
    )
    atoms = output.atoms if hasattr(output, "atoms") else output
    rows = tuple(
        (str(atom.atom_type).split(".")[-1], atom.raw_text.strip()[:60])
        for atom in atoms
    )
    return len(atoms), rows


def test_an_email_saved_as_text_yields_the_same_evidence_as_the_eml(
    tmp_path: Path,
) -> None:
    """One message, two containers, and they used to disagree.

    An Outlook "save as text" .txt carries the same From/Sent/To/Subject block
    a .eml does -- as body text rather than MIME headers -- so it reached the
    atom extractor and each header line became its own ``scope_item`` at
    ``customer_current_authored``:

        .eml   6 atoms   headers -> one deal_metadata atom
        .txt   9 atoms   headers -> four scope_item atoms at rank 90

    which is precisely the defect ``_header_atom``'s own comment describes
    having fixed -- "A From/To/Subject line is not a unit of work" -- surviving
    in the other container. Four phantom units of work in the customer-authored
    tier, per email.
    """
    eml = tmp_path / "message.eml"
    eml.write_bytes(
        ("".join(f"{k}: {v}\r\n" for k, v in _MAIL_HEADERS)
         + "\r\n" + _MAIL_BODY.replace("\n", "\r\n")).encode("utf-8")
    )
    txt = tmp_path / "message.txt"
    txt.write_text(
        "".join(f"{k}: {v}\n" for k, v in _MAIL_HEADERS) + "\n" + _MAIL_BODY,
        encoding="utf-8",
    )

    eml_count, eml_rows = _email_fingerprint(eml)
    txt_count, txt_rows = _email_fingerprint(txt)
    assert eml_rows == txt_rows, (
        f"container changed the evidence: .eml={eml_count} .txt={txt_count}"
    )
    assert not any(
        kind == "scope_item" and text.lower().startswith(("from:", "to:", "subject:", "sent:"))
        for kind, text in txt_rows
    ), "a header line was typed as a unit of work"
    assert any(kind == "deal_metadata" for kind, _ in txt_rows)


def test_a_body_that_merely_starts_with_a_colon_line_is_not_eaten(
    tmp_path: Path,
) -> None:
    """The splitter has to be conservative: the downside is eating content.

    It requires at least two header lines AND one of them to be ``from`` or
    ``subject``, so a note opening "Note: ..." over "Owner: ..." keeps every
    line.
    """
    note = tmp_path / "note.txt"
    note.write_text(
        "Note: the customer confirmed the West Wing is out.\n"
        "Owner: Cliff Creech\n"
        "Escort access is required at the Atlanta dock before 2pm.\n",
        encoding="utf-8",
    )
    _count, rows = _email_fingerprint(note)
    body = " ".join(text for _kind, text in rows)
    assert "West Wing" in body
    assert "Cliff Creech" in body


@pytest.mark.parametrize(
    "container",
    ["vtt", "srt", "txt_colon", "txt_ownline"],
)
def test_one_conversation_four_containers_same_speakers(
    tmp_path: Path, container: str
) -> None:
    """The transcript half of the same question -- and it already passes.

    Recorded rather than left implicit: .vtt voice spans, .srt cues,
    "Speaker:" lines and the Otter own-line dialect all resolve to the same
    two speakers, so the container does not decide who said what.
    """
    from app.parsers.registry import choose_parser

    turns = [
        ("Cliff Creech", "we need forty sites by Q3."),
        ("Dana Whitfield", "escort is required at the Atlanta dock."),
    ] * 8
    if container == "vtt":
        lines = ["WEBVTT", ""]
        for i, (who, what) in enumerate(turns):
            lines += [str(i + 1),
                      f"00:00:{i * 3:02d}.000 --> 00:00:{i * 3 + 2:02d}.000",
                      f"<v {who}>{what}</v>", ""]
        target = tmp_path / "a.vtt"
    elif container == "srt":
        lines = []
        for i, (who, what) in enumerate(turns):
            lines += [str(i + 1),
                      f"00:00:{i * 3:02d},000 --> 00:00:{i * 3 + 2:02d},000",
                      f"{who}: {what}", ""]
        target = tmp_path / "a.srt"
    elif container == "txt_colon":
        lines = [f"{who}: {what}" for who, what in turns]
        target = tmp_path / "a.txt"
    else:
        lines = []
        for who, what in turns:
            lines += [who, what]
        target = tmp_path / "b.txt"
    target.write_text("\n".join(lines), encoding="utf-8")

    parser, _match, _all = choose_parser(target)
    assert type(parser).__name__ == "TranscriptParser"
    output = parser.parse_artifact(project_id="p", artifact_id="a", path=target)
    atoms = output.atoms if hasattr(output, "atoms") else output
    speakers = set()
    for atom in atoms:
        locator = (atom.source_refs[0].locator if atom.source_refs else {}) or {}
        if locator.get("speaker"):
            speakers.add(locator["speaker"])
    assert speakers == {"Cliff Creech", "Dana Whitfield"}


# ── every format must agree on what a unit of evidence is ───────────────
#
# The uniform test for the parsers not covered above: write the same six facts
# into each format, parse, and compare. Content survived everywhere (6/6), so
# the defects were not loss -- they were GRANULARITY and TYPING, which are
# silent in a different way. An atom is the unit of typing: fuse an exclusion
# into a paragraph of scope and it can never be typed as an exclusion.

_FACTS = [
    "Contractor shall install forty IP cameras at the Atlanta facility.",
    "All drops terminate in IDF 3 and are certified to TIA-568.",
    "Escort access is required at the Atlanta dock before 2pm.",
    "Mid-turn jumpers are excluded from this bill of materials.",
    "The customer will provide badge access on weekdays only.",
    "Conduit above ten feet is by others.",
]
_NEEDLES = ["forty IP cameras", "TIA-568", "Escort access", "Mid-turn jumpers",
            "badge access", "Conduit above ten feet"]

_ODF_SHELL = (
    '<?xml version="1.0" encoding="UTF-8"?>'
    '<office:document-content '
    'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
    'xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" '
    'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
    'office:version="1.2"><office:body>{body}</office:body>'
    "</office:document-content>"
)


def _write_format(root: Path, fmt: str) -> Path:
    import zipfile

    if fmt == "md":
        target = root / "sow.md"
        target.write_text("# Statement of Work\n\n" + "\n\n".join(_FACTS) + "\n",
                          encoding="utf-8")
    elif fmt == "rtf":
        target = root / "sow.rtf"
        body = "".join(r"\par " + f + "\n" for f in _FACTS)
        target.write_text(r"{\rtf1\ansi\deff0 {\fonttbl{\f0 Calibri;}}" + "\n" + body + "}",
                          encoding="ascii", errors="replace")
    elif fmt == "ics":
        target = root / "kickoff.ics"
        # RFC 5545 escapes a line break inside a text value as a literal
        # backslash-n, not an actual newline -- a real newline would end the
        # DESCRIPTION property and orphan every fact after the first.
        desc = "\\n".join(_FACTS)
        target.write_bytes((
            "BEGIN:VCALENDAR\r\nVERSION:2.0\r\nPRODID:-//t//EN\r\nBEGIN:VEVENT\r\n"
            "UID:1@t\r\nDTSTAMP:20260115T090000Z\r\nDTSTART:20260115T090000Z\r\n"
            f"SUMMARY:Cabling kickoff\r\nDESCRIPTION:{desc}\r\n"
            "END:VEVENT\r\nEND:VCALENDAR\r\n"
        ).encode("utf-8"))
    elif fmt == "mbox":
        target = root / "thread.mbox"
        target.write_text(
            "From jane@acme.example Wed Jan 15 09:00:00 2026\n"
            "From: jane@acme.example\nTo: pm@purtera.example\n"
            "Subject: Scope update\n\n" + "\n".join(_FACTS) + "\n\n",
            encoding="utf-8",
        )
    elif fmt in {"odt", "ods"}:
        target = root / f"sow.{fmt}"
        if fmt == "odt":
            inner = "<office:text>" + "".join(f"<text:p>{f}</text:p>" for f in _FACTS) + "</office:text>"
            mime = "application/vnd.oasis.opendocument.text"
        else:
            rows = "".join(
                '<table:table-row><table:table-cell office:value-type="string">'
                f"<text:p>{f}</text:p></table:table-cell></table:table-row>"
                for f in _FACTS
            )
            inner = ('<office:spreadsheet><table:table table:name="Scope">'
                     + rows + "</table:table></office:spreadsheet>")
            mime = "application/vnd.oasis.opendocument.spreadsheet"
        with zipfile.ZipFile(target, "w") as z:
            z.writestr("mimetype", mime)
            z.writestr("content.xml", _ODF_SHELL.format(body=inner))
    else:  # zip
        target = root / "pack.zip"
        with zipfile.ZipFile(target, "w") as z:
            z.writestr("scope_notes.txt", "Statement of Work\n\n" + "\n".join(_FACTS) + "\n")
    return target


@pytest.mark.parametrize("fmt", ["md", "rtf", "ics", "mbox", "odt", "ods", "zip"])
def test_every_format_keeps_every_fact(tmp_path: Path, fmt: str) -> None:
    """Content preservation. All seven already passed; kept so they stay passing."""
    from app.parsers.registry import choose_parser

    target = _write_format(tmp_path, fmt)
    parser, _match, _all = choose_parser(target)
    assert parser is not None, f"{fmt} routed to NONE"
    output = parser.parse_artifact(project_id="p", artifact_id="a", path=target)
    atoms = output.atoms if hasattr(output, "atoms") else output
    blob = " || ".join((a.raw_text or "") for a in atoms).lower()
    missing = [n for n in _NEEDLES if n.lower() not in blob]
    assert not missing, f"{fmt} lost {missing}"


@pytest.mark.parametrize("fmt", ["md", "rtf", "ics", "mbox", "odt", "ods", "zip"])
def test_every_format_types_the_exclusion_as_an_exclusion(
    tmp_path: Path, fmt: str
) -> None:
    """The part that was actually broken, in three of the seven.

      * .odt / .ods -- _universal_extras._make_atom defaulted atom_type to the
        CONCLUSION scope_item, so no caller in that module ever classified:
        six atoms, all scope_item, exclusion and constraint included.
      * .mbox -- split on blank lines only, so a message arrived as one
        336-character atom while EmailParser gave the same content five typed
        atoms. mbox IS email; they must not disagree.
      * .ics -- SUMMARY, times and the whole DESCRIPTION fused into one
        394-character meeting_commitment, so working detail in an invite could
        never be read as a constraint or an exclusion.
    """
    from app.parsers.registry import choose_parser

    target = _write_format(tmp_path, fmt)
    parser, _match, _all = choose_parser(target)
    output = parser.parse_artifact(project_id="p", artifact_id="a", path=target)
    atoms = output.atoms if hasattr(output, "atoms") else output
    exclusions = [
        a for a in atoms
        if str(a.atom_type).endswith("exclusion") and "mid-turn" in (a.raw_text or "").lower()
    ]
    assert exclusions, (
        f"{fmt}: 'Mid-turn jumpers are excluded' is not an exclusion atom; "
        f"types were {sorted({str(a.atom_type).split('.')[-1] for a in atoms})}"
    )


@pytest.mark.parametrize("fmt", ["md", "rtf", "ics", "mbox", "odt", "ods"])
def test_no_format_fuses_the_whole_document_into_one_atom(
    tmp_path: Path, fmt: str
) -> None:
    """Granularity guard: six distinct facts must not arrive as one blob."""
    from app.parsers.registry import choose_parser

    target = _write_format(tmp_path, fmt)
    parser, _match, _all = choose_parser(target)
    output = parser.parse_artifact(project_id="p", artifact_id="a", path=target)
    atoms = output.atoms if hasattr(output, "atoms") else output
    longest = max((len(a.raw_text or "") for a in atoms), default=0)
    assert longest < 200, f"{fmt}: longest atom is {longest} chars -- facts are fused"


# ── msg: the fourth mail container joins the invariance contract ────────


def test_a_msg_yields_the_same_evidence_shape_as_the_eml(monkeypatch, tmp_path: Path) -> None:
    """The fourth container found carrying the header defect.

    .eml fixed it first, then .txt, then .mbox -- and .msg was still letting
    the classifier type its "From: | Subject: |" line, minting one phantom
    customer-authored scope item per message.

    No OLE2 writer exists to synthesise a real .msg, so ``extract_msg.openMsg``
    is monkeypatched -- which is the honest scope anyway: this pins OUR
    transformation (header typing, sentence granularity), not the library's
    decoding, which is extract-msg's own test suite's job.
    """
    import sys
    import types

    class _FakeMsg:
        subject = "Scope update"
        sender = "jane.customer@acme.example"
        date = "Wed, 15 Jan 2026 09:00:00 -0500"
        body = (
            "Please remove the West Wing from scope. "
            "Escort access is required at the Atlanta dock before 2pm. "
            "Mid-turn jumpers are excluded from this order."
        )
        attachments: list = []

        def __enter__(self):
            return self

        def __exit__(self, *args):
            return False

    fake_module = types.SimpleNamespace(openMsg=lambda _p: _FakeMsg())
    monkeypatch.setitem(sys.modules, "extract_msg", fake_module)

    from app.parsers._universal_extras import MsgParser

    target = tmp_path / "message.msg"
    target.write_bytes(b"\xd0\xcf\x11\xe0stub")  # routed by extension; body is faked
    output = MsgParser().parse_artifact_full(
        project_id="p", artifact_id="a", path=target
    )
    atoms = output.atoms if hasattr(output, "atoms") else output
    rows = [(str(a.atom_type).split(".")[-1], a.raw_text.strip()) for a in atoms]

    header_rows = [r for r in rows if r[1].startswith("From:")]
    assert header_rows == [
        ("deal_metadata",
         "From: jane.customer@acme.example | Subject: Scope update | "
         "Date: Wed, 15 Jan 2026 09:00:00 -0500"),
    ], "the header must be ONE deal_metadata atom, never scope"

    body_rows = [r for r in rows if not r[1].startswith("From:")]
    assert len(body_rows) == 3, "three sentences -> three atoms, not one paragraph blob"
    kinds = {k for k, _ in body_rows}
    assert "exclusion" in kinds, "'Mid-turn jumpers are excluded' must be typeable as one"
