"""Dual-screen zip-from-blob labeler for a BATCH of deals.
LEFT: the deal's PDFs (pdf.js) with the selected atom's text highlighted on its
page. RIGHT: atoms (teacher-prefilled accept/fix) + cross-context edges, in the
①-⑤ stages. Per-atom "parser issue" button -> emails griffin@optbotai.com
(mailto). All progress autosaves to localStorage (refresh-safe). Bundles the
deal source files into the zip.

Run: python _build_batch_labeler.py <slug> [<slug>...]
"""
import sys, json, glob, zipfile, os, shutil, re, html as _h
import numpy as np

def _copy_xlsx_open_first_sheet(src, dst):
    """Copy an xlsx but (a) UNHIDE every sheet and (b) make Excel OPEN it on the
    FIRST sheet (the Deal Kit). Deal kits ship with the money sheets hidden — the
    anyWAIR CALC hides Deal Kit AND Gantt Financials — but the parser reads them
    and the reviewer must see them; leaving them hidden also made the old
    activeTab=0 pointer reference a hidden sheet, which Excel treats as a corrupt
    workbook. Surgical zip edit — only sheet visibility + the activeTab/firstSheet
    pointers + per-sheet tabSelected flags change; every byte of data, formula and
    formatting is preserved (unlike an openpyxl re-save)."""
    import zipfile, shutil as _sh
    try:
        zin = zipfile.ZipFile(src)
        with zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
            for it in zin.infolist():
                data = zin.read(it.filename)
                if it.filename == "xl/workbook.xml":
                    t = data.decode("utf-8", "ignore")
                    # unhide all sheets so nothing (Deal Kit, Gantt Financials, …)
                    # is concealed from the reviewer / downloaded copy
                    t = re.sub(r'\s+state="(?:hidden|veryHidden)"', "", t)
                    # open on the first sheet (now visible); start the tab strip at it
                    if "activeTab=" in t:
                        t = re.sub(r'activeTab="\d+"', 'activeTab="0"', t)
                    elif "<workbookView" in t:
                        t = t.replace("<workbookView", '<workbookView activeTab="0"', 1)
                    t = re.sub(r'\s+firstSheet="\d+"', "", t)
                    data = t.encode("utf-8")
                elif re.match(r"xl/worksheets/sheet\d+\.xml$", it.filename):
                    # drop the "this is the open tab" flag on every sheet; activeTab=0 then wins
                    data = data.replace(b' tabSelected="1"', b"")
                zout.writestr(it, data)
        zin.close()
    except Exception:
        _sh.copy(src, dst)


def render_xlsx_sheets(path):
    """Each sheet -> {n: name, h: <table> html}, styled to resemble the real
    workbook — reads cell FILL colors + bold + number alignment from the xlsx
    (so the colored header bands / sections show), not a flat grey grid."""
    if not path or not os.path.exists(path) or not path.lower().endswith(".xlsx"):
        return []
    _numre = re.compile(r"^[-$(]?\s?[\d,]+(\.\d+)?\s?%?\)?$")

    def _fill_hex(cell):
        try:
            f = cell.fill
            if f is None or getattr(f, "patternType", None) != "solid":
                return None
            rgb = getattr(getattr(f, "fgColor", None), "rgb", None)
            if isinstance(rgb, str) and len(rgb) == 8:           # AARRGGBB, direct RGB only
                hx = rgb[2:]
                if hx.upper() not in ("000000", "FFFFFF"):
                    return "#" + hx
        except Exception:
            pass
        return None

    try:
        import openpyxl
        from openpyxl.utils import column_index_from_string
        wb = openpyxl.load_workbook(path, data_only=True)        # styles available (not read_only)
        out = []
        for sn in wb.sheetnames[:25]:
            ws = wb[sn]; rows = []
            # Excel hides these columns/rows (helper/formula scaffolding) and so
            # does the parser, so the preview hides them too — otherwise the
            # LEFT pane shows a "Country Multiplier" the atoms (RIGHT) don't,
            # which reads as a false "MISSED".
            hidden_cols = {column_index_from_string(c) - 1
                           for c, d in ws.column_dimensions.items() if d.hidden}
            hidden_rows = {i - 1 for i, d in ws.row_dimensions.items() if d.hidden}
            for ri, row in enumerate(ws.iter_rows()):
                if ri >= 200:
                    break
                if ri in hidden_rows:
                    continue
                cells = [c for ci, c in enumerate(row[:20]) if ci not in hidden_cols]
                vals = ["" if c.value is None else str(c.value) for c in cells]
                if not any(v.strip() for v in vals):
                    continue
                tds = []
                for c, v in zip(cells, vals):
                    st = ""
                    hx = _fill_hex(c)
                    if hx:
                        st += f"background:{hx};color:#111;"      # dark text on colored fill (readable)
                    try:
                        if c.font and c.font.bold:
                            st += "font-weight:600;"
                    except Exception:
                        pass
                    if v and _numre.match(v.strip()):
                        st += "text-align:right;font-variant-numeric:tabular-nums;"
                    tds.append(f'<td style="{st}">' + _h.escape(v[:160]) + "</td>")
                rows.append("<tr>" + "".join(tds) + "</tr>")
            out.append({"n": sn, "h": ('<table class="sgrid">' + "".join(rows) + "</table>") if rows else '<div class="muted">(empty sheet)</div>'})
        return out
    except Exception:
        return []


def render_docx_html(path):
    if not path or not os.path.exists(path) or not path.lower().endswith(".docx"):
        return ""
    try:
        import docx
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        from app.parsers.docx_parser import _iter_block_items
        d = docx.Document(path); out = []
        # Walk in READING ORDER and descend into w:sdt content controls, so the
        # preview shows EVERY table in document order — including the page-1
        # content-control tables (SOW version, contacts) plain python-docx skips.
        for kind, el in _iter_block_items(d.element.body):
            if kind == "p":
                t = (Paragraph(el, d).text or "").strip()
                if t:
                    out.append('<p class="sp">' + _h.escape(t) + "</p>")
            elif kind == "tbl":
                tbl = Table(el, d)
                rows = ["<tr>" + "".join("<td>" + _h.escape((c.text or "")[:120]) + "</td>" for c in r.cells[:14]) + "</tr>" for r in tbl.rows[:200]]
                if rows:
                    out.append('<table class="sgrid">' + "".join(rows) + "</table>")
        return "".join(out)
    except Exception:
        return ""

# hide parser junk from the intern (xlsx internals + float noise) — the parser
# fix is the deeper task; this keeps the labeling set clean now.
_XMLJUNK = re.compile(r"awaiting OCR.*?(\.xml|\.vml|_rels|drawings/|media/image|sharedStrings|workbook)", re.I)
_FLOATNOISE = re.compile(r"\d+\.\d{6,}")
# CAD drawing-set / site-schematic signals — we're not built to parse these yet,
# so whole drawing-set docs are dropped from the labeling set.
_SCHEM = re.compile(r"PLOTTED|LAYOUT:|\.dwg|AS-BUILT|CONSOLIDATED SET|NTICONS|SYMBOLS\s*&\s*LEGEND|RISER DIAGRAM|\b[TEM]0\d\d\b", re.I)
_SCHEM_NAME = re.compile(r"CONSOLIDATED SET|\bIFC\b|SYMBOLS?\s*&?\s*LEGEND|RISER| - T | - TA |\bT0\d\d\b|\.dwg|DRAWING|SCHEMATIC", re.I)
# deal-health gate: a deal that's mostly drawings/images/fragments isn't labelable yet.
_DRAW = re.compile(r"DO NOT SCALE|\bDWG\b|SHEET\s*\d+\s*OF|\bSCALE\s*[:=]|\bSYM\b|ELEVATION|ISOMETRIC|\bDETAIL\b|LEVELER|MAST|BASEPLATE|ANCHOR BOLT|PLOTTED|LAYOUT:|CONSOLIDATED SET|\bIFC\b|RISER|\.dwg", re.I)
def deal_unlabelable(corpus):
    n = len(corpus)
    if not n: return "empty"
    d = 100 * sum(1 for a in corpus if _DRAW.search(a["body"])) // n
    i = 100 * sum(1 for a in corpus if "awaiting OCR" in a["body"] or "image" in a["body"].lower()) // n
    f = 100 * sum(1 for a in corpus if len(a["body"].strip()) < 8 or a["body"].strip().startswith("col_")) // n
    if d >= 25: return f"drawing-heavy {d}%"
    if i >= 25: return f"image-heavy {i}%"
    if f >= 25: return f"fragment-heavy {f}%"
    return ""
def is_junk(b): return bool(_XMLJUNK.search(b or ""))
def clean_body(b): return _FLOATNOISE.sub(lambda m: ("%g" % round(float(m.group()), 2)), b or "")
from pathlib import Path
from collections import defaultdict
from _context_v2 import build_index, context
from _labeler_core import COARSE, FACETS, ADMISSION

FK = ("scope_polarity", "responsibility", "amount_kind", "obligation_modality", "metadata_kind")
slugs = [a for a in sys.argv[1:]] or ["optbot"]
OUT = Path("_batch_out")
if OUT.exists(): shutil.rmtree(OUT)
(OUT / "sources").mkdir(parents=True)


def deal_files(slug):
    for base in (f"_blob_pool/{slug}", f"_blob_deals/{slug}", "_optbot_inputs" if slug == "optbot" else None):
        if base and os.path.isdir(base):
            return sorted(glob.glob(base + "/*.pdf") + glob.glob(base + "/*.docx") + glob.glob(base + "/*.xlsx"))
    return []


def viewer_src(src):
    """prefer a derived .pdf sibling so docx/xlsx docs still get the PDF view."""
    if str(src).lower().endswith(".pdf"):
        return src if os.path.exists(src) else None
    stem = os.path.splitext(src)[0]
    if os.path.exists(stem + ".pdf"):
        return stem + ".pdf"
    return src if os.path.exists(src) else None


deals = {}
for slug in slugs:
    cf = f"_pool_cache/{slug}.json"
    pj = json.load(open(cf, encoding="utf-8")) if os.path.exists(cf) else {}
    corpus = pj.get("corpus", [])
    pathmap = pj.get("pathmap", {}) or {}
    srcdump = pj.get("srcdump", {}) or {}
    if not corpus:
        fsrc = deal_files(slug)
        if fsrc:
            from _labeler_core import parse_deal
            corpus, srcdump, pathmap = parse_deal(fsrc)
    if not corpus:
        print(f"skip {slug}: no corpus"); continue
    bad = deal_unlabelable(corpus)
    if bad:
        print(f"skip {slug}: NOT LABELABLE ({bad}) — drawings/images/junk, excluded"); continue
    # bundle ALL real (non-schematic) PDFs so the intern can browse every source doc
    sd = OUT / "sources" / slug; sd.mkdir(parents=True, exist_ok=True)
    pdf_files = []   # (stem, relpath) for the left-pane viewer
    srcdirs = [f"_blob_pool/{slug}", f"_blob_deals/{slug}"] + (["_optbot_inputs"] if slug == "optbot" else [])
    for sdir in srcdirs:
        for p in sorted(glob.glob(sdir + "/*.pdf")):
            if _SCHEM_NAME.search(Path(p).name):          # drop CAD/drawing-set PDFs
                continue
            if any(s == Path(p).stem for s, _ in pdf_files):
                continue
            try: shutil.copy(p, sd / Path(p).name); pdf_files.append((Path(p).stem, f"sources/{slug}/{Path(p).name}"))
            except Exception: pass
    def pdf_for(doc):
        for stem, rel in pdf_files:
            if doc == stem or doc in stem or stem in doc: return rel
        return ""
    # bundle the ACTUAL source files (docx/xlsx/pdf) so the intern can download + open the real thing
    dl_map = {}
    for d2, src in pathmap.items():
        if os.path.isfile(src) and src.lower().endswith((".pdf", ".docx", ".xlsx", ".xls", ".doc")):
            try:
                dst = sd / Path(src).name
                if src.lower().endswith(".xlsx"):
                    _copy_xlsx_open_first_sheet(src, dst)   # open on sheet 0 (Deal Kit), not the author's last tab
                else:
                    shutil.copy(src, dst)
                dl_map[d2] = f"sources/{slug}/{Path(src).name}"
            except Exception:
                try: shutil.copy(src, sd / Path(src).name); dl_map[d2] = f"sources/{slug}/{Path(src).name}"
                except Exception: pass
    # PARSER-EVAL: ignore the teacher type-prefill cache. It's keyed by atom
    # INDEX and was built against an earlier parse — after re-parsing, those
    # indices point at different atoms, so it (a) applies wrong types and
    # (b) silently SKIPS atoms whose new index isn't in its capped set. For
    # parser evaluation we want EVERY atom shown with clean heuristic types; the
    # teacher prefill is re-generated against the live parse for the later
    # type-labeling phase. Set SOWSMITH_USE_TEACHER=1 to re-enable.
    tA = (json.load(open(f"_teacher_cache/{slug}.json")).get("labels", {})
          if os.environ.get("SOWSMITH_USE_TEACHER") == "1" and os.path.exists(f"_teacher_cache/{slug}.json")
          else {})
    tE = json.load(open(f"_teacher_cache/{slug}_edges.json")).get("edges", []) if os.path.exists(f"_teacher_cache/{slug}_edges.json") else []
    # FAST mode: skip the embedding-backed cross-doc CONTEXT layer (build_index +
    # per-atom context()). That layer embeds every atom through qwen3/bge to build
    # a kNN similarity index — the slow part of a rebuild, and irrelevant when the
    # PM is reviewing per-sheet ORGANIZATION rather than cross-doc links. Atoms,
    # types, sections, breadcrumbs are all unchanged; only the "related context"
    # panel is empty. Set SOWSMITH_LABELER_FAST=1 to enable.
    _FAST = os.environ.get("SOWSMITH_LABELER_FAST") == "1"
    if _FAST:
        ents = types = simfn = kind = None
    else:
        ents, types, simfn, kind = build_index(corpus)

    # SEMANTIC near-duplicate collapse (embedding, not regex): a PDF often emits
    # the same fact twice — once as prose from the text layer, once as a detected
    # table_row ("Prior experience/references 20%" vs "Prior experience/references
    # | 20%"). Exact/structural dedup misses these (surface differs); embedding
    # cosine catches them. Conservative: drop a TABLE_ROW atom only when a
    # non-table atom in the SAME section is a near-identical match (cosine >=
    # 0.95), and keep the prose. Reuses the build_index matrix — zero extra cost.
    _semdrop: set[int] = set()
    if not _FAST and kind in ("embed", "bge") and simfn is not None:
        from collections import defaultdict as _dd
        _bysec = _dd(list)
        for _i, _a in enumerate(corpus):
            _bysec[_a.get("section", "")].append(_i)
        for _ids in _bysec.values():
            _prose = [j for j in _ids if (corpus[j].get("vkind") or "") != "table_row"]
            for _i in _ids:
                if (corpus[_i].get("vkind") or "") != "table_row":
                    continue
                if any(j != _i and simfn(_i, j) >= 0.95 for j in _prose):
                    _semdrop.add(_i)
        if _semdrop:
            print(f"  ~ collapsed {len(_semdrop)} table/prose semantic duplicate(s)")

    by_doc = defaultdict(lambda: defaultdict(list)); order_doc, order_tbl = [], defaultdict(list)
    for i, a in enumerate(corpus):
        if tA and str(i) not in tA:       # capped deals: only present teacher-scoped atoms
            continue
        if i in _semdrop:                 # prose twin kept; drop the table-row dup
            continue
        if is_junk(a["body"]):            # drop xlsx-internal / OCR-artifact junk
            continue
        leaf = (a["section"].split(" > ")[-1] if a.get("section") else "") or "(top)"
        tl = tA.get(str(i), {})
        xc = [] if _FAST else (context(i, corpus, ents, types, simfn, kind=kind) if i < 400 else [])
        if a["doc"] not in order_doc: order_doc.append(a["doc"])
        if leaf not in order_tbl[a["doc"]]: order_tbl[a["doc"]].append(leaf)
        _section = a.get("section", "") or ""
        by_doc[a["doc"]][leaf].append({"i": i, "t": tl.get("fine") or a["type"], "b": clean_body(a["body"])[:2000],
                                       "pg": a.get("page", 1),
                                       # sheet = sheet NAME (first path segment) for xlsx sheet-sync
                                       "sheet": (_section.split(" > ")[0] if _section else leaf),
                                       # path = FULL section path (sheet > title > ...) shown as the atom's breadcrumb
                                       "path": _section or leaf,
                                       # intro = governing lead-in (connective tissue lifted onto this atom)
                                       "intro": a.get("intro", ""),
                                       "f": tl.get("flag", "") if tl.get("flag") != "none" else "",
                                       "x": xc, "fc": {k: tl.get(k, "") for k in FK}, "src": "teacher" if tl else "heuristic"})
    docs = []
    for d in order_doc:
        # Render in true DOCUMENT order (the corpus index `i` is doc-ordered), not
        # grouped-by-section. Grouping pulled a trailing paragraph back up to where
        # its parent section first opened — so an exclusion blob written AFTER a
        # bulleted list rendered BEFORE it. Iterating by `i` keeps doc order; the
        # render still draws a section divider whenever the path changes (a reopened
        # parent just shows its divider again, which is faithful).
        rows_all = sorted((r for t in order_tbl[d] for r in by_doc[d][t]), key=lambda r: r["i"])
        schem = _SCHEM_NAME.search(d) or (rows_all and sum(1 for r in rows_all if _SCHEM.search(r["b"])) / len(rows_all) > 0.30)
        if schem:   # CAD drawing set / site schematic — not supported yet
            continue
        pf = pdf_for(d); rp = pathmap.get(d, "")
        docs.append({"doc": d, "pdf": pf, "dl": dl_map.get(d, ""),
                     # real source filename, so the download saves with the right
                     # name (the bare `download` attr can't derive a name from an
                     # inlined data: URL — every file would save as "download.xlsx").
                     "fn": Path(dl_map.get(d, "") or d).name,
                     "sheets": [] if pf else render_xlsx_sheets(rp),
                     "srchtml": "" if pf else render_docx_html(rp),
                     "tables": [{"name": t, "rows": by_doc[d][t]} for t in order_tbl[d]]})
    if not docs:
        print(f"skip {slug}: all docs were schematics/empty"); continue
    edges = [{"i": e["i"], "j": e["j"], "rel": e["rel"], "a": corpus[e["i"]]["body"][:130], "b": corpus[e["j"]]["body"][:130],
             "da": corpus[e["i"]]["doc"][:20], "db": corpus[e["j"]]["doc"][:20]}
            for e in tE if e["i"] < len(corpus) and e["j"] < len(corpus)]
    deals[slug] = {"docs": docs, "edges": edges,
                   "pdfs": [{"name": Path(rel).name, "path": rel} for _, rel in pdf_files]}
    print(f"{slug}: {sum(len(t['rows']) for d in docs for t in d['tables'])} atoms, {len(edges)} edges, {len(pdf_files)} viewable PDFs")

DATA = json.dumps({"deals": deals, "coarse": COARSE, "facets": FACETS, "admission": ADMISSION,
                   "rels": ["corroborates", "conflict", "partition_key", "no_edge"]}, ensure_ascii=False)

HTML = r"""<!doctype html><html><head><meta charset=utf-8><title>parser-os labeler</title>
<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
<style>
:root{--bg:#0e1014;--pan:#161922;--bd:#262b37;--tx:#e7eaf0;--mut:#8b93a7;--ac:#5b9dff;--ok:#3ddc84;--warn:#ffb454;--bad:#ff6b6b;--rec:#ff8fcf}
*{box-sizing:border-box}body{margin:0;font:13px/1.45 system-ui;background:var(--bg);color:var(--tx)}
.top{position:sticky;top:0;z-index:5;background:var(--pan);border-bottom:1px solid var(--bd);padding:8px 14px;display:flex;gap:10px;align-items:center;flex-wrap:wrap}
.top b{font-size:14px}.tab,.stg{padding:3px 11px;border:1px solid var(--bd);border-radius:14px;cursor:pointer;color:var(--mut);font-size:12px;user-select:none}
.tab.on{background:var(--rec);color:#111;border-color:var(--rec)}.stg.on{background:var(--ac);color:#fff;border-color:var(--ac)}
.btn{padding:4px 12px;border:1px solid var(--bd);border-radius:7px;background:var(--pan);color:var(--ac);cursor:pointer;font-size:12px}
.btn.save{background:var(--ok);color:#08130c;border:0;font-weight:600}
.wrap{display:flex;height:calc(100vh - 49px)}
.left{width:48%;border-right:1px solid var(--bd);display:flex;flex-direction:column;background:#0a0c10}
.lnav{padding:6px 10px;display:flex;gap:8px;align-items:center;border-bottom:1px solid var(--bd);background:var(--pan)}
.lnav select{flex:1;background:var(--pan);color:var(--tx);border:1px solid var(--bd);border-radius:6px;padding:3px}
#view{flex:1;overflow:auto;position:relative;padding:10px;display:flex;flex-direction:column;align-items:center}
#pgwrap{position:relative}#pglayer{position:absolute;top:0;left:0;color:transparent}
#pglayer span{position:absolute;white-space:pre;transform-origin:0 0;cursor:text}
#pglayer .hl{background:rgba(255,212,84,.45);color:transparent;border-radius:2px}
.right{flex:1;overflow:auto;padding:12px}
h4{margin:14px 0 4px;color:var(--mut);font-size:12px;text-transform:uppercase;letter-spacing:.04em}
.atom{border:1px solid var(--bd);border-radius:9px;padding:9px 10px;margin:7px 0;background:var(--pan);transition:border .15s}
.atom.sel{border-color:var(--ac);box-shadow:0 0 0 1px var(--ac)}
.atom.acc{border-left:3px solid var(--ok)}
.acpt{font-size:11px;color:var(--ok);border:1px solid var(--ok);border-radius:6px;padding:2px 10px;cursor:pointer;background:none;font-weight:600}
.acpt.done{background:var(--ok);color:#08130c}
.aty{font-size:11px;color:var(--rec);border:1px solid var(--rec);border-radius:5px;padding:0 6px}
.fl{color:var(--warn);font-size:11px;margin-left:6px}.tb{font-size:9px;color:var(--rec);border:1px solid var(--rec);border-radius:4px;padding:0 4px;margin-left:5px}
.pg{font-size:10px;color:var(--mut);margin-left:6px}
.bc{font-size:10px;color:var(--ac);opacity:.85;margin-bottom:3px;font-family:ui-monospace;letter-spacing:.2px}
.intro{font-size:10px;color:var(--ac);opacity:.7;margin-bottom:4px;font-style:italic;padding-left:6px;border-left:2px solid var(--ac)}
.shtabs{display:flex;flex-wrap:wrap;gap:4px;margin:7px 0}
.shtab{font-size:11px;padding:2px 9px;border:1px solid var(--bd);border-radius:5px;cursor:pointer;color:var(--mut);white-space:nowrap}
.shtab.on{background:var(--ac);color:#fff;border-color:var(--ac)}
.shtab:hover{border-color:var(--ac)}
.muted2{color:var(--mut);font-size:11px}
.secdiv{position:sticky;top:0;background:var(--bg);color:var(--ac);font-size:12px;font-weight:600;padding:6px 4px;margin:10px 0 2px;border-bottom:1px solid var(--bd);z-index:1}
.atext{font-family:ui-monospace,monospace;font-size:12px;margin:6px 0;white-space:pre-wrap;cursor:pointer}
.clink{font-size:11px;color:var(--mut);font-family:ui-monospace}.clink b{color:var(--ac)}.cf{color:var(--warn)}
.row{display:flex;flex-wrap:wrap;gap:6px;align-items:center;margin-top:6px}
.esel,.rel{font-size:11px;background:var(--bg);color:var(--tx);border:1px solid var(--bd);border-radius:6px;padding:2px 4px}
.pi{font-size:11px;color:var(--bad);border:1px solid var(--bad);border-radius:6px;padding:2px 9px;cursor:pointer;background:none}
.pi.on{background:var(--bad);color:#fff}
.edge{border:1px solid var(--bd);border-radius:9px;padding:9px;margin:7px 0;background:var(--pan)}.ea{font-family:ui-monospace;font-size:12px;line-height:1.5}
.muted{color:var(--mut);padding:20px;text-align:center}
.srchead{font-size:11px;color:var(--mut);padding:8px 12px;border-bottom:1px solid var(--bd);background:var(--pan);position:sticky;top:0;z-index:2}
.dlbtn{display:inline-block;background:var(--ok);color:#08130c;font-weight:700;font-size:13px;padding:6px 14px;border-radius:7px;text-decoration:none}
.dlbtn:hover{filter:brightness(1.1)}
.srcwrap{padding:6px 8px}
.shsheet{font-weight:600;color:var(--ac);padding:10px 4px 3px;font-size:12px}
.sgrid{border-collapse:collapse;font-size:12px;margin-bottom:12px}
.sgrid td{border:1px solid var(--bd);padding:3px 9px;vertical-align:top;max-width:320px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.sgrid td:hover{overflow:visible;white-space:normal;word-break:break-word;background:var(--pan)}
.sgrid tr:nth-child(even) td:not([style*=background]){background:rgba(255,255,255,.02)}
.sp{padding:4px 6px;font-size:13px;line-height:1.55;margin:0}
</style></head><body>
<div class=top><b>parser-os labeler</b><span id=tabs style="display:flex;gap:6px"></span>
<span id=stages style="display:flex;gap:6px"></span><span style=flex:1></span>
<span id=sum class=muted style=padding:0></span><button class=btn onclick=acceptAll()>✓ accept all shown</button><button class="btn save" onclick=dl()>⬇ download gold</button></div>
<div class=wrap>
 <div class=left><div class=lnav><select id=docsel onchange=setDoc(this.value)></select>
   <button class=btn onclick=pg(-1)>◀</button><span id=pgnum class=muted style=padding:0>—</span><button class=btn onclick=pg(1)>▶</button>
   <label style="font-size:11px;color:var(--mut);display:flex;align-items:center;gap:4px"><input type=checkbox id=pgonly checked onchange="pageOnly=this.checked;render()"> page only</label></div>
   <div id=view><div class=muted>pick an atom to view its source</div></div></div>
 <div class=right id=main></div>
</div>
<script id=ld type=application/json>__DATA__</script>
<script>
pdfjsLib.GlobalWorkerOptions.workerSrc="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js";
const R=JSON.parse(ld.textContent),D=R.deals,C=R.coarse,FAC=R.facets,AD=R.admission,REL=R.rels,DEALS=Object.keys(D);
const KEY="parseros_gold_v2";let store=JSON.parse(localStorage.getItem(KEY)||"{}");
let deal=DEALS[0],stage="② type",curDoc=null,pdf=null,page=1,pending=null,mode="doc",sheetI=0,pageOnly=true,selI=null;
const $=i=>document.getElementById(i),esc=s=>(s||"").replace(/&/g,"&amp;").replace(/</g,"&lt;").replace(/"/g,"&quot;");
const K=i=>deal+"::"+i,EK=(i,j)=>deal+"::e"+i+"_"+j;
function save(){localStorage.setItem(KEY,JSON.stringify(store));upd();}     // every change -> refresh-safe
function upd(){let tot=0;D[deal].docs.forEach(dc=>dc.tables.forEach(t=>tot+=t.rows.length));
 let acc=Object.keys(store).filter(k=>k.startsWith(deal+"::")&&k.indexOf("::e")<0&&store[k].label).length;
 $("sum").textContent=acc+" / "+tot+" labeled ("+(tot?Math.round(100*acc/tot):0)+"%)";}
const STG=[["all",a=>1],["① parser",a=>a.f],["② type",a=>1],["③ conflicts",a=>(a.x||[]).some(c=>c.cf)],["④ edges",null],["⑤ todo",a=>!(store[K(a.i)]&&store[K(a.i)].label)]];
function docOf(i){for(const dc of D[deal].docs)for(const t of dc.tables)for(const r of t.rows)if(r.i==i)return dc;}
function rowOf(i){for(const dc of D[deal].docs)for(const t of dc.tables)for(const r of t.rows)if(r.i==i)return r;}

// ---- PDF (left) ----
async function openDoc(path){pdf=null;mode='pdf';
 if(!path){$("view").innerHTML='<div class=muted>no PDF for this doc — label from the text on the right.</div>';return;}
 try{pdf=await pdfjsLib.getDocument(path).promise;page=1;await drawPage();}catch(e){$("view").innerHTML='<div class=muted>PDF could not load. Label from the text on the right.</div>';}}
async function drawPage(hlText){if(!pdf)return;page=Math.max(1,Math.min(page,pdf.numPages));$("pgnum").textContent="p"+page+"/"+pdf.numPages;
 const p=await pdf.getPage(page),base=p.getViewport({scale:1}),W=($("view").clientWidth||720)-26,vp=p.getViewport({scale:Math.max(0.5,Math.min(2.2,W/base.width))});
 $("view").innerHTML='<div id=pgwrap><canvas id=cv></canvas><div id=pglayer></div></div>';
 const cv=$("cv");cv.width=vp.width;cv.height=vp.height;$("pgwrap").style.width=vp.width+"px";$("pgwrap").style.height=vp.height+"px";
 await p.render({canvasContext:cv.getContext("2d"),viewport:vp}).promise;
 const tc=await p.getTextContent(),L=$("pglayer");L.style.width=vp.width+"px";L.style.height=vp.height+"px";
 const norm=s=>(s||"").toLowerCase().replace(/[^a-z0-9]/g,"");
 // pass 1: find the atom's CONTIGUOUS region in the page text (precise, not scattered tokens)
 let full="",map=[];tc.items.forEach((it,idx)=>{const n=norm(it.str);if(n){map.push({s:full.length,e:full.length+n.length,idx});full+=n;}});
 const tgt=norm(hlText),hit=new Set();
 if(tgt.length>=10){let probe=tgt.slice(0,45),pos=full.indexOf(probe);
   while(pos<0&&probe.length>14){probe=probe.slice(0,probe.length-6);pos=full.indexOf(probe);}
   if(pos>=0){const e=pos+probe.length;map.forEach(m=>{if(m.s<e&&m.e>pos)hit.add(m.idx);});}}
 // pass 2: render text layer, highlight only the matched contiguous run
 tc.items.forEach((it,idx)=>{const tx=pdfjsLib.Util.transform(vp.transform,it.transform),sp=document.createElement("span");
  sp.textContent=it.str;sp.style.left=tx[4]+"px";sp.style.top=(tx[5]-it.height*vp.scale)+"px";sp.style.fontSize=(it.height*vp.scale)+"px";
  if(hit.has(idx))sp.className="hl";L.appendChild(sp);});
 const h=L.querySelector(".hl");if(h)h.scrollIntoView({block:"center"});
 if(pageOnly)renderMain();}
function pg(d){if(mode=='pdf'){page+=d;drawPage();return;}
 if(mode=='sheet'){const dc=D[deal].docs.find(x=>x.doc==curDoc);if(!dc||!dc.sheets)return;sheetI=Math.max(0,Math.min(dc.sheets.length-1,sheetI+d));renderSheet(dc);}}
async function focusAtom(r){document.querySelectorAll(".atom.sel").forEach(e=>e.classList.remove("sel"));
 selI=r.i;const el=$("a"+r.i);if(el)el.classList.add("sel");const dc=docOf(r.i),pf=dc&&dc.pdf;
 if(pf){if(!pdf)await openDoc(pf);page=r.pg||1;await drawPage(r.b);}
 else if(!pdf)$("view").innerHTML='<div class=muted>this atom is from a Word/Excel doc (no PDF) — label from the text on the right. Pick a PDF in the dropdown to browse it.</div>';
 /* else: a PDF is already open and this atom isn\x27t in it — keep the PDF up, just select the atom */ }

// ---- atom + edge editors (right) ----
function opt(v,c){return "<option"+(v===c?" selected":"")+">"+v+"</option>";}
function tsel(c){let h='<select class=esel data-f=__type__><optgroup label=admission>';AD.filter(x=>x!="keep").forEach(t=>h+=opt(t,c));h+="</optgroup>";Object.keys(C).forEach(f=>{h+="<optgroup label='"+f.replace(/_/g,' ')+"'>";C[f].forEach(t=>h+=opt(t,c));h+="</optgroup>";});return h+"</select>";}
function fsel(d,c){let h='<select class=esel data-f='+d+'>';FAC[d].forEach(v=>h+=opt(v,c));return h+"</select>";}
function atomHTML(a){const e=store[K(a.i)]||{},lab=e.label||a.t,pi=e.parser_issue,acc=!!e.label;
 let h='<div class="atom'+(acc?' acc':'')+'" id=a'+a.i+' data-i='+a.i+'>'+(a.path?'<div class=bc title="section / title this atom belongs to">'+esc(a.path)+'</div>':'')+(a.intro?'<div class=intro title="connective lead-in this atom belongs under (lifted from the doc)">↳ '+esc(a.intro)+'</div>':'')+'<div><span class=aty>'+lab+'</span>'+(a.f?'<span class=fl>['+a.f+']</span>':'')+(a.src=="teacher"?'<span class=tb>T</span>':'')+'<span class=pg>p'+a.pg+'</span></div>';
 h+='<div class=atext onclick=focusRow('+a.i+')>'+esc(a.b)+'</div>';
 if((a.x||[]).length)h+='<div style=margin:3px 0>'+a.x.map(c=>'<div class=clink><b>'+esc(c.d)+'</b> '+esc(c.sh)+(c.cf?' <span class=cf>⚠'+esc(c.cf)+'</span>':'')+'</div>').join('')+'</div>';
 h+='<div class=row><button class="acpt'+(acc?' done':'')+'" onclick="accept('+a.i+')">'+(acc?'✓ saved':'✓ accept')+'</button> type '+tsel(lab);Object.keys(FAC).forEach(d=>h+=' '+(d=="metadata_kind"?"meta":d.split("_")[0])+' '+fsel(d,(e.facets&&e.facets[d])||a.fc[d]));
 h+='<button class="pi'+(pi?' on':'')+'" onclick="pissue('+a.i+')">⚠ parser issue</button></div></div>';return h;}
function edgeHTML(e){const s=store[EK(e.i,e.j)]||{},r=s.rel||e.rel;
 let h='<div class=edge><div class=ea><b>'+esc(e.da)+'</b>: '+esc(e.a)+'<br><b>'+esc(e.db)+'</b>: '+esc(e.b)+'</div><div class=row>relation <select class=rel data-ei='+e.i+' data-ej='+e.j+'>';
 REL.forEach(v=>h+=opt(v,r));return h+'</select></div></div>';}
function focusRow(i){focusAtom(rowOf(i));}
function accept(i){const A=$("a"+i);if(!A)return;const e={label:"",facets:{}};A.querySelectorAll(".esel").forEach(s=>{if(s.dataset.f=="__type__")e.label=s.value;else e.facets[s.dataset.f]=s.value;});store[K(i)]={...store[K(i)],...e};save();const b=A.querySelector(".acpt");if(b){b.classList.add("done");b.textContent="✓ saved";}A.classList.add("acc");}
function acceptAll(){document.querySelectorAll("#main .atom").forEach(A=>accept(+A.dataset.i));}
document.addEventListener("keydown",e=>{if(e.key=="a"){const s=document.querySelector(".atom.sel");if(s){accept(+s.dataset.i);e.preventDefault();}}});
function renderMain(){let H="";const dd=D[deal];
 if(stage=="④ edges"){H=dd.edges.length?'<h4>cross-context edges ('+dd.edges.length+') — accept/fix relation</h4>'+dd.edges.map(edgeHTML).join(''):'<div class=muted>no edges</div>';}
 else{const f=STG.find(s=>s[0]==stage)[1];const dc=dd.docs.find(d=>d.doc==curDoc)||dd.docs[0];
   let rows=[];(dc?dc.tables:[]).forEach(t=>t.rows.forEach(r=>{if(f(r))rows.push(r);}));
   rows.sort((a,b)=>(a.i||0)-(b.i||0));  /* DOCUMENT order: tables are section-grouped, but a reopened parent section must not pull a trailing paragraph above an earlier sub-section's bullets */
   let scope='';
   if(mode=='pdf'&&pdf&&pageOnly){rows=rows.filter(r=>(r.pg||1)==page);scope=' · page '+page+' only';}
   if(mode=='sheet'&&pageOnly&&dc&&dc.sheets&&dc.sheets[sheetI]){const sn=dc.sheets[sheetI].n;rows=rows.filter(r=>r.sheet==sn);scope=' · sheet: '+sn;}
   let body='',prevp=null;
   rows.forEach(r=>{if(r.path!==prevp){body+='<div class=secdiv>'+esc(r.path||'(no section)')+'</div>';prevp=r.path;}body+=atomHTML(r);});
   H=(dc?'<h4>'+esc(dc.doc)+(dc.pdf?'  · PDF':'  · Word/Excel')+esc(scope)+'</h4>':'')+(rows.length?body:'<div class=muted>no atoms on this '+(mode=='sheet'?'sheet':mode=='pdf'&&pageOnly?'page':'doc')+' in this stage ✓ — flip with ◀ ▶</div>');}
 $("main").innerHTML=H;if(selI!=null){const e=$("a"+selI);if(e)e.classList.add("sel");}upd();}
function render(){renderMain();
 $("tabs").innerHTML=DEALS.map(s=>'<span class="tab'+(s==deal?' on':'')+'" onclick="deal=\''+s+'\';fillDocs();render()">'+s+'</span>').join('');
 $("stages").innerHTML=STG.map(s=>'<span class="stg'+(s[0]==stage?' on':'')+'" onclick="stage=\''+s[0]+'\';render()">'+s[0]+'</span>').join('');}
function setSheet(k){const dc=D[deal].docs.find(x=>x.doc==curDoc);if(!dc||!dc.sheets)return;sheetI=Math.max(0,Math.min(dc.sheets.length-1,k));renderSheet(dc);}
function renderSheet(dc){const s=dc.sheets[sheetI];$("pgnum").textContent='sheet '+(sheetI+1)+'/'+dc.sheets.length;
 const btn=dc.dl?'<a class=dlbtn href="'+dc.dl+'" download="'+esc(dc.fn||'deal-kit.xlsx')+'">⬇ Download &amp; open the real Excel file</a>':'';
 const tabs='<div class=shtabs>'+dc.sheets.map((x,k)=>'<span class="shtab'+(k==sheetI?' on':'')+'" onclick="setSheet('+k+')">'+esc(x.n)+'</span>').join('')+'</div>';
 $("view").innerHTML='<div class=srchead>'+btn+tabs+'<div style="margin-top:5px" class=muted2>sheet <b>'+esc(s.n)+'</b> ('+(sheetI+1)+' of '+dc.sheets.length+') — click a tab above or use ◀ ▶. raw values; open the file for real formatting.</div></div><div class=srcwrap>'+s.h+'</div>';renderMain();}
function showLeft(dc){if(dc&&dc.pdf){mode='pdf';openDoc(dc.pdf);return;}
 pdf=null;const dl=(dc&&dc.dl)||'';const fn=(dc&&dc.fn)||'document';
 if(dc&&dc.sheets&&dc.sheets.length){mode='sheet';sheetI=0;renderSheet(dc);return;}
 mode='doc';$("pgnum").textContent='—';const h=(dc&&dc.srchtml)||'';
 const btn=dl?'<a class=dlbtn href="'+dl+'" download="'+esc(fn)+'">⬇ Download &amp; open the real file (Word)</a>':'';
 $("view").innerHTML='<div class=srchead>'+btn+'<div style="margin-top:5px">↓ quick preview (raw values — open the file above for the real formatted view)</div></div>'+(h?'<div class=srcwrap>'+h+'</div>':'<div class=muted>(no preview — download above)</div>');}
function setDoc(name){curDoc=name;showLeft(D[deal].docs.find(d=>d.doc==name));render();}
function fillDocs(){const ds=D[deal].docs;curDoc=ds.length?ds[0].doc:null;pdf=null;
 $("docsel").innerHTML=ds.map(d=>'<option value="'+esc(d.doc)+'">'+esc(d.doc)+(d.pdf?'  · PDF':'  · Word/Excel')+'</option>').join('');
 showLeft(ds[0]);}
$("main").addEventListener("change",ev=>{const A=ev.target.closest("[data-i]");
 if(A&&ev.target.classList.contains("esel")){const i=+A.dataset.i,e=store[K(i)]||{label:"",facets:{}};A.querySelectorAll(".esel").forEach(s=>{if(s.dataset.f=="__type__")e.label=s.value;else e.facets[s.dataset.f]=s.value;});store[K(i)]={...store[K(i)],...e};save();}
 if(ev.target.classList.contains("rel")){const t=ev.target;store[EK(+t.dataset.ei,+t.dataset.ej)]={...store[EK(+t.dataset.ei,+t.dataset.ej)],rel:t.value};save();}});
function pissue(i){const r=rowOf(i),e=store[K(i)]||{label:"",facets:{}};
 const note=prompt("Parser issue for Griffin (what's wrong with how this was extracted?)",(e.parser_issue&&e.parser_issue.note)||"");
 if(note===null)return;e.parser_issue={note:note,doc:docOf(i)?docOf(i).doc:"",page:r.pg};store[K(i)]=e;save();render();
 const body="Deal: "+deal+"%0ADoc: "+(e.parser_issue.doc)+" (page "+r.pg+")%0AAtom: "+encodeURIComponent(r.b.slice(0,240))+"%0A%0AIssue: "+encodeURIComponent(note);
 window.location.href="mailto:griffin@optbotai.com?subject="+encodeURIComponent("Parser issue: "+deal+"/"+e.parser_issue.doc)+"&body="+body;}
function dl(){const out={atoms:[],edges:[],parser_issues:[]};DEALS.forEach(s=>{D[s].docs.forEach(dc=>dc.tables.forEach(t=>t.rows.forEach(r=>{const e=store[s+"::"+r.i];if(!e)return;if(e.label)out.atoms.push({deal:s,doc:dc.doc,page:r.pg,atom:r.b,label:e.label,facets:e.facets||{},teacher:r.t});if(e.parser_issue)out.parser_issues.push({deal:s,...e.parser_issue,atom:r.b});})));D[s].edges.forEach(e=>{const g=store[s+"::e"+e.i+"_"+e.j];if(g&&g.rel)out.edges.push({deal:s,a:e.a,b:e.b,rel:g.rel,teacher:e.rel});});});
 const bl=new Blob([JSON.stringify(out,null,1)],{type:"application/json"}),u=URL.createObjectURL(bl),x=document.createElement("a");x.href=u;x.download="gold_"+DEALS.join("_").slice(0,40)+".json";x.click();}
fillDocs();render();
</script></body></html>"""
(OUT / "labeler.html").write_text(HTML.replace("__DATA__", DATA), encoding="utf-8")
zf = OUT / "batch.zip"
with zipfile.ZipFile(zf, "w", zipfile.ZIP_DEFLATED) as z:
    for f in OUT.rglob("*"):
        if f.is_file() and f.name != "batch.zip":
            z.write(f, f.relative_to(OUT))
print(f"WROTE {zf} ({os.path.getsize(zf)/1e6:.2f} MB) · {len(deals)} deals")
