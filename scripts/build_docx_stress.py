"""DOCX adversarial bundle — managed-services SOW shapes."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def da_simple_sow(out: Path) -> Path:
    p = out / "da_simple_sow.docx"
    doc = Document()
    doc.add_heading("Statement of Work — Acme 2026 Refresh", level=1)
    doc.add_heading("1. Scope of Work", level=2)
    doc.add_paragraph(
        "PurTera will furnish and install 50 wireless access points at ATL-HQ-01 "
        "and 30 access points at ATL-WEST-02 by Mar 15, 2026."
    )
    doc.add_heading("2. Pricing", level=2)
    doc.add_paragraph("Total contract value: USD $245,000.00 fixed-price.")
    doc.add_heading("3. Acceptance", level=2)
    doc.add_paragraph(
        "Final acceptance (ATP) by June 30, 2026 with 30-day stabilization."
    )
    doc.save(p)
    return p


def db_table_in_docx(out: Path) -> Path:
    p = out / "db_table_in_docx.docx"
    doc = Document()
    doc.add_heading("SOW with Embedded BOM Table", level=1)
    doc.add_paragraph("Hardware breakdown:")
    table = doc.add_table(rows=4, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Site ID"
    hdr[1].text = "Part Number"
    hdr[2].text = "Description"
    hdr[3].text = "Qty"
    hdr[4].text = "Unit Price"
    rows = [
        ["ATL-HQ-01",   "WAP-9180AX-K9", "Wi-Fi 6E AP",   "50", "$1,200"],
        ["ATL-WEST-02", "WAP-9180AX-K9", "Wi-Fi 6E AP",   "30", "$1,200"],
        ["ATL-HQ-01",   "C9300-48P-A",   "48-port switch", "5", "$3,500"],
    ]
    for i, r in enumerate(rows, 1):
        for j, val in enumerate(r):
            table.rows[i].cells[j].text = val
    doc.save(p)
    return p


def dc_bullet_list_scope(out: Path) -> Path:
    p = out / "dc_bullet_list_scope.docx"
    doc = Document()
    doc.add_heading("Scope of Work (Bulleted)", level=1)
    for line in [
        "Install 50 wireless access points at ATL-HQ-01.",
        "Install 30 wireless access points at ATL-WEST-02.",
        "Install 12 SFP-10G-LR-S= modules across ATL-AIR-03.",
        "Provide 12 months of Silver-tier managed service post-cutover.",
        "Response SLA: P1 within 2 hours, P2 within 4 business hours.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.save(p)
    return p


def dd_multi_section_msa(out: Path) -> Path:
    p = out / "dd_multi_section_msa.docx"
    doc = Document()
    doc.add_heading("Master Services Agreement", level=1)
    for h, body in [
        ("Section 12: Indemnification", "Each party shall indemnify, defend, and hold harmless the other party..."),
        ("Section 13: Force Majeure", "Neither party shall be liable for failures arising from acts of God..."),
        ("Section 14: Confidentiality", "All non-public information shall be treated as confidential for 5 years."),
        ("Section 15: Term & Termination", "Initial term is 36 months with 90-day renewal."),
        ("Section 16: Governing Law", "Georgia state law governs this Agreement."),
    ]:
        doc.add_heading(h, level=2)
        doc.add_paragraph(body)
    doc.save(p)
    return p


def de_signature_block(out: Path) -> Path:
    p = out / "de_signature_block.docx"
    doc = Document()
    doc.add_heading("Approvals", level=1)
    doc.add_paragraph("OPTBOT — Director of Workplace Technology: Jane Roe")
    doc.add_paragraph("OPTBOT — VP IT Infrastructure: John Smith")
    doc.add_paragraph("PurTera — Program Manager: Maria Lopez")
    doc.add_paragraph("PurTera — Solutions Architect: T. Nguyen")
    doc.save(p)
    return p


def df_site_roster_docx(out: Path) -> Path:
    p = out / "df_site_roster_docx.docx"
    doc = Document()
    doc.add_heading("Authoritative Site Roster", level=1)
    doc.add_paragraph("kind=physical_site for all rows below.")
    table = doc.add_table(rows=6, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Site ID", "Facility name", "Street address", "MDF / IDF", "Access window", "Escort owner"]):
        hdr[i].text = h
    rows = [
        ["ATL-HQ-01",   "OPTBOT Atlanta HQ",        "1200 Peachtree St NE, Atlanta GA 30309",     "MDF-3A / IDF 2-7",  "Mon-Fri 07:00-18:00", "OPTBOT Facilities"],
        ["ATL-WEST-02", "OPTBOT West Campus",       "3100 Interstate N Pkwy, Atlanta GA 30339",   "MDF-W1 / IDF W2-3", "Mon-Fri 07:00-18:00", "OPTBOT Facilities"],
        ["ATL-AIR-03",  "OPTBOT Airport Logistics", "6000 N Terminal Pkwy, Atlanta GA 30320",     "MDF-A / IDF A1",     "Mon-Sat 06:00-22:00", "OPTBOT Security"],
        ["ATL-047-04",  "OPTBOT Brady Training",    "047 Brady Ave NW, Atlanta GA 30318",         "MDF-B / IDF B1-2",   "Mon-Fri 08:00-17:00", "OPTBOT Facilities"],
        ["ATL-CP-05",   "OPTBOT College Park",      "1850 Sullivan Rd, College Park GA 30337",    "MDF-CP / staging",   "Mon-Fri 07:00-15:00", "OPTBOT Logistics"],
    ]
    for i, r in enumerate(rows, 1):
        for j, val in enumerate(r):
            table.rows[i].cells[j].text = val
    doc.save(p)
    return p


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: python scripts/build_docx_stress.py <out_dir>", file=sys.stderr)
        return 2
    out = Path(sys.argv[1]).resolve() / "artifacts"
    out.mkdir(parents=True, exist_ok=True)
    for b in [da_simple_sow, db_table_in_docx, dc_bullet_list_scope, dd_multi_section_msa, de_signature_block, df_site_roster_docx]:
        p = b(out)
        print(f"  -> {p.name}")
    print(f"\n6 DOCX files in {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
