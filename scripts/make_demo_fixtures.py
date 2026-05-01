from __future__ import annotations

import json
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from openpyxl import Workbook

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": WORD_NS}


def _inject_deleted_tracked_change(docx_path: Path, deleted_text: str) -> None:
    # python-docx does not expose tracked-change APIs directly, so fixture generation
    # inserts a minimal w:del block into word/document.xml for deterministic tests.
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        with zipfile.ZipFile(docx_path, "r") as source:
            source.extractall(temp_root)

        doc_xml = temp_root / "word" / "document.xml"
        root = ET.fromstring(doc_xml.read_bytes())
        body = root.find(".//w:body", NSMAP)
        if body is None:
            raise ValueError("Invalid DOCX structure: missing word body")

        p = ET.Element(f"{{{WORD_NS}}}p")
        r = ET.SubElement(p, f"{{{WORD_NS}}}r")
        del_node = ET.SubElement(r, f"{{{WORD_NS}}}del")
        del_run = ET.SubElement(del_node, f"{{{WORD_NS}}}r")
        del_text_node = ET.SubElement(del_run, f"{{{WORD_NS}}}delText")
        del_text_node.text = deleted_text
        body.append(p)

        doc_xml.write_bytes(
            ET.tostring(root, encoding="utf-8", xml_declaration=True),
        )

        with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as target:
            for file in temp_root.rglob("*"):
                if file.is_file():
                    target.write(file, file.relative_to(temp_root).as_posix())


def create_demo_project(base_dir: Path) -> Path:
    project_dir = base_dir / "tests" / "fixtures" / "demo_project"
    expected_dir = base_dir / "tests" / "fixtures" / "expected"
    project_dir.mkdir(parents=True, exist_ok=True)
    expected_dir.mkdir(parents=True, exist_ok=True)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "site_roster"
    sheet.append(["Site", "Floor", "Hostname", "Quantity", "Access Window", "Scope"])
    sheet.append(["Main Campus", "1", "IP Camera", "50", "Weekdays 8am-5pm", "Install"])
    sheet.append(["West Wing", "2", "IP Camera", "41", "Escort required", "Install"])
    sheet.append(["TOTAL", "", "", "91", "", ""])
    workbook.save(project_dir / "site_list.xlsx")

    quote_workbook = Workbook()
    quote_sheet = quote_workbook.active
    quote_sheet.title = "quote"
    quote_sheet.append(["Part Number", "Description", "Quantity", "Unit Price", "Lead Time"])
    quote_sheet.append(["CAM-IP-001", "IP Camera", "72", "300.00", "2 weeks"])
    quote_workbook.save(project_dir / "vendor_quote.xlsx")

    email_thread = (
        "From: jane.customer@example.com\n"
        "Sent: 2026-01-15 09:00\n"
        "Subject: Scope update\n"
        "\n"
        "Please remove West Wing from scope. Main Campus requires escort access after 5pm.\n"
        "\n"
        "On 2026-01-10, Jane Customer wrote:\n"
        "> Please include West Wing in the camera rollout.\n"
    )
    (project_dir / "customer_email.txt").write_text(email_thread, encoding="utf-8")

    kickoff_transcript = (
        "[00:00:01] Purtera PM: Starting kickoff for the camera rollout.\n"
        "[00:01:12] Jane Customer: Main Campus requires escort access after 5pm.\n"
        "[00:02:05] Jane Customer: Please remove West Wing from scope for now.\n"
        "[00:03:10] Purtera PM: Confirmed, West Wing will be treated as excluded pending written confirmation.\n"
        "[00:04:22] Bob Customer: Customer will provide lift access for lobby camera installation.\n"
        "[00:05:30] Purtera PM: Open question: confirm whether MDF room requires badge access.\n"
        "[00:06:15] Jane Customer: We may add 5 more IP cameras at Main Campus after the pilot.\n"
    )
    (project_dir / "kickoff_transcript.txt").write_text(kickoff_transcript, encoding="utf-8")

    sow_doc = Document()
    sow_doc.add_heading("Statement of Work Draft", level=1)
    sow_doc.add_paragraph("Scope includes installation of IP Cameras at Main Campus and West Wing.")
    sow_doc.add_paragraph("AV displays are excluded from scope.")
    sow_doc.add_paragraph("Customer is responsible for providing lift access.")
    sow_path = project_dir / "sow_draft.docx"
    sow_doc.save(sow_path)
    _inject_deleted_tracked_change(sow_path, "Install AV displays in conference rooms.")

    expected_summary = {
        "packet_families": [
            "quantity_conflict",
            "vendor_mismatch",
            "scope_exclusion",
            "site_access",
            "scope_inclusion",
            "meeting_decision",
            "action_item",
            "missing_info",
        ],
        "anchors": {
            "quantity_conflict": ["device:ip_camera"],
            "vendor_mismatch": ["device:ip_camera"],
            "scope_exclusion": ["site:west_wing", "site:unknown"],
            "site_access": ["site:main_campus", "site:west_wing", "unknown"],
            "scope_inclusion": ["device:unknown"],
            "meeting_decision": ["meeting_decision:confirmed_west_wing_will_be_treated_as_excluded_pending_written_confirmation"],
            "action_item": [
                "action_item:customer:customer_will_provide_lift_access_for_lobby_camera_installation",
                "action_item:internal:open_question_confirm_whether_mdf_room_requires_badge_access",
            ],
            "missing_info": [
                "missing_info:open_question_confirm_whether_mdf_room_requires_badge_access",
                "missing_info:confirmed_west_wing_will_be_treated_as_excluded_pending_written_confirmation",
            ],
        },
        "notes": [
            "IDs are deterministic but content-sensitive; test families and anchors instead of hard-coded IDs."
        ],
    }
    (expected_dir / "demo_summary.json").write_text(
        json.dumps(expected_summary, indent=2),
        encoding="utf-8",
    )

    return project_dir


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    created = create_demo_project(root)
    print(f"Demo fixtures ready at: {created}")
